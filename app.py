import base64
import os
import fitz
import pandas as pd
import streamlit as st
import docx2txt
import re
import google.generativeai as genai
# from pyresparser import ResumeParser
from io import BytesIO
from datetime import datetime
from threading import Thread
from queue import Queue
from bokeh.models import ColumnDataSource, CustomJS
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
#from win32com import client as wc
import time
import shutil
import tempfile
# import paramiko
# from scp import SCPClient
#import textract
import subprocess
import aspose.words as aw
from spire.doc import *
from spire.doc.common import *

# Below 3 lines of code are used to declare the Gemini Engine
GOOGLE_API_KEY = "AIzaSyDWoz6YZOX_Qa2deOkmo8mg7oCW1L-yIVw"
#GOOGLE_API_KEY = "AIzaSyCduHrzs1_aC_d0FFk-UaGi-qsQ14IjocY"
genai.configure(api_key=GOOGLE_API_KEY)
model = genai.GenerativeModel('gemini-1.5-flash') 
#model = genai.GenerativeModel('gemini-1.5-flash-8b')

DS_Skills = ["Probability and Statistics", "Hypothesis Testing", "Regression Analysis", "Time Series Analysis",
             "Data Cleaning and Preprocessing",
             "Exploratory Data Analysis (EDA)", "Data Visualization (Matplotlib, Seaborn, ggplot2)",
             "Feature Engineering",
             "Supervised Learning (Classification, Regression)",
             "Unsupervised Learning (Clustering, Dimensionality Reduction)",
             "Ensemble Methods (Random Forests, Gradient Boosting)", "Text Processing", "Tokenization",
             "Named Entity Recognition", "Natural Language Processing", "Sentiment Analysis"]

DL_Skills = ["Neural Networks", "Convolutional Neural Networks (CNN)", "Recurrent Neural Networks (RNN)",
             "Long Short-Term Memory (LSTM)",
             "Gated Recurrent Units (GRU)", "Autoencoders", "Reinforcement Learning", "Transfer Learning",
             "Hyperparameter Tuning", "TensorFlow",
             "Keras", "PyTorch", "Theano", "Natural Language Processing (NLP) with Deep Learning Techniques",
             "Computer Vision with Deep Learning Techniques"]

MLOps_Skills = ["Continuous Integration/Continuous Deployment (CI/CD)", "Model Deployment", "Model Monitoring",
                "Model Versioning",
                "Model Orchestration", "Containerization (Docker, Kubernetes)", "Automated Testing",
                "Scalability", "Performance Optimization", "Model Governance", "Data Versioning", "MLOps Practices",
                "Cloud Computing Platforms (AWS, Azure, Google Cloud)",
                "Deployment Automation Tools (TensorFlow Serving, Seldon Core, Kubeflow)",
                "MLOps Platforms (MLflow, TFX, DataRobot)", "Workflow Automation Tools (Airflow, Luigi)"]

GenAI_Skills = ["Generative Adversarial Networks (GAN)", "Variational Autoencoders (VAE)", "Transformers", "DALL-E",
                "Langchain", "Hugging Face",
                "GPT", "Gemini", "OpenAI GPT 3.5", "Palm", " LLMs", "Llama2", "Mistral 8*7", "Prompt Engineering",
                "Transformers", "BERT", "Fine tune LLM using LoRA, QLoRA"]

Example1 = """
         Resume content = "POOJA UGALE JR.DATA SCIENTIST A tech lover Data Scientist with strong math background and 3+ years of experience using predictive modelling, data processing, and data mining algorithms to solve challenging business problems. 
                         poojaugale777@gmail.com +91 7057444357 Pune, Maharashtra WORK EXPERIENCE JR. DATA SCIENTIST IFS INDIA MARCANTILE Pvt. Ltd. NOV 2019 -Present Pune, Maharashtra Project Name: Built predictive based piecewise degradation model of remaining useful life prediction of an Turbofan engine 
                         Domain:Aviation Skills/Tools Used: LSTM, Train test split, CNN,GRU, Matplotlib, RNN, Encoder, Decoder, Attention Based Encoder Decoder GRU Description: • Built a machine learning and deep learning model that can predicts the failure of Machine in upcoming number of days. 
                         • The objective of this project is to implement various predictive maintenance methods and assess the performance of each equipment • Predictive maintenance techniques are used to determine the condition of an equipment to plan the maintenance /Failure ahead of its time. 
                         By doing this the equipment downtime cost can be reduced significantly. Project Name:On Field Survice Resource Monitoring System Domain: Construction Field Skills/Tools Used: Logestic egression ,SVM,KNN,MLP Classifier 
                         Description: • To Built a machine learning model that can used to classifying the performance is in working state,idle or other . • Installing IoT devices to generate the data and creat resouces management mining in construction business onsite and offsite and creat an automated dashboard to view the report. 
                         • Create web application to fetch the data from DB Project Name: Legal Document Curation using NLP Domain: Legal System Skills/Tools Used: • Tokenization,stemming,Lemmatization,train test split, Word cloud,TF-IDF,Gensim,Sumy, Transformer. 
                         Description: • An Important challenge in legal text summarization is identifying the information part while avoiding the irrelevant one. • The goal of thise project is to develop a web application for meaningful legal text summarization an extraction of entities 
                         • Summarization is performed by using Spacy,Gensim,Nltk etc KEY SKILLS Regression Analysis Natural Language Processing SQL Machine Learning Data Pre-processing Visualization Artificial Intelligence (AI) Deep Learning Data Cleaning Social Media Mining Feature Engineering Computer Vision Microsoft Excel Advanced Statistics Python Hypothesis Testing AWS 
                         TECHNICAL EXPERTISE • Programming languages: Python • Packages: Numpy, SciPy, Pandas, NLTK, Matplotlib, Scikit- learn, Tensorflow, Keras • Tools : Jupyter Notebook, Anaconda, MS Word and MS Excel. • Modeling : Regression, Classification like Random forest, Logistic, SVM,Bagging, Clustering -K-Means. • Deep Learning: Natural Language Processing, ANN, CNN, RNN, LSTM, etc. 
                         EDUCATIONAL DETAILS • Bachelor Degree Civil Engineering with 74% (Distinction)from Savitribai Phule Pune University. • HSC with 75.69% in 2014 from State Board, Maharashtra. • SSC with 82.73% in 2012 from State Board, Maharashtra. PERSONAL DETAILS • Date of Birth : 07/07/1996 • Current Address : Panchvati Society, R.no.401, B-wing, Sinhgad campus, Ambegaon,Pune-46 
                         • Languages : English, Hindi, Marathi SOFT SKILLS Complex Problem Solver Critical Thinking Team Player Good Product Understanding Quick Learner INTERESTS Learning New Skills Investing Spirituality Sports "

        Here are the details of the candidates:

        1. Programming Languages: Python
        2. Recent Experience: 4 years
        3. Machine learning skills: Logestic regression ,SVM, Tokenization,stemming,Lemmatization,train test split, Word cloud,TF-IDF,Gensim,Sumy,Transformer
        4. Deep Learning skills: LSTM, Train test split, CNN,GRU, Matplotlib, RNN, Encoder, Decoder, Attention Based Encoder Decoder GRU
        5. MLOps skills: None
        6. Cloud skills: AWS
        7. GenAI skills: None
        8. List of projects: 
        1) Built predictive based piecewise degradation model of remaining useful life prediction of an Turbofan engine
        2) On Field Survice Resource Monitoring System 
        3) Legal Document Curation using NLP
        9. Recommendation: Recommend for interview. The candidate has a knowledge of Python and has 4 years of recent experience. 
                           Candidate has Machine learning skills and Deep learning skills and done at least 3 projects, in addition they have cloud skills.
        10. Name of the candidate: POOJA UGALE
        11. count of projects: 3
        12. additional skills: None
        13. email: poojaugale777@gmail.com
        14. phone: +91 7057444357

        """

Example2 = """
        Resume content = "Jayshree Sharma Indian Institute of Information Technology Guwahati (+91) 8254812553/(+91)7099215869 bpsdworld@gmail.com 
                         OBJECTIVE WORK EXPERIENCES  Working as Data Scientist: Advanced Analytics in IBM, Bengaluru from 16/02/2021 to present. 
                         Project Digital Solutions Asset Performance Management Description In this project, we need to bring out an Artificial Intelligence based solution for asset management for the power grid company. 
                         Employer IBM India Private Limited Responsibilities  Working as Data Scientist, implementing business rules through IBM ODM (Operational Decision Management). \
                          Creating models and deployment with docker. Environment Python, Machine learning Models, ODM.  Worked as full time Data Science-intern at TNG Innovation Labs Pvt Ltd, Bangalore from June 2020 to 30 November 2020. 
                          Worked for fetching and analyzing data for live AI projects like fleet management, routing etc  Worked for Machine learning, AI projects for videos: Getting camera pose from dash cam video and finding path of vehicle including left/right turns and detecting accident/non accident. 
                         Used Random forest, CNN and other models.  Worked for Machine learning, AI projects for time series data: Used LSTM, Bi-LSTM, CNN-LSTM models. Project Accident Détection(Jido) Description As part of this project, we detect the accident or non-accidenttrips and type of accident (front/back/side(left/right) hit) using the sensor data collected from vehicle and using dashcam video. 
                         Employer TNG Innovation Labs Pvt. Ltd. Responsibilities  Getting camera pose from dash cam video using CNN and other models.  Finding path of vehicle including left/right turns and detecting accident/non accident using different approaches and finding best suitable one. 
                          Getting high performance using supervised and unsupervised approaches along with statistical methods. Environment Python, Machine learning Models Project Vehicle route tracking Description The objective of this project is to track the vehicle live route using sensor data collected from vehicle. 
                         Employer TNG Innovation Labs Pvt. Ltd. Responsibilities  Using reverse geo coding and finding the total length travelled by a vehicle within a state. 
                          Writing a script for finding at what time vehicle crosses the border and how many litres of fuel consumed by vehicle within a state Environment Python, Machine learning Models Project Employee Transportation system Description As part of this project, we have find best route for employee transportation with minimum time and least fuel consumption. Employer TNG Innovation Labs Pvt. Ltd. Responsibilities 
                          Analysing the Speed limit data in third party APIs like Mapbox navigation app, OSM, Tomtom Navigation App, Bing navigation app and finding the best app suitable as per use case.  Representing and comparing the data from all the possible app available in the market for navigation with speed limit.  Presenting the above details to seniors.  Doing research work on getting algorithms for bus stop selection with given constrains and delivering same to teammates. 
                         Environment Python, Machine learning Models  Worked as Computer Science Faculty in Jawahar Navodaya Vidyalaya, Nongstoin, Meghalaya from Aug 2016 to April 2017 
                          Worked as a Database Administrator at ISS (Integrated System Services) from May 2014 to May 2016. Worked on projects like SOCIO-ECONOMIC CASTE CENSUS (SECC). My responsibility was to manage and supervise a team of 50 resources across 12 districts of Assam under the SECC project working on solving technical issues like database issues, software problems and other non-technicalissues. 
                         Project Socio-Economic Caste census 2011 Description The objective of this project is to digitize the census data and correct it as per requirement. Employer Integrated Systems & Services. Responsibilities  My responsibility was to manage and supervise a team of 50 resources across 12 districts of Assam under the SECC project working on solving technical issues like database issues, software problems and other non-technicalissues. Environment Mysql. PAPER PUBLISHED 
                          An IoT based Smart Assistive Devicefor the Visually Impaired. 2020 IEEERegion 10 Symposium (TENSYMP) TECHNICAL SKILLS  PRIMARY SKILLS: MACHINE LEARNING, DEEP LEARNING, PYTHON, COMPUTER VISION, AI  SECONDARY SKILLS: DATA SCIENCE, BIG DATA, SPARK, AWS  Programming Skills: C, Python, C++, Scala (basics)  Other languages: MySQL,HTML, CSS  Operating System: Windows, Ubuntu  Software Tools/Frameworks: Jupiter Notebook, Anaconda, Pycharm, Eclipse, MatLab, Netbeans, Spark, Keras, Tensorflow, AWS. 
                          Data Science, Machine learning, AI using Numpy, Pandas, Scikit learn, Regression techniques, Clustering algorithms, SVM, Random forest, CNN, LSTM, geopandas, geocoding etc , . ACADEMIC DETAILS  M.Tech in Computer Science and Engineering [2020]: Indian Institute of Information Technology Guwahati (IIIT). CGPA:7.6  Bachelor of Technology (B. Tech IT) [ 2012]: Department of Information Technology, Assam University Silchar (Central University) Percentage: 66.25% 
                          Higher Secondary Examination [2006]: Jawahar Navodaya Vidyalaya (CBSE), Tinsukia, Assam Percentage: 72.00%  High School Leaving Certificate Examination [2004]: Jawahar Navodaya Vidyalaya (CBSE), Tinsukia Percentage: 79.40% PERSONAL SKILLS Comprehensive problem-solving abilities, excellent verbal and written communication skills, ability to interact with people efficiently, willingness to learn. Brilliant team player, self-driven, well at collaborating with the team all the time. CERTIFICATION 
                          Certification from NPTEL on Cloud Computing under Prof. Adrijit Goswami, IITKharagpur. EXTRA CURRICULAR ACTIVITIES AND ACHIEVEM ENTS  Qualified GATE 2011,2018,2019.  NCC “B” Certificate holder as “Senior Under Officer”  Secured “A” grade in All India Level Intelligence, Aptitude and Scholastic test.  Secured “A” grade in Indian Institute of personality development test at all India level.  Qualified 43rdAnnual All India UN Information Test.  Got “NEHRU AWARD” by Tea Board India.
                           Got 3rd prize in 8th National „Youth Parliament‟ Competition at Regional level.  Secured position in various competitions and eventslike racing, essay writing, song, badminton etc. ACADEMIC PROJECT  Human Activity Recognition using feature learning and Transfer learning (June 2019 –June 2020) M.Tech- Project- Under the guidance of Dr. Dip Sankar Banerjee, IIITG The project primarily aims at finding a good approach for recognizing human activity on sensor dataset. 
                          Focus on dynamic feature learning and Transfer Learning for activity recognition. Technology Used: Python3, keras with Tensorflow backend.  Anomaly Detection using MLlib (Jan -Apr 2019) Instructor Dr. Dip Sankar Banerjee, IIIT Guwahati The Credit Card Fraud Detection by applying different Machine Learning Techniqueson past credit card transaction dataset and analyzing the result. The aim here is to detect the fraudulent transactions by using different machine learning techniques. 
                          Technology Used: pySpark  IoT (Internet of Thing) project “Smart blind stick” Instructor Dr. Rakesh Matam, IIIT Guwahati (Jan-April 2019) The aim of the project is to build an intelligent walking stick that provide an efficient movement aid for the blind persons which gives a sense of vision by providing alert when any object comes in front of them and also notify family or friend when help is needed. Technology Used: Raspberry pi, Raspbian OS. 
                           Network on Chip ( 2012) Under the guidanceof Asst. Prof. Abhijit Biswas, TS School of Technology, Assam University Silchar. The project involved the study of delay optimization in Network on Chip.  Medical shop Management System ( 2011) Under the guidanceof Prof. Ajoy Kumar Khan, TS Schoolof Technology, Assam University Silchar. The project involved management of all records related to a medical shop. Technology Used: Java, MySQL INTERNSHIP 
                           Undertaken Industrial Training a t IOCL, Guwahati Refinery, Title: Networking REPOSITORIES AND NETWORK  LinkedIn: https://www.linkedin.com/in/jayshree -sharma-022bb672/  Github: https://github.com/Sharmajayshree REFERENCES References are available on request."

        Here are the details of the candidates:

        1. Programming Languages: Python
        2. Recent Experience: 3 years
        3. Machine learning skills: Python, Machine learning Models
        4. Deep Learning skills: None
        5. MLOps skills: None
        6. Cloud skills: AWS
        7. GenAI skills: None
        8. List of projects: 
        1) Digital Solutions Asset Performance Management                  
        9. Recommendation: Not Recommended for interview. Candidate have very less Machine learning skills and don't have deep learning, MLOps, cloud and genai skills.
                           Candidate don't have any good Machine learning projects.
        10. Name of the candidate: Jayshree Sharma 
        11. Count of projects: 1      
        12. additional skills: None
        13. email: bpsdworld@gmail.com
        14. phone: (+91) 8254812553/(+91)7099215869


        """

#This function is used to extract text from the PDF.
def extract_text_from_pdf(pdf_path):
    try:

        pdf_document = fitz.open(stream=pdf_path.read(), filetype="pdf")
        text = "".join([page.get_text() for page in pdf_document])

        return text

    except Exception as e:

        # st.error(f"Error processing {pdf_path}: {e}")

        return None

#This function is used to extract text from the doc
def extract_text_from_docx(docx_path):
    try:
        text = docx2txt.process(docx_path)
        # doc = Document(docx_path)
        # text = ""
        # for paragraph in doc.paragraphs:
        #     text += paragraph.text + "\n"
        return text

    except Exception as e:
        print(f"Error processing {docx_path}: {e}")
        return None


def extract_text_from_doc(doc_path):
    try:
        #doc = Document(docx_path)
        #text = ""
        #for paragraph in doc.paragraphs:
        #    text += paragraph.text + "\n"
        #text = textract.process(doc_path)
        # Create a Document object
        
        text = ""
        return text

    except Exception as e:
        print(f"Error processing {doc_path}: {e}")
        return None

# def convert_doc_to_docx(doc_file_path, docx_file_path):
#     word = wc.Dispatch('Word.Application')
#     doc = word.Documents.Open(doc_file_path)

#     doc.SaveAs(docx_file_path, FileFormat=16)
#     doc.Close()
#     word.Quit()


#This function is used to store the resumes in the form of dictionary from the provided folder path.
def load_resumes_as_dict(folder_path):
    resumes_dict = {}

    for file_name in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file_name)
        if file_name.endswith('.pdf'):
            text = extract_text_from_pdf(file_path)

        elif file_name.endswith('.docx'):
            text = extract_text_from_docx(file_path)

        #elif file_name.endswith('.doc'):
            #docx_file_path = file_path.replace('.doc', '.docx')
            #convert_doc_to_docx(file_path, docx_file_path)
            #text = extract_text_from_docx(docx_file_path)
            #os.remove(docx_file_path)

        else:
            continue

        if text:
            resumes_dict[file_name] = (file_path, text)

    return resumes_dict


#This is the batch-processing version of the above function.
def load_resumes_as_dict2(folder_path, batch_files):
    resumes_dict2 = {}

    for file_name in batch_files:
        file_path = os.path.join(folder_path, file_name)
        if file_name.endswith('.pdf'):
            text = extract_text_from_pdf(file_path)

        elif file_name.endswith('.docx'):
            text = extract_text_from_docx(file_path)

        #elif file_name.endswith('.doc'):
            #docx_file_path = file_path.replace('.doc','.docx')
            #convert_doc_to_docx(file_path, docx_file_path)
            #text = extract_text_from_docx(docx_file_path)
            #os.remove(docx_file_path)

        else:
            continue

        if text:
            resumes_dict2[file_name] = (file_path, text)

    return resumes_dict2


# data = []
# def generate_resume_summary(resumes_dict, Optional_skills, DS_Skills, DL_Skills, MLOps_Skills,GenAI_Skills, resume_folder,user_score,weight_ml,weight_dl,weight_genai,weight_mlops,weight_cloud):
def generate_resume_summary(resumes_dict, Optional_skills,DS_Skills,DL_Skills,MLOps_Skills,GenAI_Skills,selectfield,GOOGLE_API_KEY):
    data = []
    questions_dict = {}
    genai.configure(api_key=GOOGLE_API_KEY)
    model = genai.GenerativeModel('gemini-1.5-flash')
    for file_name, ( text_content) in resumes_dict.items():

        prompt5 = f"""
                Given is the resume content of the candidate. Analyze the resume.

                STEP 1:
                    Your task is to search for the technical skills section or similar sections where candidate usually write their skills as comma separated.
                    Along with this search for the professional summary or similar to summary where candidate gives small introduction about themselves. sometimes it is just a paragraph.
                    Here, use your understanding that you need to avoid excluding project details or information about projects.
                STEP 2:
                    You have to exclude that whole section containing skills as comma separated, paragraph of summary or introduction of candidate and give everything other than this in proper format.
                    Include name, Email ID and phone number.


                Resume content:{text_content} 
                """

        try:

            response_text5 = model.generate_content(prompt5, generation_config=genai.types.GenerationConfig(
                candidate_count=1,
                temperature=0.3)
                                                    )

            #print(response_text5.text)
            # model_info = genai.get_model("models/gemini-1.5-flash")
            # print(f"{model_info.input_token_limit=}")
            # print(f"{model_info.output_token_limit=}")
            # print(response_text5.usage_metadata)

            prompt4 = f"""
                Your task is to analyze the given resume. You have to think step by step as mentioned and give answer of the questions in just 1 to 2 words asked on the resume.

                I am providing you example of both recommended candidate as well as not recommended candidate. Please go through it and understand how skills and other things are extracted.
                I want you to follow the same logic while giving answers for the following questions. 

                Example of recommended candidate = {Example1}

                Example of not recommended candidate = {Example2}. Here Candidate has experience as an intern so we excluded everything.

                In the above both examples, user search for powerbi as optional skill but both candidate haven't worked on it that's why additional skills are marked as 'No'.

                Ensure accuracy while providing the answers. Avoid giving skills mentioned in the questions, used it only for your reference.

                STEP 1:
                    Question: Does candidate mentioned or used 'Python' or 'R' programming language? Provide programming langauge used only.

                    If answer is 'Yes' then only move to next part else simply return 'No programming language'.

                STEP 2:
                    Question:  Does the candidate have at least 2 years of recent experience in NLP, Machine learning, Machine learning, Deep learning or GenAI ?

                    If yes then only move to next question else provide 'Less experience'.

                    Question: If 'Yes', does the candidate have used any of the NLP, Machine learning, Machine learning, Deep learning or GenAI skills during this recent experience? 

                    Question: How many years of recent experience candidate have? Recent experience means last company or present company in which candidate is/was working. Consider present as Feb 2024.

                STEP 3:
                    For this step only consider resume content which talks about candidates actual work experience in Machine learning, deep learning, genai, mlops. Please focus only to answer skills related questions below based on such work experience/projects only. Don't extract skills which are mentioned explicitly under technical skills or primary skills or secondary skills and which are not part of sections related to work experience and project experience.
                    For your understanding following are the examples of skills mapping.
                    You should analyze the content of the resume and extract skills based on patterns and keywords commonly associated with these domains.
                    Avoid giving skills mentioned below, used it only for your reference. 

                    Machine learning skills : {DS_Skills}
                    Deep Learning skills : {DL_Skills}
                    MLOps skills : {MLOps_Skills} 
                    GenAI skills : {GenAI_Skills}

                    Your task here is to go through each project done by the candidate, extract relevant skills which are mentioned by the candidate in the project. 
                    If you are not able to extract any skill strictly provide none.

                    For following questions, Only extract projects/work experience/professional experience  related to machine learning, deep learning, mlops, genai ONLY if he/she 'developed' or 'build' something in the project. 

                    Question: What are the skills mentioned in sections like technical skills or primary skills or secondary skills? 
                    Please do not include this skills in the analysis.

                    Question: Does the candidate has any work experience as an intern? What skills he used as an intern ?
                    Please do not include this skills in the analysis.

                    Question: What are the specific Machine learning skills candidate has utilized or developed while working on projects except as an intern?
                    question: Are you sure ? Ensure accuracy while giving response and include only technical skills.
                    Question: What are the specific deep learning skills candidate has utilized or developed while working on projects except as an intern?
                    question: Are you sure ? Ensure accuracy while giving response and include only technical skills.
                    Question: What are the specific MLOps skills candidate has utilized or developed while working on projects?
                    question: Are you sure ? Ensure accuracy while giving response and include only technical skills.
                    Question: What are the specific cloud skills candidate has utilized or developed while working on projects?
                    question: Are you sure ? Ensure accuracy while giving response and include only technical skills.
                    Question: What are the specific GenAI skills candidate has utilized or developed while working on projects?
                    question: Are you sure ? Ensure accuracy while giving response and include only technical skills.
                    
                    If {Optional_skills} is empty then provide additional skills as 'None'.
                    check if candidate has mentioned or used skills {Optional_skills} in the resume, if yes then provide that skill as additional skills else return 'None'.
                    Avoid extracting skills other than {Optional_skills} in additional skills.
                    
                    Always verify your answer and check above mentioned conditions are met, then provide updated answers.

                STEP 4:
                    Question: Provide the machine learning/deep learning/genai related project names done by the candidate?
                    Question: Provide the total number of  machine learning/deep learning/genai related project done by the candidate as mentioned above?

                STEP 5:
                    You are SME in Machine learning and genai and your task is to evaluate given resume based on step 1 to step 5. 
                    If the candidate has the skills in Machine learning and deep learning in step 3 and at least 3 projects as calculated in step 4, recommend that candidate for job with detailed reasoning referring to the responses from step 1 to step 3. If the candidate is rejected give the proper reasoning.
                    Use your understanding while giving recommendation.

                STEP 6:
                    Provide the name of the candidate. If you are unable to find name then provide None.
                    Provide the Email ID of the candidate. If you are unable to find Email ID then provide None.
                    Provide the Phone number of the candidate. If you are unable to find Phone number then provide None.


                Please follow the below format (Use Plain text only, avoid using Bold format) for providing the response, Use the latest skills fetched for related asked question:


                1. Programming Languages:
                2. Recent Experience:
                3. Machine learning skills:
                4. Deep Learning skills:
                5. MLOps skills:
                6. Cloud skills:
                7. GenAI skills:
                8. List of projects:
                9. Recommendation:
                10. Name of the candidate:
                11. count of projects:
                12. additional skills:
                13. email:
                14. phone:


                Resume content:{response_text5.text}
                """

            response_text = model.generate_content(prompt4, generation_config=genai.types.GenerationConfig(
                candidate_count=1,
                temperature=0.3)
                                                   )

            # st.write(response_text.text)

            language_pattern = r"\*{0,2}Programming Languages:\*{0,2}\s(.*?)(?=\d+\. |\*{0,2}Recent Experience:\*{0,2}|\n)"

            duration_pattern = r"\*{0,2}Recent Experience:\*{0,2}\s(.*?)(?=\*{0,2}Machine learning skills:\*{0,2}|\n)"

            ds_pattern = r"\*{0,2}Machine learning skills:\*{0,2}\s(.*?)(?=\*{0,2}Deep Learning skills:\*{0,2}|\n)"

            dl_pattern = r"\*{0,2}Deep Learning skills:\*{0,2}\s(.*?)(?=\*{0,2}MLOps skills:\*{0,2}|\n)"

            MLOps_pattern = r"\*{0,2}MLOps skills:\*{0,2}\s(.*?)(?=\*{0,2}Cloud skills:\*{0,2}|\n)"

            Cloud_pattern = r"\*{0,2}Cloud skills:\*{0,2}\s(.*?)(?=\*{0,2}GenAI skills:\*{0,2}|\n)"

            genai_pattern = r"\*{0,2}GenAI skills:\*{0,2}\s(.*?)(?=\*{0,2}List of projects:\*{0,2}|\n)"

            projects_pattern = r"\*{0,2}List of projects:\*{0,2}\s(.*?)(?=\*{0,2}\d+\. Recommendation:\*{0,2}|\Z)"

            recommendation_pattern = r"\*{0,2}Recommendation:\*{0,2}\s(.*?)(?=\*{0,2}Name of the candidate:\*{0,2}|\n)"

            name_pattern = r"\*{0,2}Name of the candidate:\*{0,2}\s(.*?)(?=\*{0,2}count of projects:\*{0,2}|\n)"

            count_pattern = r"\*{0,2}count of projects:\*{0,2}\s(.*?)(?=\*{0,2}additional skills:\*{0,2}|\n)"

            additional_pattern = r"\*{0,2}additional skills:\*{0,2}\s(.*?)(?=\*{0,2}email:\*{0,2}|\n)"

            email_pattern = r"\*{0,2}email:\*{0,2}\s(.*?)(?=\*{0,2}phone:\*{0,2}|\n)"
            
            phone_pattern = r"\*{0,2}phone:\*{0,2}\s(.*?)$"


            # Extract text following each label using regular expressions

            extract_text = lambda pattern, response: re.search(pattern, response, re.DOTALL).group(

                1).strip() if re.search(pattern, response, re.DOTALL) else "Could not read"

            language = extract_text(language_pattern, response_text.text)

            duration = extract_text(duration_pattern, response_text.text)

            ds_skills = extract_text(ds_pattern, response_text.text)

            dl_skills = extract_text(dl_pattern, response_text.text)

            mlops_skills = extract_text(MLOps_pattern, response_text.text)

            cloud_skills = extract_text(Cloud_pattern, response_text.text)

            genai_skills = extract_text(genai_pattern, response_text.text)

            projects = extract_text(projects_pattern, response_text.text)
            # st.write(projects)

            recommendation = extract_text(recommendation_pattern, response_text.text)

            additional = extract_text(additional_pattern, response_text.text)

            name = extract_text(name_pattern, response_text.text)
            
            email = extract_text(email_pattern, response_text.text)

            phone = extract_text(phone_pattern, response_text.text)
            
            # print(phone)
            phone = phone.strip()
            phone = phone.replace('+91',"")
            phone = phone.replace('-',"")
            phone = phone.replace('+',"").replace('(',"").replace(")","")
            phone = phone.strip()
            # print(phone)

            if str(name).lower() in ['none', 'not provided', 'not mentioned', 'could not read', 'not available',
                                     'not mentioned in the resume']:
                name = ' '.join(
                    re.sub(r'\[.*?\]', '', os.path.splitext(file_name)[0]).replace('+', ' ').replace('_', ' ').replace(
                        'Naukri', '').replace('PHENOM', '').split()[:2]).title()
            else:
                name = name.title()

            count = extract_text(count_pattern, response_text.text)

            if 'None' in str(count) or 'Not' in str(count) or 'No' in str(count) or 'not' in str(count) or str(
                    count) == "":
                count = 0
            else:
                count = count

            ds_list = ds_skills.split(',')
            ds_list = [item.strip() for item in ds_list]
            dl_list = dl_skills.split(',')
            dl_list = [item.strip() for item in dl_list]
            genai_list = genai_skills.split(',')
            genai_list = [item.strip() for item in genai_list]

            llm_list = ['large language models (LLMs)', 'LLM', 'GPT', 'GPT 3.5', 'GPT-3', 'BERT', 'ChatGPT',
                        'GPT-4', 'Generative AI', 'GenAI', 'Prompts', 'AzureOpenAI',
                        'langchain', 'Llamaindex framework', 'LLMs', 'Vector DB', 'Open-AI Da-Vinci LLM',
                        'Development of LLMs and Generative AI applications', "LLM's",
                        'PaLM', 'Gen AI', 'Prompt Engineering', 'Fine tune LLM using LoRA', 'QLoRA']

            for item in llm_list:
                if item in ds_list:
                    ds_list.remove(item)
                    if item not in genai_list:
                        genai_list.append(item)
                if item in dl_list:
                    dl_list.remove(item)
                    if item not in genai_list:
                        genai_list.append(item)

            ds_skills = ','.join(ds_list)
            dl_skills = ','.join(dl_list)
            genai_skills = ','.join(genai_list)
            genai_skills = genai_skills.replace('None,', '')

            def calculate_skill_score(skills_str):
                skill_count = 0
                skill_count = skills_str.count(',') + 1

                score = 0
                if skill_count == 1 and (
                        'None' in skills_str or 'Not' in skills_str or 'No' in skills_str or 'not' in skills_str or skills_str == ""):
                    skill_count = 0
                    score = 0


                elif skill_count >= 1 and skill_count <= 8:
                    score = 0.25 * skill_count
                else:
                    score = 2

                return [skill_count, score]

            def calculate_ds_skill_score(skills_str):
                skill_count = 0
                skill_count = skills_str.count(',') + 1

                score = 0
                if skill_count == 1 and (
                        'None' in skills_str or 'Not' in skills_str or 'No' in skills_str or 'not' in skills_str or skills_str == ""):
                    skill_count = 0
                    score = 0


                elif skill_count >= 1 and skill_count <= 10:
                    score = 0.2 * skill_count
                else:
                    score = 2

                return [skill_count, score]

            def calculate_dl_skill_score(skills_str):
                skill_count = 0
                skill_count = skills_str.count(',') + 1

                score = 0
                if skill_count == 1 and (
                        'None' in skills_str or 'Not' in skills_str or 'No' in skills_str or 'not' in skills_str or skills_str == ""):
                    skill_count = 0
                    score = 0


                elif skill_count >= 1 and skill_count <= 10:
                    score = 0.25 * skill_count
                else:
                    score = 2

                return [skill_count, score]

            def calculate_MLOps_skill_score(skills_str):
                skill_count = 0
                skill_count = skills_str.count(',') + 1

                score = 0
                if skill_count == 1 and (
                        'None' in skills_str or 'Not' in skills_str or 'No' in skills_str or 'not' in skills_str or skills_str == ""):
                    skill_count = 0
                    score = 0


                elif skill_count >= 1 and skill_count <= 3:
                    score = 0.2 * skill_count
                else:
                    score = 0.6

                return [skill_count, score]

            def calculate_cloud_skill_score(skills_str):

                skill_count = 0
                skill_count = skills_str.count(',') + 1

                score = 0
                if skill_count == 1 and (
                        'None' in skills_str or 'Not' in skills_str or 'No' in skills_str or 'not' in skills_str or skills_str == ""):
                    skill_count = 0
                    score = 0


                elif skill_count >= 1 and skill_count <= 3:
                    score = 0.2 * skill_count
                else:
                    score = 0.6

                return [skill_count, score]

            def calculate_combine_skill_score(combine_count):

                if combine_count == 0:
                    score = 0

                elif combine_count >= 1 and combine_count <= 4:
                    score = 0.25 * combine_count
                else:
                    score = 1

                return score

            def calculate_project_score(count):

                if count == 0:
                    score = 0

                elif count >= 1 and count <= 10:
                    score = 0.1 * count
                else:
                    score = 1

                return score

            genAI_count = calculate_skill_score(str(genai_skills))
            ds_count = calculate_ds_skill_score(str(ds_skills))
            dl_count = calculate_dl_skill_score(str(dl_skills))
            mlops_count = calculate_MLOps_skill_score(str(mlops_skills))
            cloud_count = calculate_cloud_skill_score(str(cloud_skills))

            sum1 = dl_count[0] + ds_count[0]+genAI_count[0]
            # st.write(sum1)

            if genAI_count[0] == 0:
                genai_status = "No"
            else:
                genai_status = "Yes"

            if ds_count[0] == 0:
                ds_status = "No"
            else:
                ds_status = "Yes"

            if dl_count[0] == 0:
                dl_status = "No"
            else:
                dl_status = "Yes"

            if mlops_count[0] == 0:
                mlops_status = "No"
            else:
                mlops_status = "Yes"

            if cloud_count[0] == 0:
                cloud_status = "No"
            else:
                cloud_status = "Yes"

            # ds_list = ds_skills.split(',')
            # dl_list = dl_skills.split(',')
            # combine_skills = ds_list + dl_list
            # combine_count = len(combine_skills)

            combine_score = ds_count[1] + dl_count[1]

            project_score = calculate_project_score(int(count))

            total_score = combine_score + project_score + cloud_count[1] + genAI_count[1] + mlops_count[1]
            # st.write(total_score)

            rerun = False

            if total_score == 0 and not rerun:
                prompt5 = f"""
                                Given is the resume content of the candidate. Analyze the resume.

                                STEP 1:
                                    Your task is to search for the technical skills section or similar sections where candidate usually write their skills as comma separated.
                                    Along with this search for the professional summary or similar to summary where candidate gives small introduction about themselves. sometimes it is just a paragraph.
                                    Here, use your understanding that you need to avoid excluding project details or information about projects.
                                STEP 2:
                                    You have to exclude that whole section containing skills as comma separated, paragraph of summary or introduction of candidate and give everything other than this in proper format.
                                    Include name, Email ID and phone number.


                                Resume content:{text_content} 
                                """

                try:

                    response_text5 = model.generate_content(prompt5, generation_config=genai.types.GenerationConfig(
                        candidate_count=1,
                        temperature=0.3)
                                                            )

                    # st.write(response_text5.text)

                    prompt4 = f"""
                                Your task is to analyze the given resume. You have to think step by step as mentioned and give answer of the questions in just 1 to 2 words asked on the resume.

                                I am providing you example of both recommended candidate as well as not recommended candidate. Please go through it and understand how skills and other things are extracted.
                                I want you to follow the same logic while giving answers for the following questions. 

                                Example of recommended candidate = {Example1}

                                Example of not recommended candidate = {Example2}. Here Candidate has experience as an intern so we excluded everything.

                                In the above both examples, user search for powerbi as optional skill but both candidate haven't worked on it that's why additional skills are marked as 'No'.

                                Ensure accuracy while providing the answers. Avoid giving skills mentioned in the questions, used it only for your reference.

                                STEP 1:
                                    Question: Does candidate mentioned or used 'Python' or 'R' programming language? Provide programming langauge used only.

                                    If answer is 'Yes' then only move to next part else simply return 'No programming language'.

                                STEP 2:
                                    Question:  Does the candidate have at least 2 years of recent experience in NLP, Machine learning, Machine learning, Deep learning or GenAI ?

                                    If yes then only move to next question else provide 'Less experience'.

                                    Question: If 'Yes', does the candidate have used any of the NLP, Machine learning, Machine learning, Deep learning or GenAI skills during this recent experience? 

                                    Question: How many years of recent experience candidate have? Recent experience means last company or present company in which candidate is/was working. Consider present as Feb 2024.

                                STEP 3:
                                    For this step only consider resume content which talks about candidates actual work experience in Machine learning, deep learning, genai, mlops. Please focus only to answer skills related questions below based on such work experience/projects only. Don't extract skills which are mentioned explicitly under technical skills or primary skills or secondary skills and which are not part of sections related to work experience and project experience.
                                    For your understanding following are the examples of skills mapping.
                                    You should analyze the content of the resume and extract skills based on patterns and keywords commonly associated with these domains. 
                                    Avoid giving skills mentioned below, used it only for your reference.

                                    Machine learning skills : {DS_Skills}
                                    Deep Learning skills : {DL_Skills}
                                    MLOps skills : {MLOps_Skills} 
                                    GenAI skills : {GenAI_Skills}

                                    Your task here is to go through each project done by the candidate, extract relevant skills which are mentioned by the candidate in the project. 
                                    If you are not able to extract any skill strictly provide none.

                                    For following questions, Only extract projects/work experience/professional experience  related to machine learning, deep learning, mlops, genai ONLY if he/she 'developed' or 'build' something in the project. 

                                    Question: What are the skills mentioned in sections like technical skills or primary skills or secondary skills? 
                                    Please do not include this skills in the analysis.

                                    Question: Does the candidate has any work experience as an intern? What skills he used as an intern ?
                                    Please do not include this skills in the analysis.

                                    Question: What are the specific Machine learning skills candidate has utilized or developed while working on projects except as an intern?
                                    question: Are you sure ? Ensure accuracy while giving response and include only technical skills. 
                                    Question: What are the specific deep learning skills candidate has utilized or developed while working on projects except as an intern?
                                    question: Are you sure ? Ensure accuracy while giving response and include only technical skills.
                                    Question: What are the specific MLOps skills candidate has utilized or developed while working on projects?
                                    question: Are you sure ? Ensure accuracy while giving response and include only technical skills.
                                    Question: What are the specific cloud skills candidate has utilized or developed while working on projects?
                                    question: Are you sure ? Ensure accuracy while giving response and include only technical skills.
                                    Question: What are the specific GenAI skills candidate has utilized or developed while working on projects?
                                    question: Are you sure ? Ensure accuracy while giving response and include only technical skills.
                                    
                                    If {Optional_skills} is empty then provide additional skills as 'None'.
                                    check if candidate has mentioned or used skills {Optional_skills} in the resume, if yes then provide that skill as additional skills else return 'None'.
                                    Avoid extracting skills other than {Optional_skills} in additional skills.
                    
                                    Always verify your answer and check above mentioned conditions are met, then provide updated answers.

                                STEP 4:
                                    Question: Provide the machine learning/deep learning/genai related project names done by the candidate?
                                    Question: Provide the total number of  machine learning/deep learning/genai related project done by the candidate as mentioned above?

                                STEP 5:
                                    You are SME in Machine learning and genai and your task is to evaluate given resume based on step 1 to step 5. 
                                    If the candidate has the skills in Machine learning and deep learning in step 3 and at least 3 projects as calculated in step 4, recommend that candidate for job with detailed reasoning referring to the responses from step 1 to step 3. If the candidate is rejected give the proper reasoning.
                                    Use your understanding while giving recommendation.

                                STEP 6:
                                    Provide the name of the candidate. If you are unable to find name then provide None.
                                    Provide the Email ID of the candidate. If you are unable to find Email ID then provide None.
                                    Provide the Phone number of the candidate. If you are unable to find Phone number then provide None.

                                Please follow the below format (Use Plain text only, avoid using Bold format) for providing the response, Use the latest skills fetched for related asked question:

                                1. Programming Languages:
                                2. Recent Experience:
                                3. Machine learning skills:
                                4. Deep Learning skills:
                                5. MLOps skills:
                                6. Cloud skills:
                                7. GenAI skills:
                                8. List of projects:
                                9. Recommendation:
                                10. Name of the candidate:
                                11. count of projects:
                                12. additional skills:
                                13. email:
                                14. phone:


                                Resume content:{response_text5.text}
                                """

                    response_text = model.generate_content(prompt4, generation_config=genai.types.GenerationConfig(
                        candidate_count=1,
                        temperature=0.3)
                                                           )

                    # st.write(response_text.text)

                    language_pattern = r"\*{0,2}Programming Languages:\*{0,2}\s(.*?)(?=\d+\. |\*{0,2}Recent Experience:\*{0,2}|\n)"

                    duration_pattern = r"\*{0,2}Recent Experience:\*{0,2}\s(.*?)(?=\*{0,2}Machine learning skills:\*{0,2}|\n)"

                    ds_pattern = r"\*{0,2}Machine learning skills:\*{0,2}\s(.*?)(?=\*{0,2}Deep Learning skills:\*{0,2}|\n)"

                    dl_pattern = r"\*{0,2}Deep Learning skills:\*{0,2}\s(.*?)(?=\*{0,2}MLOps skills:\*{0,2}|\n)"

                    MLOps_pattern = r"\*{0,2}MLOps skills:\*{0,2}\s(.*?)(?=\*{0,2}Cloud skills:\*{0,2}|\n)"

                    Cloud_pattern = r"\*{0,2}Cloud skills:\*{0,2}\s(.*?)(?=\*{0,2}GenAI skills:\*{0,2}|\n)"

                    genai_pattern = r"\*{0,2}GenAI skills:\*{0,2}\s(.*?)(?=\*{0,2}List of projects:\*{0,2}|\n)"

                    projects_pattern = r"\*{0,2}List of projects:\*{0,2}\s(.*?)(?=\*{0,2}\d+\. Recommendation:\*{0,2}|\Z)"

                    recommendation_pattern = r"\*{0,2}Recommendation:\*{0,2}\s(.*?)(?=\*{0,2}Name of the candidate:\*{0,2}|\n)"

                    name_pattern = r"\*{0,2}Name of the candidate:\*{0,2}\s(.*?)(?=\*{0,2}count of projects:\*{0,2}|\n)"

                    count_pattern = r"\*{0,2}count of projects:\*{0,2}\s(.*?)(?=\*{0,2}additional skills:\*{0,2}|\n)"

                    additional_pattern = r"\*{0,2}additional skills:\*{0,2}\s(.*?)(?=\*{0,2}email:\*{0,2}|\n)"

                    email_pattern = r"\*{0,2}email:\*{0,2}\s(.*?)(?=\*{0,2}phone:\*{0,2}|\n)"
                    
                    phone_pattern = r"\*{0,2}phone:\*{0,2}\s(.*?)$"

                    # Extract text following each label using regular expressions

                    extract_text = lambda pattern, response: re.search(pattern, response, re.DOTALL).group(

                        1).strip() if re.search(pattern, response, re.DOTALL) else "Could not read"

                    language = extract_text(language_pattern, response_text.text)

                    duration = extract_text(duration_pattern, response_text.text)

                    ds_skills = extract_text(ds_pattern, response_text.text)

                    dl_skills = extract_text(dl_pattern, response_text.text)

                    mlops_skills = extract_text(MLOps_pattern, response_text.text)

                    cloud_skills = extract_text(Cloud_pattern, response_text.text)

                    genai_skills = extract_text(genai_pattern, response_text.text)

                    projects = extract_text(projects_pattern, response_text.text)
                    # st.write(projects)

                    recommendation = extract_text(recommendation_pattern, response_text.text)

                    additional = extract_text(additional_pattern, response_text.text)

                    name = extract_text(name_pattern, response_text.text)
                    
                    email = extract_text(email_pattern, response_text.text)

                    phone = extract_text(phone_pattern, response_text.text)
                    
                    # print(phone)
                    phone = phone.strip()
                    phone = phone.replace('+91',"")
                    phone = phone.replace('-',"")
                    phone = phone.replace('+',"").replace('(',"").replace(")","")
                    phone = phone.strip()
                    # print(phone)


                    if str(name).lower() in ['none', 'not provided', 'not mentioned', 'could not read', 'not available',
                                             'not mentioned in the resume']:
                        name = ' '.join(
                            re.sub(r'\[.*?\]', '', os.path.splitext(file_name)[0]).replace('+', ' ').replace('_',
                                                                                                             ' ').replace(
                                'Naukri', '').replace('PHENOM', '').split()[:2]).title()
                    else:
                        name = name.title()

                    count = extract_text(count_pattern, response_text.text)

                    if 'None' in str(count) or 'Not' in str(count) or 'No' in str(count) or 'not' in str(count) or str(
                            count) == "":
                        count = 0
                    else:
                        count = count

                    ds_list = ds_skills.split(',')
                    ds_list = [item.strip() for item in ds_list]
                    dl_list = dl_skills.split(',')
                    dl_list = [item.strip() for item in dl_list]
                    genai_list = genai_skills.split(',')
                    genai_list = [item.strip() for item in genai_list]

                    llm_list = ['large language models (LLMs)', 'LLM', 'GPT', 'GPT 3.5', 'GPT-3', 'BERT', 'ChatGPT',
                                'GPT-4', 'Generative AI', 'GenAI', 'Prompts', 'AzureOpenAI',
                                'langchain', 'Llamaindex framework', 'LLMs', 'Vector DB', 'Open-AI Da-Vinci LLM',
                                'Development of LLMs and Generative AI applications', "LLM's",
                                'PaLM', 'Gen AI', 'Prompt Engineering', 'Fine tune LLM using LoRA', 'QLoRA']

                    for item in llm_list:
                        if item in ds_list:
                            ds_list.remove(item)
                            if item not in genai_list:
                                genai_list.append(item)
                        if item in dl_list:
                            dl_list.remove(item)
                            if item not in genai_list:
                                genai_list.append(item)

                    ds_skills = ','.join(ds_list)
                    dl_skills = ','.join(dl_list)
                    genai_skills = ','.join(genai_list)
                    genai_skills = genai_skills.replace('None,', '')

                    def calculate_skill_score(skills_str):
                        skill_count = 0
                        skill_count = skills_str.count(',') + 1

                        score = 0
                        if skill_count == 1 and (
                                'None' in skills_str or 'Not' in skills_str or 'No' in skills_str or 'not' in skills_str or skills_str == ""):
                            skill_count = 0
                            score = 0


                        elif skill_count >= 1 and skill_count <= 8:
                            score = 0.25 * skill_count
                        else:
                            score = 2

                        return [skill_count, score]

                    def calculate_ds_skill_score(skills_str):
                        skill_count = 0
                        skill_count = skills_str.count(',') + 1

                        score = 0
                        if skill_count == 1 and (
                                'None' in skills_str or 'Not' in skills_str or 'No' in skills_str or 'not' in skills_str or skills_str == ""):
                            skill_count = 0
                            score = 0


                        elif skill_count >= 1 and skill_count <= 10:
                            score = 0.2 * skill_count
                        else:
                            score = 2

                        return [skill_count, score]

                    def calculate_dl_skill_score(skills_str):
                        skill_count = 0
                        skill_count = skills_str.count(',') + 1

                        score = 0
                        if skill_count == 1 and (
                                'None' in skills_str or 'Not' in skills_str or 'No' in skills_str or 'not' in skills_str or skills_str == ""):
                            skill_count = 0
                            score = 0


                        elif skill_count >= 1 and skill_count <= 10:
                            score = 0.2 * skill_count
                        else:
                            score = 2

                        return [skill_count, score]

                    def calculate_MLOps_skill_score(skills_str):
                        skill_count = 0
                        skill_count = skills_str.count(',') + 1

                        score = 0
                        if skill_count == 1 and (
                                'None' in skills_str or 'Not' in skills_str or 'No' in skills_str or 'not' in skills_str or skills_str == ""):
                            skill_count = 0
                            score = 0


                        elif skill_count >= 1 and skill_count <= 3:
                            score = 0.2 * skill_count
                        else:
                            score = 0.6

                        return [skill_count, score]

                    def calculate_cloud_skill_score(skills_str):

                        skill_count = 0
                        skill_count = skills_str.count(',') + 1

                        score = 0
                        if skill_count == 1 and (
                                'None' in skills_str or 'Not' in skills_str or 'No' in skills_str or 'not' in skills_str or skills_str == ""):
                            skill_count = 0
                            score = 0


                        elif skill_count >= 1 and skill_count <= 3:
                            score = 0.2 * skill_count
                        else:
                            score = 0.6

                        return [skill_count, score]

                    def calculate_combine_skill_score(combine_count):

                        if combine_count == 0:
                            score = 0

                        elif combine_count >= 1 and combine_count <= 4:
                            score = 0.25 * combine_count
                        else:
                            score = 1

                        return score

                    def calculate_project_score(count):

                        if count == 0:
                            score = 0

                        elif count >= 1 and count <= 10:
                            score = 0.1 * count
                        else:
                            score = 1

                        return score

                    genAI_count = calculate_skill_score(str(genai_skills))
                    ds_count = calculate_ds_skill_score(str(ds_skills))
                    dl_count = calculate_dl_skill_score(str(dl_skills))
                    mlops_count = calculate_MLOps_skill_score(str(mlops_skills))
                    cloud_count = calculate_cloud_skill_score(str(cloud_skills))

                    sum1 = dl_count[0] + ds_count[0]+genAI_count[0]
                    # st.write(sum1)

                    if genAI_count[0] == 0:
                        genai_status = "No"
                    else:
                        genai_status = "Yes"

                    if ds_count[0] == 0:
                        ds_status = "No"
                    else:
                        ds_status = "Yes"

                    if dl_count[0] == 0:
                        dl_status = "No"
                    else:
                        dl_status = "Yes"

                    if mlops_count[0] == 0:
                        mlops_status = "No"
                    else:
                        mlops_status = "Yes"

                    if cloud_count[0] == 0:
                        cloud_status = "No"
                    else:
                        cloud_status = "Yes"

                    # ds_list = ds_skills.split(',')
                    # dl_list = dl_skills.split(',')
                    # combine_skills = ds_list + dl_list
                    # combine_count = len(combine_skills)

                    combine_score = ds_count[1] + dl_count[1]

                    project_score = calculate_project_score(int(count))

                    total_score = combine_score + project_score + cloud_count[1] + genAI_count[1] + mlops_count[1]
                    # st.write(total_score)

                except Exception:
                    total_score = combine_score + project_score + cloud_count[1] + genAI_count[1] + mlops_count[1]

                rerun = True

            if selectfield == 'GenAI/DataScience':

                if total_score >= 1.8 and sum(
                        x == 'None' for x in [ds_skills, dl_skills, mlops_skills, cloud_skills, genai_skills]) <= 4:
                    if (genAI_count[0]) > 2:
                        new_recommend_genai = "Yes.Recommended for GenAI profile."
                        if (combine_score) >= 1:
                            new_recommend_ds = "Yes.Recommended for Data Science profile."
                        else:
                            new_recommend_ds = "No.Candidate doesn't have essential skills."
                        if (mlops_count[0]) >=2 and cloud_count[0] > 0:
                            mlops_recommend = "Yes.Recommended for MLOps profile."
                        else:
                            mlops_recommend = "No.Candidate doesn't have essential skills."
                    else:
                        new_recommend_ds = "Yes.Recommended for Data Science profile."
                        new_recommend_genai = "No.Candidate doesn't have essential skills."
                        if (mlops_count[0]) >=2 and cloud_count[0] > 0:
                            mlops_recommend = "Yes.Recommended for MLOps profile."
                        else:
                            mlops_recommend = "No.Candidate doesn't have essential skills."

                elif total_score >= 1 and sum(
                        x == 'None' for x in [ds_skills, dl_skills, mlops_skills, cloud_skills, genai_skills]) <= 2:
                    if (genAI_count[0]) > 2:
                        new_recommend_genai = "Yes.Recommended for GenAI profile."
                        if (combine_score) >= 1:
                            new_recommend_ds = "Yes.Recommended for Data Science profile."
                        else:
                            new_recommend_ds = "No.Candidate doesn't have essential skills."
                        if (mlops_count[0]) >=2 and cloud_count[0] > 0:
                            mlops_recommend = "Yes.Recommended for MLOps profile."
                        else:
                            mlops_recommend = "No.Candidate doesn't have essential skills."
                    else:
                        new_recommend_ds = "Yes.Recommended for Data Science profile."
                        new_recommend_genai = "No.Candidate doesn't have essential skills."
                        if (mlops_count[0]) >=2 and cloud_count[0] > 0:
                            mlops_recommend = "Yes.Recommended for MLOps profile."
                        else:
                            mlops_recommend = "No.Candidate doesn't have essential skills."

                elif sum1 >= 5 and total_score >= 1.5 and sum(
                        x == 'None' for x in [ds_skills, dl_skills, mlops_skills, cloud_skills, genai_skills]) <= 3:
                    if (genAI_count[0]) > 2:
                        new_recommend_genai = "Yes.Recommended for GenAI profile."
                        if (combine_score) >= 1:
                            new_recommend_ds = "Yes.Recommended for Data Science profile."
                        else:
                            new_recommend_ds = "No.Candidate doesn't have essential skills."
                        if (mlops_count[0]) >=2 and cloud_count[0] > 0:
                            mlops_recommend = "Yes.Recommended for MLOps profile."
                        else:
                            mlops_recommend = "No.Candidate doesn't have essential skills."
                    else:
                        new_recommend_ds = "Yes.Recommended for Data Science profile."
                        new_recommend_genai = "No.Candidate doesn't have essential skills."
                        if (mlops_count[0]) >=2 and cloud_count[0] > 0:
                            mlops_recommend = "Yes.Recommended for MLOps profile."
                        else:
                            mlops_recommend = "No.Candidate doesn't have essential skills."

                elif (genAI_count[0]) > 4:
                    new_recommend_genai = "Yes.Recommended for GenAI profile."
                    if (combine_score) >= 1:
                        new_recommend_ds = "Yes.Recommended for Data Science profile."
                    else:
                        new_recommend_ds = "No.Candidate doesn't have essential skills."
                    if (mlops_count[0]) >= 2 and cloud_count[0] > 0:
                        mlops_recommend = "Yes.Recommended for MLOps profile."
                    else:
                        mlops_recommend = "No.Candidate doesn't have essential skills."

                else:
                    new_recommend_ds = "No.Candidate doesn't have essential skills."
                    new_recommend_genai = "No.Candidate doesn't have essential skills."
                    if (mlops_count[0]) >= 2 and cloud_count[0] > 0:
                        mlops_recommend = "Yes.Recommended for MLOps profile."
                    else:
                        mlops_recommend = "No.Candidate doesn't have essential skills."

                prompt2 = f"""
                Your task is to generate positive summary if  {new_recommend_genai} =="Recommended for GenAI profile." or {new_recommend_ds} =="Recommended for Data Science profile." for the candidate based on the given skills.
                Your task is also to generate negative summary if {new_recommend_genai} == "No. Candidate doesn't have essential skills." or {new_recommend_ds} == "No. Candidate doesn't have essential skills." for the candidate based on the given skills if candidate has less or no skills or worked on less than 2 projects.

                provide in detail positive,negative summary in following format(Use Plain text only, avoid using Bold format) .

                For Example:

                 1. Positive Summary:
                 2. Negative Summary: 


                Machine learning skills = {ds_skills}
                Deep learning skills = {dl_skills}
                MLOps skills = {mlops_skills}
                Cloud skills = {cloud_skills}
                GenAI skills = {genai_skills}

                """

                response_text1 = model.generate_content(prompt2, generation_config=genai.types.GenerationConfig(
                    candidate_count=1,
                    temperature=0.3)
                                                        )

                # st.write(response_text1.text)

            elif selectfield == 'Machine Learning':

                if (ds_count[0]) > 6 and (int(count)) > 2:
                    new_recommend_ds = "Yes.Recommended for Machine Learning profile."
                    new_recommend_genai =""
                    if (mlops_count[0]) >= 2 and cloud_count[0] > 0:
                        mlops_recommend = "Yes.Recommended for MLOps profile."
                    else:
                        mlops_recommend = "No.Candidate doesn't have essential skills."

                else:
                    new_recommend_ds = "No.Candidate doesn't have essential skills."
                    new_recommend_genai = ""
                    if (mlops_count[0]) >= 2 and cloud_count[0] > 0:
                        mlops_recommend = "Yes.Recommended for MLOps profile."
                    else:
                        mlops_recommend = "No.Candidate doesn't have essential skills."

                prompt2 = f"""
                Your task is to generate positive summary if  {new_recommend_ds} =="Recommended for Machine Learning profile." for the candidate based on the given skills.
                Your task is also to generate negative summary if {new_recommend_ds} == "No. Candidate doesn't have essential skills." for the candidate based on the given skills if candidate has less or no skills or worked on less than 2 projects.

                provide in detail positive,negative summary in following format(Use Plain text only, avoid using Bold format) .

                For Example:

                 1. Positive Summary:
                 2. Negative Summary: 


                Machine learning skills = {ds_skills}

                """

                response_text1 = model.generate_content(prompt2, generation_config=genai.types.GenerationConfig(
                    candidate_count=1,
                    temperature=0.3)
                                                        )

                # st.write(response_text1.text)

            else:

                if (dl_count[0]) > 6 and (int(count)) > 2:
                    new_recommend_ds = "Yes.Recommended for Deep Learning profile."
                    new_recommend_genai = ""
                    if (mlops_count[0]) >= 2 and cloud_count[0] > 0:
                        mlops_recommend = "Yes.Recommended for MLOps profile."
                    else:
                        mlops_recommend = "No.Candidate doesn't have essential skills."

                else:
                    new_recommend_ds = "No.Candidate doesn't have essential skills."
                    new_recommend_genai = ""
                    if (mlops_count[0]) >= 2 and cloud_count[0] > 0:
                        mlops_recommend = "Yes.Recommended for MLOps profile."
                    else:
                        mlops_recommend = "No.Candidate doesn't have essential skills."

                prompt2 = f"""
                Your task is to generate positive summary if  {new_recommend_ds} =="Recommended for Deep Learning profile." for the candidate based on the given skills.
                Your task is also to generate negative summary if {new_recommend_ds} == "No. Candidate doesn't have essential skills." for the candidate based on the given skills if candidate has less or no skills or worked on less than 2 projects.

                provide in detail positive,negative summary in following format(Use Plain text only, avoid using Bold format) .

                For Example:

                 1. Positive Summary:
                 2. Negative Summary: 

                Deep learning skills = {dl_skills}

                """

                response_text1 = model.generate_content(prompt2, generation_config=genai.types.GenerationConfig(
                    candidate_count=1,
                    temperature=0.3)
                                                        )

                # st.write(response_text1.text)

            positive_pattern = r"\*{0,2}Positive Summary:\*{0,2}\s(.*?)(?=\*{0,2}Negative Summary:\*{0,2}|\Z)"

            negative_pattern = r"\*{0,2}Negative Summary:\*{0,2}\s(.*?)$"

            extract_text = lambda pattern, response: re.search(pattern, response, re.DOTALL).group(

                1).strip() if re.search(pattern, response, re.DOTALL) else "Could not read"

            positive = extract_text(positive_pattern, response_text1.text)
            # st.write(positive)

            negative = extract_text(negative_pattern, response_text1.text)
            # st.write(negative)

            if 'No.' in new_recommend_ds or 'No.' in new_recommend_genai:
                recommendation = negative

            else:
                recommendation = positive
            # st.write(recommendation)

            # if 'None' in additional:
            #     additional = ''
            additional_list = additional.split(',')

            for i in range(len(additional_list)):
                additional_list[i] = additional_list[i].lower()

            for i in range(len(Optional_skills)):
                Optional_skills[i] = Optional_skills[i].lower()

            if ('Recommended' in new_recommend_ds or 'Recommended' in new_recommend_genai) and any(
                    skill in additional_list for skill in Optional_skills):
                alter = 'Yes'
                Optional_skills = ','.join(Optional_skills).capitalize()
                if 'Recommended' in new_recommend_ds:
                    new_recommend_ds = new_recommend_ds + f" Also recommended for {Optional_skills}"
                if 'Recommended' in new_recommend_genai:
                    new_recommend_genai = new_recommend_genai + f" Also recommended for {Optional_skills}"
            else:
                alter = 'No'
                additional = ''
            # alter = 'No'

            # if total_score == 0:
            #     language = duration = ds_skills = ds_status = dl_skills = dl_status = mlops_skills = mlops_status = cloud_skills = cloud_status = genai_skills = genai_status = additional = projects = recommendation = alter = ""
            #     new_recommend = "Need Human Intervention"

            if 'Recommended' in new_recommend_ds or 'Recommended' in new_recommend_genai:
                prompt99 = f"""
                        Consider you are an interviewer, your task is to ask potential questions to be asked from the candidate whose resume is {text_content}.

                        Generate a set of 10 comprehensive interview questions based on resume provided. Questions must be related to Data Science skills and/or GenAI skills for a candidate applying for a role either of Data Science and/or Generative AI. These questions should assess the candidate's understanding, experience, and technical skills across candidate's major skills:

                        1. Projects and Technical Aspects (5 Questions): Questions should delve into the candidate's practical experience and the technical challenges they've navigated in their projects. Generate case study based technical questions as an interviewer based on projects/work experiences done by the candidates. The dept of Questions should increase complexity wise from question 1 to 5. Each question should also have a section to ask specific technical package/skills used related to python language.

                        2. Python Specific Technical Skills (3 Questions): Ask conceptual programming questions to check candidate's programming ability and problem solving power in python.

                        3. Additional questions (2 questions): Ask questions based on other skills/cloud if mentioned in resume. for example, "How would you deploy machine learning model on AWS?"

                        Provide response in the tabular format. First row should contain only one column where name of the candidate should be mentioned.
                        Next rows should contain 2 columns , first column for questions with question number like Q1,Q2,etc, second column for score and it should be empty always.

                        Exclude the provided example from the analysis, Use the example only for the formatting.

                        For example:
                        "
                        | Name of the candidate |
                        | --- |
                        | DEBASIS SAMAL |
                        | Questions | Score |
                        | --- | --- |
                        | Q1.Walk me through the process of building the movie recommender system that you build. How did you train the NLP vectorizer model and predict similar movies based on cosine similarity? |  |
                        | Q2.Describe your experience in developing the multilingual chatbot using Generative AI (OpenAI) on AWS SageMaker. How did you handle the challenges of multilingual processing and knowledge extraction from PDF documents? |  |
                        | --- | --- |
                        | ... |  |
                        "


                        Name of the candidate = {name}
                        Machine learning skills = {ds_skills}
                        Deep learning skills = {dl_skills}
                        MLOps skills = {mlops_skills}
                        Cloud skills = {cloud_skills}
                        GenAI skills = {genai_skills}

                        """
                questions_text = model.generate_content(prompt99, generation_config=genai.types.GenerationConfig(
                    candidate_count=1,
                    temperature=0.3)
                                                        )

                # st.write(questions_text.text)
                # st.write(type(questions_text.text))
                
                questions_dict[file_name]= questions_text.text

                # Parse response and extract data
                lines = questions_text.text.strip().split('\n')
                candidate_name = lines[2].strip().split("|")[1].strip()
                questions_and_scores = [line.strip() for line in lines[4:]]

                # Create a new Document
                doc = Document()

                sections = doc.sections
                for section in sections:
                    section.left_margin = Pt(35)  # Adjust left margin as needed
                    section.right_margin = Pt(35)  # Adjust right margin as needed
                    section.top_margin = Pt(35)  # Adjust top margin as needed
                    section.bottom_margin = Pt(35)  # Adjust bottom margin as needed

                timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                # doc.add_paragraph(f"Report Generated At: {timestamp}",style = 'Heading 1')

                timestamp_paragraph = doc.add_paragraph(f"Report Generated At: {timestamp}")
                timestamp_run = timestamp_paragraph.runs[0]
                timestamp_run.bold = True
                timestamp_run.font.size = Pt(11)  # Adjust the font size as needed

                # Add candidate name
                doc.add_heading('Candidate Name:', level=1)
                doc.add_paragraph(candidate_name)

                # Add table
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'

                table.alignment = WD_TABLE_ALIGNMENT.LEFT

                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Questions'
                hdr_cells[1].text = 'Score'

                # Make column names bold
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                # Populate table
                for line in questions_and_scores:
                    parts = line.split('|')
                    if len(parts) > 1 and parts[1].strip() != '---':  # Check if the line is not just '---'
                        question = parts[1].strip()
                        score = parts[2].strip() if len(
                            parts) > 2 else ""  # Handle the case where score is not provided
                        row_cells = table.add_row().cells
                        row_cells[0].text = question
                        row_cells[1].text = score

                        # Add space after each question
                        row_cells[0].paragraphs[0].runs[-1].add_break()

                        # for cell in row_cells:
                        #     for paragraph in cell.paragraphs:
                        #         paragraph.space_after = Pt(10)

                # Adjust row height to fit content
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Set alignment to left
                            paragraph.space_after = Pt(12)  # Add space after each paragraph (adjust as needed)
                        cell.height = Pt(60)  # Adjust row height to fit content

                # Add a border around the entire document
                for paragraph in doc.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    paragraph.space_after = Pt(12)
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black for better visibility
                        run.font.name = 'Cambria'
                        run.font.size = Pt(11)

                # Add space after the headings "Questions" and "Score"
                for cell in table.rows[0].cells:
                    cell.paragraphs[0].runs[-1].add_break()

                for cell in table.columns[1].cells:
                    cell.width = Pt(30)  # Adjust cell width as needed

                for cell in table.columns[0].cells:
                    cell.width = Pt(700)

                table.border_collapse = True

                # new_folder = os.path.join(resume_folder, 'Generated_Docx')
                # os.makedirs(new_folder, exist_ok=True)
                # docx_filename = os.path.join(new_folder, f'{name}_reports.docx')
                # doc.save(docx_filename)
                # download_link = f"file:///{docx_filename}"
                # docx_link = f'=HYPERLINK("{download_link}", "Open Report")'
                docx_link = ''
            else:
                docx_link = ''

            prompt1 = f"""
            Based on the following data,return 'Yes' or 'No' to indicate whether the profile is suitable for a data science-related position.

            Following are the conditions for positive recommendation:

                1.Candidate must have any of the 'Python' or 'R' in 'Skills' or 'Programming Languages'.

                2.Candidate must have either of 'GenAI Related' or 'Data Science Related' skills.

            Indicate 'Yes' only if the candidate satisfies two conditions above else 'No'.   
            Indicate 'No' if only 'Machine Learning' is present in 'Data Science Related' skills and has 'None' in 'GenAI Related' skills.

            Also, provide a short reason for your recommendation.
            If 'GenAI Related' skills present then recommend candidate for GenAI profile.
            If only 'Data Science Related' skills present then recommend candidate for Data Science profile.
            Avoid using phrases like 'Answer:'.
            If Recommendation is 'No', provide reason as 'Candidate doesn't have essential skills'.

            Please focus exclude the provided example from the analysis.

            For example:
            Yes. Recommended for Data Science profile.
            Reason: Given the strong foundation and candidate demonstrates expertise in Data Science skills, the candidate is well-suited for positions related to Data Science.
            The candidate possesses proficiency in the Programming Languages, specifically in 'Python.' 
            Notably, the candidate has cloud skills of AWS,Azure.

            Focus solely on the relevant information requested and avoid additional details.

            Summary: {response_text1.text}

            """

            # recommendation = openai.Completion.create(

            #     engine="text-davinci-003",

            #     prompt=prompt1,

            #     max_tokens=500,

            #     temperature=0

            # )

            # recommendation_text = recommendation.choices[0].text.strip()

            # recommendation_text = model.generate_content(prompt1, generation_config=genai.types.GenerationConfig(
            #     candidate_count=1,
            #     top_p=0.6,
            #     top_k=5,
            #     temperature=0)
            #                                              )

            # recommendation_text=process1(response_text,prompt1)

            # role, experience,datascience_skills[1], recommendation,pdf_link,docx_link

            pdf_link =''
            data.append([file_name, name, language, duration, ds_skills, ds_status, dl_skills, dl_status, mlops_skills,
                         mlops_status, cloud_skills, cloud_status, genai_skills, genai_status, additional, alter,
                         projects, recommendation, new_recommend_genai,new_recommend_ds,mlops_recommend, genAI_count[0], ds_count[0], dl_count[0],
                         mlops_count[0], cloud_count[0], count,
                         combine_score, project_score, cloud_count[1], genAI_count[1], mlops_count[1], total_score,
                         pdf_link, docx_link, email, phone])

            # Print statements to check progress
            # st.write(data)
            print(f"Processed {file_name}")
            time.sleep(9)

        except Exception as e:

            # st.error(f"Error processing summary for {file_name}: {e}")
            print(f"Error processing summary for {file_name}: {e}")

            data.append(
                [file_name, "Error in processing", "Error in processing", "Error in processing", "Error in processing",
                 "Error in processing", "Error in processing", "Error in processing",
                 "Error in processing","Error in processing","Error in processing",
                 "Error in processing", "Error in processing", "Error in processing", "Error in processing",
                 "Error in processing", "Error in processing","Error in processing",
                 "Error in processing", "Error in processing", "Error in processing", "Error in processing",
                 "Error in processing", "Error in processing","Error in processing",
                 "Error in processing", "Error in processing", "Error in processing", "Error in processing",
                 "Error in processing", "Error in processing", "Error in processing",
                 "Error in processing", "Error in processing", 0, "Error in processing", "Error in processing"])

    return data,questions_dict


def generate_questions(recommended_dict):
    questions = []
    for file_name, content in recommended_dict.items():

        prompt9 = f"""
                Generate a set of 15 comprehensive interview questions based on resume provided for a candidate applying for a role involving Data Science and/or Generative AI. These questions should assess the candidate's understanding, experience, and technical skills across three distinct areas:

                1. Fundamentals and Reasoning-Based Questions (First 5 Questions): Questions in this section should test the candidate's grasp of core concepts in data science and generative AI, including their ability to explain fundamental principles, methodologies, and theoretical aspects.Generate questions based on projects/work experiences done by the candidates. 

                2. Projects and Technical Aspects (Next 5 Questions): This section should delve into the candidate's practical experience and the technical challenges they've navigated in their projects. Generate questions based on projects/work experiences done by the candidates. 

                3. Python or R Specific Technical Skills (Last 5 Questions): Focus on assessing the candidate's proficiency in programming languages essential for data science and AI, such as Python or R. Generate questions based on projects/work experiences done by the candidates.

                Resume Content = {content}
                """

        try:
            questions_text = model.generate_content(prompt9, generation_config=genai.types.GenerationConfig(
                candidate_count=1,
                temperature=0.3)
                                                    )

            # st.write(questions_text.text)

            questions.append([file_name, questions_text.text])


        except Exception as e:

            # st.error(f"Error processing summary for {file_name}: {e}")

            questions.append([file_name], "Error in processing")

    return questions

#This function is rudimentary.
def rank_resume(data):
    prompt7 = f"""
        Analyze all the resumes given below for skills related to data science, genai,ai,ml.
        provide Confidence score out of 10 for each of the resume based on the GenAI related skills and Data Science related skills.

        resume data={data}
"""
    Rank = model.generate_content(prompt7, generation_config=genai.types.GenerationConfig(
        candidate_count=1,
        top_p=0.6,
        top_k=5,
        temperature=0))

    return Rank.text

#This function is rudimentaryt.
def resume_filter(resumes_dict, Optional_skills):
    fill = []
    for file_name, (file_path, text_content) in resumes_dict.items():
        # parsed_resume = ResumeParser(file_path).get_extracted_data()

        # skills_used = ', '.join(parsed_resume.get('skills', 'Not Provided'))
        # st.write(skills_used)
        # skills_used_list = [skill.strip().lower() for skill in skills_used.split(',')]
        # # st.write(skills_used_list)
        # optional_skills_lower = [skill.strip().lower() for skill in Optional_skills]
        # for skill in optional_skills_lower:
        #     if skill in skills_used_list:
        #         fill.append(file_name)

        prompt5 = f"""
                Please analyze the resume content for searching skills {Optional_skills} in the resume.

                Please check if candidate has mentioned or used skills {Optional_skills} in the resume, if yes then return {file_name} of the candidate, else return 'None'.
                if {Optional_skills} id empty then return 'None'.

                Focus solely on the relevant information requested and avoid additional details.

                Here is the resume content: {text_content}.

                """

        response = model.generate_content(prompt5, generation_config=genai.types.GenerationConfig(
            candidate_count=1,
            top_p=0.6,
            top_k=5,
            temperature=0)
                                          )
        fill.append(response.text)

    return fill

#This is a rudimentary function.
def Shortlisted_Resumes1(data):
    prompt3 = f"""
                Utilize the provided resume data containing 'File Name' and 'Recommendation' for multiple candidates.

                Total number of Resumes: (Extract the total count of PDF files (e.g Chandanverma.pdf) mentioned in 'File Name'.

                Please analyze the provided json format data, use your understanding and extract the following information:

                    - Number of Resumes suited for Data Science job profile: (Calculate the total count of 'File Name' where 'Recommendation' starts with word 'Yes'.)
                    - Number of Not suitable resumes: (Calculate the total count of 'File Name' where 'Recommendation' starts with word 'No')
                    - File names suitable for Data Science job profile: (Provide a list of all file names where 'Recommendation' starts with word 'Yes'.)

                Please ensure accuracy in counting, always verify your answer.
                Exclude the provided example from the analysis.

                For example:

                Total number of Resumes: 3
                Number of Resumes suited for Data Science job profile: 3
                Number of Not suitable resumes: 0
                File names suitable for Data Science job profile: Chandanverma.pdf, Chethan N.pdf, Yashaswini Kulkarni.pdf

                Provide the analysis in a structured format as outlined above. 
                Please focus solely on extracting the requested information. 

                Here is the provided data: {data}

                """

    response_text = model.generate_content(prompt3, generation_config=genai.types.GenerationConfig(
        candidate_count=1,
        temperature=0.3)
                                           )

    return response_text.text

#This is a rudimentary function.
def Shortlisted_Resumes(data):
    prompt3 = f"""
            Utilize the provided resume data containing details of multiple candidates in the list inside the list.

            Total number of Resumes: (Extract the total count of PDF files (e.g Chandanverma.pdf) mentioned in the data.Please focus for Counting 'Total number of resumes', avoid considering index of the list as it starts from 0.)

            Please analyze the provided data, use your understanding and extract the following information:

                - Number of Resumes suited for Data Science job profile: (Extract the total count of resumes where the last element in each resume's list starts with word 'Yes.'.)
                - Number of Not suitable resumes: (Extract the total count of resumes where the last element in each resume's list starts with word 'No.' or 'Unknown'.)
                - File names suitable for Data Science job profile: (Provide a list of file names for resumes where the last element in each resume's list starts with word 'Yes.'.)

            Exclude the provided example from the analysis.

            For example:

            Total number of Resumes: 3
            Number of Resumes suited for Data Science job profile: 3
            Number of Not suitable resumes: 0
            File names suitable for Data Science job profile: Chandanverma.pdf, Chethan N.pdf, Yashaswini Kulkarni.pdf

            Please focus to go through each of the resume details then only provide requested information.
            Provide the analysis in a structured format as outlined above. 
            Please focus solely on extracting the requested information. 
            Ensure accuracy in counting the 'Total number of resumes','Number of Resumes suited for Data Science job profile','Number of Not suitable resumes'.

            Here is the provided data: {data}

            """

    response_text = model.generate_content(prompt3, generation_config=genai.types.GenerationConfig(
        candidate_count=1,
        temperature=0)
                                           )

    return response_text.text

#This functionality is to be trigerred from frontend for downloading csv.
def download_button(object_to_download, download_filename, button_text='Download as CSV'):
    csv = object_to_download.to_csv(index=False)

    b64 = base64.b64encode(csv.encode()).decode()

    href = f'<a href="data:file/csv;base64,{b64}" download="{download_filename}">{button_text}</a>'

    return href


def download_button_excel(object_to_download, download_filename, button_text='Download'):
    object_to_download.seek(0)
    timestamp = datetime.now().strftime('%Y/%m/%d/%H:%M:%S')
    download_filename = f'resume_summaries_{timestamp}.xlsx'
    b64 = base64.b64encode(object_to_download.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="{download_filename}">{button_text}</a>'

#Function to process the firstbox
def prathambox(firstbox):
    x = firstbox
    return x

#Function to process the secondbox
def dutiyabox(skills_to_search):
    y = skills_to_search
    return y
#Function to process the thirdbox
def tritiyabox(folder_path):
    global resume_folder
    if folder_path:
        resume_folder = rf"{folder_path}"
    return resume_folder
#function to split the optionalskills
def fourthbox(skills_to_search):
    global Optional_skills
    Optional_skills = [skill.strip() for skill in skills_to_search.split(",")]

    return Optional_skills

#This function shows userInput summary.
def userInput_summary(firstbox, skills_to_search,folder_path):
    # if folder_path:
    #     resume_folder = rf"{folder_path}"


    # resumes_dict = load_resumes_as_dict(resume_folder)

    # Calculate number of files in folder
    num_files = len(folder_path)
    # Display number of files to user
    xy = firstbox + ", " + skills_to_search

    library = []
    library.append([num_files, xy])
    df_library = pd.DataFrame(library, columns=["# Resumes to Screen", "Required Skillset to Filter"])

    return df_library.to_json(),num_files

#Function to return df_summaries
# def dataFr(Optional_skills,resume_folder,user_score,weight_ml,weight_dl,weight_genai,weight_mlops,weight_cloud):
def dataFr(selectfield,Optional_skills,resume_dict,GOOGLE_API_KEY):

    # resumes_dict = load_resumes_as_dict2(resume_folder)
    # # st.write(resumes_dict)
    summary_data,questions_dict = generate_resume_summary(resume_dict, Optional_skills,DS_Skills,DL_Skills,MLOps_Skills,GenAI_Skills,selectfield,GOOGLE_API_KEY)

    # batch_size = 10
    # files = os.listdir(resume_folder)
    # num_filess = len(files)
    # num_batches = (num_filess + batch_size - 1) // batch_size

    # for i in range(num_batches):
    #     start_i = i * batch_size
    #     end_i = min((i + 1) * batch_size, num_filess)
    #     batch_files = files[start_i:end_i]
    #     resumes_dict = load_resumes_as_dict2(resume_folder, batch_files)
    #     # st.write(resumes_dict)
    #     summary_data = generate_resume_summary(resumes_dict, Optional_skills,DS_Skills,DL_Skills,MLOps_Skills,GenAI_Skills,resume_folder,selectfield)
        # summary_data = generate_resume_summary(resumes_dict, Optional_skills, DS_Skills, DL_Skills, MLOps_Skills,GenAI_Skills, resume_folder,user_score,weight_ml,weight_dl,weight_genai,weight_mlops,weight_cloud)

    df_summaries = pd.DataFrame(summary_data,
                                      columns=['File Name','Name','Programming Language','Recent Experience','Machine learning','ML Status',
                                               'Deep Learning','DL Status','MLOps','MLops Status','Cloud','Cloud Status','GenAI','GenAI Status',
                                               'Additional Skills','Alternate Recommendation','Projects','Candidate Summary','GenAI Recommendation','DS Recommendation','MLOps Recommendation','GenAI Count','ML Count',
                                               'DL Count','MLops Count','Cloud Count','Project Count','ML DL Score','Project Score','Cloud Score','GenAI Score',
                                               'MLOps Score','Total Score', 'View PDF','View Docx', 'Email', 'Phone'])

    df_summaries['Total Score'] = pd.to_numeric(df_summaries['Total Score'], errors='coerce')

    df_summaries = df_summaries.sort_values(by=['Total Score'], ascending=False)
    # df_summaries['Additional Skills'] = ''
    # df_summaries['Alternate Recommendation'] = 'No'  # Default to 'No'
    #
    # for i in Optional_skills:
    #     mask = df_summaries['Skills'].str.contains(i, case=False)
    #     recommendation_mask = df_summaries['Recommendation'].str.startswith('Yes', na=False)
    #
    #     # Check if the skill is not already present in 'Additional Skills'
    #     not_present_mask = ~df_summaries['Additional Skills'].str.contains(i, case=False) | df_summaries[
    #         'Additional Skills'].isna()
    #
    #     # Update 'Alternate Recommendation' based on conditions
    #     df_summaries.loc[mask & recommendation_mask & not_present_mask, 'Alternate Recommendation'] = 'Yes'
    #
    #     # Append the skill only if it's not already present
    #     df_summaries.loc[mask & recommendation_mask & not_present_mask, 'Additional Skills'] = df_summaries.loc[
    #         mask & recommendation_mask & not_present_mask, 'Additional Skills'].apply(
    #         lambda x: x + ',' + i if pd.notna(x) and x != '' else i)
    #
    # # Update 'Recommendation' to 'No' if 'Additional Skills' is empty
    # df_summaries.loc[df_summaries['Additional Skills'].eq('') | df_summaries[
    #     'Additional Skills'].isna(), 'Alternate Recommendation'] = 'No'

    return  df_summaries,questions_dict

#function to return genai_count
def genAICount(noob):
    genai_count = noob['GenAI Recommendation'].str.contains('Recommended for GenAI profile').sum()

    return genai_count


# function to return datscience_count
def dsCount(noob):
    data_science_count = noob['DS Recommendation'].str.contains(
        'Recommended for Data Science profile').sum()

    return data_science_count


def mlCount(noob):
    ml_count = noob['DS Recommendation'].str.contains('Yes.Recommended for Machine Learning profile.').sum()

    return ml_count


def dlCount(noob):
    dl_count = noob['DS Recommendation'].str.contains('Yes.Recommended for Deep Learning profile.').sum()

    return dl_count


# function to return preffered_count
def altrCount(noob):
    alternate_count = noob['Alternate Recommendation'].str.contains('Yes').sum()

    return alternate_count


# function to return borderline_count
def bdrCount(noob):
    borderline_count = noob['DS Recommendation'].str.contains('Borderline').sum()

    return borderline_count


# function to return resume summary
def resuSumm(noob, skill, additionalskill):
    summary = []
    total_resumes = 0
    total_resumes = len(noob)

    recommended_count = noob.apply(lambda row: any(
        'Yes' in str(val) for val in [row['GenAI Recommendation'], row['DS Recommendation'], row['MLOps Recommendation']]),
                                   axis=1).sum()

    notrecommended_count = noob.apply(lambda row: all(
        'No' in str(val) for val in [row['GenAI Recommendation'], row['DS Recommendation'], row['MLOps Recommendation']]),
                                      axis=1).sum()

    data_count = {"Total": str(total_resumes), "Recommended": str(recommended_count),
                  "Not Recommended": str(notrecommended_count)}
    data_count = {"Resume Classification": data_count}
    json_data = json.dumps(data_count)

    genai_count = noob['GenAI Recommendation'].str.contains('Yes.Recommended for GenAI profile.').sum()
    # genai_borderline_count = noob['Recommendation'].str.contains('Borderline. Recommended for GenAI profile.').sum()
    filtered_skills_genai = noob['GenAI'].dropna().replace('None', '').replace('Not Mentioned', '').replace(
        'Error in processing', '').replace('Could not read', '')
    all_genai_skills = filtered_skills_genai.str.split(',').explode().str.strip().str.rstrip(';')
    all_genai_skills = all_genai_skills[all_genai_skills != '']
    trend_genai_df = pd.DataFrame({'skills': all_genai_skills})
    top5_ganai_skills = trend_genai_df['skills'].value_counts().nlargest(5).index
    top5_ganai_skills_dict = top5_ganai_skills.tolist()
    top5_ganai_skills_dict = ', '.join(top5_ganai_skills_dict)

    genai_ui = {"Recommended": str(genai_count), "Trend": top5_ganai_skills_dict}
    genai_ui = {"GenAI count": genai_ui}
    genai_ui_json = json.dumps(genai_ui)

    data_science_count = noob['DS Recommendation'].str.contains('Yes.Recommended for Data Science profile.').sum()
    # datascience_borderline_count = noob['Recommendation'].str.contains('Borderline. Recommended for Data Science profile.').sum()
    filtered_skills_ds = noob['Machine learning'].dropna().replace('None', '').replace('Not Mentioned', '').replace(
        'Error in processing', '').replace('Could not read', '')
    all_datascience_skills = filtered_skills_ds.str.split(',').explode().str.strip().str.rstrip(';')
    all_datascience_skills = all_datascience_skills[all_datascience_skills != '']
    trend_datascience_df = pd.DataFrame({'skills': all_datascience_skills})
    top5_datasceince_skills = trend_datascience_df['skills'].value_counts().nlargest(5).index
    top5_datasceince_skills_dict = top5_datasceince_skills.tolist()
    top5_datasceince_skills_dict = ', '.join(top5_datasceince_skills_dict)

    ds_ui = {"Recommended": str(data_science_count), "Trend": top5_datasceince_skills_dict}
    ds_ui = {"Data Science count": ds_ui}
    ds_ui_json = json.dumps(ds_ui)

    mlops_count = noob['MLOps Recommendation'].str.contains('Yes.Recommended for MLOps profile.').sum()
    # genai_borderline_count = noob['Recommendation'].str.contains('Borderline. Recommended for GenAI profile.').sum()
    filtered_skills_mlops = noob['MLOps'].dropna().replace('None', '').replace('Not Mentioned', '').replace(
        'Error in processing', '').replace('Could not read', '')
    all_mlops_skills = filtered_skills_mlops.str.split(',').explode().str.strip().str.rstrip(';')
    all_mlops_skills = all_mlops_skills[all_mlops_skills != '']
    trend_mlops_df = pd.DataFrame({'skills': all_mlops_skills})
    top5_mlops_skills = trend_mlops_df['skills'].value_counts().nlargest(5).index
    top5_mlops_skills_dict = top5_mlops_skills.tolist()
    top5_mlops_skills_dict = ', '.join(top5_mlops_skills_dict)

    mlops_ui = {"Recommended": str(mlops_count), "Trend": top5_mlops_skills_dict}
    mlops_ui = {"MLOps count": mlops_ui}
    mlops_ui_json = json.dumps(mlops_ui)

    ml_count = noob['DS Recommendation'].str.contains('Yes.Recommended for Machine Learning profile.').sum()
    filtered_skills_ml = noob['Machine learning'].dropna().replace('None', '').replace('Not Mentioned', '').replace(
        'Error in processing', '').replace('Could not read', '')
    all_ml_skills = filtered_skills_ml.str.split(',').explode().str.strip().str.rstrip(';')
    all_ml_skills = all_ml_skills[all_ml_skills != '']
    trend_ml_df = pd.DataFrame({'skills': all_ml_skills})
    top5_ml_skills = trend_ml_df['skills'].value_counts().nlargest(5).index
    top5_ml_skills_dict = top5_ml_skills.tolist()
    top5_ml_skills_dict = ', '.join(top5_ml_skills_dict)

    ml_ui = {"Recommended": str(ml_count), "Trend": top5_ml_skills_dict}
    ml_ui = {"Machine learning count": ml_ui}
    ml_ui_json = json.dumps(ml_ui)

    dl_count = noob['DS Recommendation'].str.contains('Yes.Recommended for Deep Learning profile.').sum()
    filtered_skills_dl = noob['Deep Learning'].dropna().replace('None', '').replace('Not Mentioned', '').replace(
        'Error in processing', '').replace('Could not read', '')
    all_dl_skills = filtered_skills_dl.str.split(',').explode().str.strip().str.rstrip(';')
    all_dl_skills = all_dl_skills[all_dl_skills != '']
    trend_dl_df = pd.DataFrame({'skills': all_dl_skills})
    top5_dl_skills = trend_dl_df['skills'].value_counts().nlargest(5).index
    top5_dl_skills_dict = top5_dl_skills.tolist()
    top5_dl_skills_dict = ', '.join(top5_dl_skills_dict)

    dl_ui = {"Recommended": str(dl_count), "Trend": top5_dl_skills_dict}
    dl_ui = {"Deep learning count": dl_ui}
    dl_ui_json = json.dumps(dl_ui)

    alternate_count = noob['Alternate Recommendation'].str.contains('Yes').sum()
    genai_additional_count = noob[noob['Alternate Recommendation'].str.contains('Yes')]
    genai_additional_count = genai_additional_count[
        genai_additional_count['GenAI Recommendation'].str.contains('Recommended for GenAI profile')]
    genai_additional_count = len(genai_additional_count)
    ds_additional_count = noob[noob['Alternate Recommendation'].str.contains('Yes')]
    ds_additional_count = ds_additional_count[
        ds_additional_count['DS Recommendation'].str.contains('Recommended for Data Science profile.')]
    ds_additional_count = len(ds_additional_count)
    ml_additional_count = noob[noob['Alternate Recommendation'].str.contains('Yes')]
    ml_additional_count = ml_additional_count[
        ml_additional_count['DS Recommendation'].str.contains('Recommended for Machine Learning profile.')]
    ml_additional_count = len(ml_additional_count)
    dl_additional_count = noob[noob['Alternate Recommendation'].str.contains('Yes')]
    dl_additional_count = dl_additional_count[
        dl_additional_count['DS Recommendation'].str.contains('Recommended for Deep Learning profile.')]
    dl_additional_count = len(dl_additional_count)
    mlops_additional_count = noob[noob['Alternate Recommendation'].str.contains('Yes')]
    mlops_additional_count = mlops_additional_count[
        mlops_additional_count['MLOps Recommendation'].str.contains('Recommended for MLOps profile.')]
    mlops_additional_count = len(mlops_additional_count)
 
    additional_ui = {"Recommended": str(alternate_count), "GenAI Recommended": str(genai_additional_count),
                     "Data Science Recommended": str(ds_additional_count),
                     "Machine learning Recommended": str(ml_additional_count),
                     "Deep learning Recommended": str(dl_additional_count),
                     "MLOps Recommended": str(mlops_additional_count)}
    additional_ui = {"Additional count": additional_ui}
    additional_ui_json = json.dumps(additional_ui)

    table_ui = json.dumps(
        {"serial": str(total_resumes), "skillset": str(skill), "additionalSkills": str(additionalskill)})

    # summary.append([total_resumes, m1, m2, m3])
    #
    # df_new = pd.DataFrame(summary, columns=['Total Number of Resumes', 'GenAI', 'Data Science',
    #                                         'Profile screened for Additional Skills'])

    return json_data, genai_ui_json, ds_ui_json,ml_ui_json, dl_ui_json, additional_ui_json, table_ui,mlops_ui_json


# function for GenAI resume results
def genAIres(noob, m1):
    # global gensa
    genai_df = noob[noob['GenAI Recommendation'].str.contains('Recommended for GenAI profile')]
    show_genai_df = genai_df[
        [ 'File Name', 'Name','GenAI Recommendation', 'MLOps Recommendation','GenAI', 'Machine learning','Deep Learning','MLOps',
          'Cloud', 'Programming Language', 'View PDF']]
    if m1 == 0:
        print("No GenAI related resumes found !")
    else:
        show_genai_df
    return show_genai_df.to_json(orient='records')


# function for dataScience resume results
def DSres(noob, m1):
    # global dssa
    data_science_df = noob[
        noob['DS Recommendation'].str.contains('Recommended for Data Science profile')]
    show_data_science_df = data_science_df[
        ['File Name','Name', 'DS Recommendation', 'MLOps Recommendation','GenAI', 'Machine learning', 'Deep Learning','MLOps', 'Cloud'
         , 'Programming Language', 'View PDF']]
    if m1 == 0:
        print("No Data Science related resumes found !")
    else:
        show_data_science_df
    return show_data_science_df.to_json(orient='records')


# function for MLOps resume results
def mlopsres(noob, m1):
    # global gensa
    mlops_df = noob[noob['MLOps Recommendation'].str.contains('Recommended for MLOps profile')]
    show_mlops_df = mlops_df[
        ['File Name','Name','MLOps Recommendation', 'GenAI', 'Machine learning', 'Deep Learning','MLOps', 'Cloud', 
         'Programming Language', 'View PDF']]
    if m1 == 0:
        print("No MLOps related resumes found !")
    else:
        show_mlops_df
    return show_mlops_df.to_json(orient='records')


def MLres(noob, m1):
    # global dssa
    ml_df = noob[
        noob['DS Recommendation'].str.contains('Recommended for Machine Learning profile')]
    show_ml_df = ml_df[
        ['File Name','Name','DS Recommendation', 'MLOps Recommendation','GenAI', 'Machine learning',  'Deep Learning','MLOps',
          'Cloud', 'Programming Language', 'View PDF']]
    if m1 == 0:
        print("No Machine Learning related resumes found !")
    else:
        show_ml_df
    return show_ml_df.to_json(orient='records')


def DLres(noob, m1):
    # global dssa
    dl_df = noob[
        noob['DS Recommendation'].str.contains('Recommended for Deep Learning profile')]
    show_dl_df = dl_df[
        ['File Name', 'Name', 'DS Recommendation', 'MLOps Recommendation','GenAI', 'Machine learning',  'Deep Learning',
         'MLOps', 'Cloud', 'Programming Language', 'View PDF']]
    if m1 == 0:
        print("No Deep Learning related resumes found !")
    else:
        show_dl_df
    return show_dl_df.to_json(orient='records')


# function for alternate resume results
def altrres(noob, m1):
    global altsa
    alternate_df = noob[noob['Alternate Recommendation'].str.contains('Yes')]
    show_alternate_df = alternate_df[
        [ 'File Name','Name','GenAI Recommendation', 'MLOps Recommendation', 'GenAI', 'Machine learning', 'Deep Learning','MLOps', 
         'Cloud', 'Programming Language', 'View PDF']]
    if m1 == 0:
        print("No GenAI related resumes found !")
    else:
        show_alternate_df
    return show_alternate_df.to_json(orient='records')


# Function for column colors
def color_columns(s):
    df = pd.DataFrame('', index=s.index, columns=s.columns)
    df[['File Name', 'Name']] = 'background-color: #FFFF00'
    df['Programming Language'] = 'background-color: orange'

    return df

def cleanup_folder(folder_path):
    """Deletes the specified folder and its contents."""
    try:
        if os.path.exists(folder_path):
            shutil.rmtree(folder_path)
            print(f"Deleted folder: {folder_path}")
    except Exception as e:
        print(f"Error during folder cleanup: {e}")
 
 
#JD Codes begin here
def normalize_json_output(json_output):
    """
    Normalize the keys and convert list values to comma-separated strings.
    """
    normalized_output = {}
    for key, value in json_output.items():
        normalized_key = key.lower().replace(" ", "_")
        if isinstance(value, list):
            normalized_output[normalized_key] = ', '.join(value)
        else:
            normalized_output[normalized_key] = value
    return normalized_output
 
 
def clean_json(json_string):
    # Remove any text before the first '{' and after the last '}'
    cleaned_json = re.sub(r'^[^{]*|[^}]*$', '', json_string)
    return cleaned_json
 
 

#This below function helps in generating eligiblity questions
jd_questions = " "
def question_generator(jd_dict,text_inp,GOOGLE_API_KEY):
  genai.configure(api_key=GOOGLE_API_KEY)
  model = genai.GenerativeModel('gemini-1.5-flash')
  json_outputs = []
  df = pd.DataFrame()
#   jd_dict = load_resumes_as_dict(file_input)
#   jd_dict.update({"additional_input": text_inp})
  #print(jd_dict)
  Format = '''
            {Filename1:[Questions1, Questions2,...], Filname2: [Questions1, Questions2,...]}
            '''
  jd_title = " "
  for file_name in jd_dict:
    jd_title += os.path.basename(file_name) + ", "
 
  prompt_temp = f"""
                    Based on the provided Job Description in {jd_dict} and {text_inp},
                    Let's think step by step.
                    here job description is in the form of dictionary where key is filename of the job description and value is the extracted text of job description.
 
 
                    Consider and generate each and every point in {jd_dict} and {text_inp} as a yes and no type question for each {jd_title} from the job description to test the eligiblity of the candidate.
                    Note, each point in the job description should be taken in consideration for above qusestions, no point has to be missed.
                    Do not ask the same question.
                   
                  """
 
  response_text = model.generate_content(prompt_temp, generation_config=genai.types.GenerationConfig(
              candidate_count=1,
              temperature=0.3))
 
  jd_questions = response_text.text
 
  prompt_temp1 = f"""
                    Based on the provided Job Description in {jd_dict},
                    Let's think step by step.
                    here job description is in the form of dictionary where key is filename of the job description and value is the extracted text of job description.
 
 
                    Consider and generate each and every point in {jd_dict} as a yes and no type question for each {jd_title} from the job description to test the eligiblity of the candidate.
                    Note, each point in the job description should be taken in consideration for above questions, no point has to be missed.
                    Do not ask the same question.
 
                    Provide the questions in one json format only. The key-value pair of json are as follows:
                     {Format}
                    Don't add up things by yourself, just follow my instruction and the key-value pair
                    should be under one json format. Give response in json format only, do not provide values in list.
                   
                  """
 
  response_text = model.generate_content(prompt_temp1, generation_config=genai.types.GenerationConfig(
              candidate_count=1,
              temperature=0.3))
 
  jd_ques = response_text.text
 
  #print(jd_ques)
 
  cleaned_paragraph = clean_json(jd_ques)
  #print(cleaned_paragraph)
  json_data = json.loads(cleaned_paragraph)
  return json_data
 
 
#This functon is for JD bucket
main_json = None
def driver_code(jd_dict, resume_text, text_inp,GOOGLE_API_KEY):
  genai.configure(api_key=GOOGLE_API_KEY)
  model = genai.GenerativeModel('gemini-1.5-flash')
  json_outputs = []
  df = pd.DataFrame()
#   jd_dict = load_resumes_as_dict(file_input)
#   jd_dict.update({"additional_input": text_inp})
  jd_title = " "
  for file_name in jd_dict:
    jd_title += os.path.basename(file_name) + ", "
 
#   resume_text = load_resumes_as_dict(resume_input)
 
  outputs = {}
  for file_name, textt in resume_text.items():
 
    bestfit_prompt = f"""
                        Based on the the job description following questions are generated
                        {jd_questions}. Treat each contents of {jd_title} respective keys for the JSON.
                        Analyze the resume in {textt} and answer yes and no for each question present in {jd_questions} for each {jd_title}.
                        For each {jd_title}, calculate percent with the formula as followed:
                        percent = ((total no. of questions with 'yes' answer)/ (total number of questions)) * 100
 
                        Follow the below instructions and follow them with importance
                        NOTE: How to decide yes and no for each content of {jd_title}?
                              If percent of a content is greater than or equal to 80 and highest of all other
                              percent of contents of {jd_title}, then write Yes under that content or else write NO.
                        Also write answer of each question.
                        Note: 1)Ensure accuracy while providing the response.
                              2)Always verify your answer.
                              3) Use provided jd questions only.
                              4) Don't return {jd_questions} and their answers in response.
                              5) Don't provide "YES" response more than one for content of {jd_title}.
                     """
    score_prompt = f"""
                        Create the following 7 keys only. Make sure key names are same as provided below.
                        =>Name: This will contain candidate's Name that is extracted from the text of user input.
                        {bestfit_prompt}
                        =>Recommended: This key should contain Yes and No response. If the percent of any
                                       content of {jd_title} present in {bestfit_prompt} is
                                       greater than or equal to 80 then provide "Yes" or else provide "No". Please recommend the candidate if percent is 80.        
                        =>Experience: This will contain candidate's relevant experience in years only for {jd_title}.
                        =>Percent:
                        =>Other Skills: Extract all the relevant skills as per JD candidate has.
                        =>Gen AI Skills: Extract GenAI skills candidate has.
                       
                       
                        Guidelines:
                        -Please provide all the keys that is written above. Always provide value for 'Percent', do not provide 'unable to process'.
                        - Please don't return any extra key other than the keys mentioned above in output of the JSON.
                        - Note: Please note that all keys should be mandetorily present in your response.
                        - Very important: Please do not make up anything. If the information of a required field is not available, output ‘Unable to process’ for it.
                        - Output in JSON format. The JSON should contain all the above keys.
                        - Please don't return any scores value in output of the JSON.
                        - Do not use any special characters in the response. Use only "\n" for new line
                        - Please do not repeat any output.
                        - Please do not use "\xa0" as escape sequence in the response.
                        - Please follow the below format for providing response(Do not provide "```" in the response):
                         
                  """
    response_res = model.generate_content(score_prompt, generation_config=genai.types.GenerationConfig(
              candidate_count=1,
              temperature=0))
 
    final_response = response_res.text
   
    cleaned_paragraph = clean_json(final_response)
    cleaned_paragraph= cleaned_paragraph.replace("Unable to process","No Skills")
    #print(cleaned_paragraph)
    raw_data = json.loads(cleaned_paragraph)
    # print(raw_data)
    outputs[file_name] = raw_data
   
  main_json = outputs
  return outputs
 
def reco_count():
    global main_json
    data = main_json
    # Count the number of "Recommended" values that are "Yes"
    yes_count = sum(1 for item in data.values() if item.get("Recommended") == "Yes")
    print(yes_count)
    return yes_count
 