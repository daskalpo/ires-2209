from flask import Flask,json, render_template, session, request, redirect, url_for, send_file,jsonify, g
import os
from flask_cors import CORS
from app import prathambox
from app import dutiyabox
from app import tritiyabox
from app import userInput_summary
from app import fourthbox
# from app import dataFr
from app import genAICount, dsCount, altrCount, mlCount, dlCount
from app import resuSumm
from app import genAIres, DSres, altrres, MLres, DLres,mlopsres
from app import color_columns
from app import download_button_excel
from io import BytesIO
import pandas as pd
from datetime import datetime
import secrets
from app import extract_text_from_pdf,extract_text_from_docx,extract_text_from_doc
from flask import Response
from io import BytesIO,StringIO
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
# from app import no_of_files
from app import question_generator
from app import driver_code
from app import reco_count
import boto3
from boto3.s3.transfer import TransferConfig, S3Transfer
import threading
from concurrent.futures import ThreadPoolExecutor
import uuid
from concurrent.futures import ThreadPoolExecutor, as_completed
import concurrent
from botocore.exceptions import NoCredentialsError
import base64
import json
import aspose.words as aw
import chardet
from flask_socketio import SocketIO
from app import clean_json

import numpy as np
import faiss
from sentence_transformers import SentenceTransformer
import json
import pandas as pd
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Alignment, PatternFill,Font
import spacy

import base64
import os
import fitz
import pandas as pd
import streamlit as st
import docx2txt
import re
import google.generativeai as genai
# from pyresparser import ResumeParser
from io import BytesIO
from datetime import datetime
from threading import Thread
from queue import Queue
from bokeh.models import ColumnDataSource, CustomJS
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
#from win32com import client as wc
import time
import shutil
import tempfile
from flask import Flask, session
from flask_session import Session
# import paramiko
# from scp import SCPClient
#import textract
import subprocess
from dotenv import load_dotenv


load_dotenv()
model = genai.GenerativeModel('gemini-1.5-flash')
#model = genai.GenerativeModel('gemini-1.5-flash-8b')

nlp = spacy.load("en_core_web_sm")

# Initializing flask app.,
app = Flask(__name__)
app.secret_key = secrets.token_hex(16)
app.config['SESSION_TYPE'] = 'filesystem'
Session(app)
CORS(app, resources={r"/backend/*": {"origins":"*"}},supports_credentials=True)
socketio = SocketIO(app, cors_allowed_origins="*")

#flask_app begins here
#global selectfield, additionalSkills, resumePath
#Route for fetching the user-input from the frontend.
#@app.route('/backend/input', methods = ['GET'])
# global selectfield
# selectfield = "GenAI/DataScience"
# #selectfield = request.json['pass']
# global additionalSkills
# additionalSkills = "PowerBI"
# #additionalSkills = request.json['pass']
# global resumePath
# resumePath = r"C:\Users\2160945\Documents\GenAI POC's\AITalentScan\MLOps Resumes"
#resumePath = request.json['pass']
# return selectfield, additionalSkills, resumePath

# @app.route('/backend/input', methods = ['POST','GET'])
# def input():
#     global selectfield, additionalSkills,resumePath
#     # json_data = request.data
#     # parsed_data = json.loads(json_data)
#     # selectfield = parsed_data['selectfield']
#     selectfield = "GenAI/DataScience"
#     # additionalSkills = parsed_data['additionalSkillset']
#     additionalSkills = "PowerBI"
#     # resumePath = parsed_data['resumepath']
#     resumePath = r"C:\Users\2160945\Documents\GenAI POC's\AITalentScan\MLOps Resumes"
#     print(selectfield, additionalSkills, resumePath)
#     return jsonify(selectfield=selectfield, additionalSkills=additionalSkills, resumePath=resumePath)

# from flask import session

def schedule_lambda(session_id):
    client = boto3.client('events',aws_access_key_id='AKIAVAYAJDYX7MDDJKGL',
        aws_secret_access_key='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',
        region_name='us-east-1')
    lambda_client = boto3.client('lambda', aws_access_key_id='AKIAVAYAJDYX7MDDJKGL',
        aws_secret_access_key='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',
        region_name='us-east-1')

    lambda_function_name = 'Practice_ires'

    try:
        response = lambda_client.get_function(FunctionName=lambda_function_name)

        rule_name = f"DeleteS3Bucket_{session_id}"
        lambda_arn = response['Configuration']['FunctionArn']

        # Add permission to lambda to allow invocation from CloudWatch Events
        lambda_client.add_permission(
            FunctionName=lambda_function_name,
            StatementId=f'{rule_name}-Event',
            Action='lambda:InvokeFunction',
            Principal='events.amazonaws.com',
            SourceArn=f'arn:aws:events:{'us-east-1'}:{345208593967}:rule/{rule_name}'
        )

        # Create a Cloudwatch rule
        response = client.put_rule(
            Name=rule_name,
            ScheduleExpression='rate(30 minutes)',
            State='ENABLED'
        )

        # Add target to the rule
        response = client.put_targets(
            Rule=rule_name,
            Targets=[
                {
                    'Id': '1',
                    'Arn': lambda_arn,
                    'Input': json.dumps({'session_id': session_id})
                },
            ]
        )
    except Exception as e:
        print(f"Failed to schedule lambda function. Reason: {e}")


@app.route('/backend/start', methods=['POST', 'GET'])
def start_session():
    session['id'] = str(uuid.uuid4())  # Generate a unique session id
    session['api_key']= ''
    session['selectfield'] = ''
    session['additionalSkills'] = ''
    session['resumePath'] = ''
    session['noob'] = None
    session['questions_dict'] = {}
    session['files']= ''
    session['files_jd']= ''
    session['resume_ner_data_list'] = [] #each element is a dict of the resume's NER output
    session['dataframe_list'] = []
    session['job_title_list'] = []
    session['summary_data'] = []
    session['jd_ner_list']=[]
    session['fetched_dataframes'] = []
    print(session['id'])
    print(session)
    return  "Session started"

@app.route('/backend/input', methods=['POST', 'GET'])
def input():
    json_data = request.data
    parsed_data = json.loads(json_data)
    session['selectfield'] = parsed_data['selectfield']
    session['additionalSkills'] = parsed_data['additionalSkillset']
    session['resumePath'] = parsed_data['resumepath']
    print(session['selectfield'], session['additionalSkills'], session['resumePath'])
    return json_data
    
s3_client = boto3.client(
        's3',
        aws_access_key_id='AKIAVAYAJDYX7MDDJKGL',
        aws_secret_access_key='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',
        region_name='us-east-1'
    )
 
s3_resource = boto3.resource(
        's3',
        aws_access_key_id='AKIAVAYAJDYX7MDDJKGL',
        aws_secret_access_key='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',
        region_name='us-east-1'
    )
 
bucket= 'ires-v1'
log='log.json'

@app.route('/backend/upload', methods=['POST', 'GET'])
def upload_files():
    session['selectfield'] = request.form.get('selectfield')
    session['additionalSkills'] = request.form.get('additionalSkillset')
    session['api_key'] = request.form.get('apikey')
    # session['difficultylevel'] = request.form.get('difficultylevel')
    files = request.files.getlist('files')
    files  = [f for f in files if f.filename.count('/') < 2]
    session['files']= [f.filename for f in files if f.filename.count('/') < 2]
    #print(session['files'])
    print(session)
    #session['files'] = files
    answer1,count_num = userInput_summary(session['selectfield'],session['additionalSkills'],files )

    resume_dictionary = {}
    download_dictionary = {}

    # Create a ThreadPoolExecutor for multi-threading
    with ThreadPoolExecutor(max_workers=10) as executor:
        future_to_file = {executor.submit(upload_file_to_s3, file, session['id']): file for file in files}

    for future in concurrent.futures.as_completed(future_to_file):
        file = future_to_file[future]
        try:
            result = future.result()
        except Exception as exc:
            print(f'{file} generated an exception: {exc}')
        else:
            resume_dictionary.update(result[0])
            download_dictionary.update(result[1])
            
    #print(resume_dictionary)
    #print(download_dictionary)

    save_dictionary_to_s3(resume_dictionary, session['id'], 'resume')
    save_dictionary_to_s3(download_dictionary, session['id'], 'download')

    return jsonify({'message': 'Files successfully uploaded', 'selectfield': session['selectfield'], 'additionalSkillset': session['additionalSkills'], 'numberofresumes': count_num})


def upload_file_to_s3(file, session_id):
    resume_dictionary = {}
    download_dictionary = {}

    try:
        folder_name, file_name = file.filename.split('/', 1)
        #print(folder_name,file_name)
        s3_resource.Bucket('ires-v1').upload_fileobj(file, f"""{session_id}/{folder_name}/{file_name}""")
    except Exception as e:
        print(e)

    file_object = s3_resource.Object(bucket, f"""{session_id}/{folder_name}/{file_name}""")
    file_content2 = file_object.get()['Body'].read()
    

    if file.filename.endswith('.pdf'):
        file_data = extract_text_from_pdf(BytesIO(file_content2))
        #print(file_data)
    elif file.filename.endswith('.doc'):
        
        file_data = extract_text_from_doc(BytesIO(file_content2))
        #print(file_data)
    elif file.filename.endswith('.docx'):
        file_data = extract_text_from_docx(BytesIO(file_content2))
        #print(file_data)
    else:
        try:
            file_data = file.read().decode('utf-8')
        except UnicodeDecodeError:
            file_data = file.read().decode('latin-1')
    
    resume_dictionary[file.filename] = file_data  
    download_dictionary[file.filename] = base64.b64encode(file_content2).decode("utf-8")
    

    return resume_dictionary, download_dictionary
    
def save_dictionary_to_s3(dictionary, session_id, file_name_suffix):
    try:
        s3_resource.Object('ires-v1', f'{session_id}/{session_id}_{file_name_suffix}.json').put(
            Body=json.dumps(dictionary))
    except NoCredentialsError:
        print("No AWS credentials were found.")
    except Exception as e:
        print(e)

@app.route('/backend/user_summary')
def user_summary():
    answer1 = userInput_summary(session['selectfield'], session['additionalSkills'], session['resumePath'])
    return answer1

@app.route('/backend/res_summary')
def res_summary():
    global json_data, genai_ui_json, ds_ui_json, additional_ui_json, table_ui
    x = prathambox(session['selectfield'])
    y = fourthbox(session['additionalSkills'])
    session['noob'], session['questions_dict'] = dataFr(x, y, session['resume_dictionary'])
    gcount = genAICount(session['noob'])
    dcount = dsCount(session['noob'])
    acount = altrCount(session['noob'])
    mlcount = mlCount(session['noob'])
    dlcount = dlCount(session['noob'])
    json_data, genai_ui_json, ds_ui_json, ml_ui_json, dl_ui_json, additional_ui_json, table_ui, mlops_ui_json = resuSumm(session['noob'], session['selectfield'], session['additionalSkills'])
    return '[{},{},{},{},{},{},{},{}]'.format(json_data, genai_ui_json, ds_ui_json, additional_ui_json, table_ui, ml_ui_json, dl_ui_json, mlops_ui_json)

@app.route('/backend/get_resume_count', methods=['GET'])
def get_resume_count():
    try:
        # Assuming you have a way to determine the number of resumes
        resume_count = len(session['files'])  # Example: counting the number of files in the session
        return jsonify({'count': resume_count}), 200
    except Exception as e:
        print(f'Error getting resume count: {e}')
        return jsonify({'error': 'Unable to get resume count'}), 500
    
@app.route('/backend/new_res_summary')
def new_res_summary():

    resume_count = len(session['files'])
    
    global json_data, genai_ui_json, ds_ui_json, additional_ui_json, table_ui
    x = prathambox(session['selectfield'])
    y = fourthbox(session['additionalSkills'])
    resume_dictionary=fetch_resume_dict_from_s3(session['id'])
    print(resume_dictionary)
    noob,questions_dict = dataFr(x, y, resume_dictionary,session['api_key'])
    save_dataframe_to_s3(noob,bucket , f'{session["id"]}/noob.csv')
    save_dictionary_to_s3(questions_dict, session['id'], 'questions')
    session['gcount'] = str(genAICount(noob))
    session['dcount'] = str(dsCount(noob))
    session['acount'] = str(altrCount(noob))
    session['mlcount'] = str(mlCount(noob))
    session['dlcount'] = str(dlCount(noob))
    # if request.method == 'POST':
    json_data, genai_ui_json, ds_ui_json, ml_ui_json, dl_ui_json, additional_ui_json, table_ui, mlops_ui_json= resuSumm(noob,
                                                                                                          session['selectfield'], session['additionalSkills'])
                                                                                                          
    fetch_append_update_json_in_s3(bucket, log, session['id'], "{resume_count} resumes Run successfully")                                                                                                      
                                                                                                       
    # Delete files after processingq
    with ThreadPoolExecutor(max_workers=10) as executor:
        
        for file in session['files'] :
            executor.submit(delete_file_from_s3, file,session['id'])
            
    schedule_lambda(session['id'])
    
    return '[{},{},{},{},{},{},{},{}]'.format(json_data, genai_ui_json, ds_ui_json, additional_ui_json, table_ui,
                                           ml_ui_json, dl_ui_json,mlops_ui_json)
                                           
def delete_file_from_s3(file,session_id):
    try:
        folder_name, file_name = file.split('/', 1)
        s3_resource.Object('ires-v1', f"""{session_id}/{folder_name}/{file_name}""").delete()
    except Exception as e:
        print(e)

def fetch_resume_dict_from_s3(session_id):
    
    file_name = f'{session_id}/{session_id}_resume.json'
    s3_client.download_file('ires-v1', file_name, 'local_resume.json')

    with open('local_resume.json') as json_file:
        resume_dict = json.load(json_file)
    return resume_dict
    
def fetch_JD_dict_from_s3(session_id):
    
    file_name = f'{session_id}/{session_id}_JD.json'
    s3_client.download_file('ires-v1', file_name, 'local_JD.json')

    with open('local_JD.json') as json_file:
        JD_dict = json.load(json_file)
    return JD_dict

def fetch_questions_dict_from_s3(session_id):
    
    file_name = f'{session_id}/{session_id}_questions.json'
    s3_client.download_file('ires-v1', file_name, 'local_questions.json')

    with open('local_questions.json') as json_file:
        questions_dict = json.load(json_file)
    return questions_dict

def save_dataframe_to_s3(df, bucket, file_path):
    
    csv_buffer = StringIO()
    df.to_csv(csv_buffer)
    s3_client.put_object(Bucket=bucket, Body=csv_buffer.getvalue(), Key=file_path)

def fetch_dataframe_from_s3(bucket, file_path):
   
    obj = s3_client.get_object(Bucket=bucket, Key=file_path)
    data = obj['Body'].read().decode('utf-8')
    df = pd.read_csv(StringIO(data))
    return df

@app.route('/backend/genAI_sum')
def genAI_sum():
    noob = fetch_dataframe_from_s3(bucket, f'{session["id"]}/noob.csv')
    gcount = int(session['gcount'])
    answer3 = genAIres(noob, gcount)
    return answer3

@app.route('/backend/dataScience_sum')
def dataScience_sum():
    noob = fetch_dataframe_from_s3(bucket, f'{session["id"]}/noob.csv')
    dcount = int(session['dcount'])
    answer4 = DSres(noob, dcount)
    return answer4

@app.route('/backend/MachineLearning_sum')
def MachineLearning_sum():
    noob = fetch_dataframe_from_s3(bucket, f'{session["id"]}/noob.csv')
    mlcount = int(session['mlcount'])
    answer6 = MLres(noob, mlcount)
    return answer6

@app.route('/backend/DeepLearning_sum')
def DeepLearning_sum():
    noob = fetch_dataframe_from_s3(bucket, f'{session["id"]}/noob.csv')
    dlcount = int(session['dlcount'])
    answer7 = DLres(noob, dlcount)
    return answer7

@app.route('/backend/mlops_sum')
def mlops_sum():
    noob = fetch_dataframe_from_s3(bucket, f'{session["id"]}/noob.csv')
    dcount = int(session['dcount'])
    answer8 = mlopsres(noob, dcount)
    return answer8

@app.route('/backend/pref_sum')
def preferred_sum():
    noob = fetch_dataframe_from_s3(bucket, f'{session["id"]}/noob.csv')
    acount = int(session['acount'])
    answer5 = altrres(noob, acount)
    return answer5
    
    
@app.route('/backend/download', methods = ['POST','GET'])
def download_file():
    
    file_name = request.args.get('file_name')
    #resumePath2 = resumePath + '/'
    #file_path = os.path.join(resumePath2, file_name)
    #if os.path.exists(file_path):
        #return send_file(file_path, as_attachment=True)
    #else:
        #return 'File Not Found'
    # Read the download dictionary from the S3 bucket
    session_id=session['id']
    print(session_id)
    s3_object = s3_resource.Object('ires-v1', f'{session["id"]}/{session["id"]}_download.json')
    download_dictionary_str = s3_object.get()['Body'].read().decode('utf-8')
    download_dictionary = json.loads(download_dictionary_str)

    if file_name in download_dictionary:
        file_content_base64 = download_dictionary[file_name]

        # Decode the base64 string back to bytes
        file_content = base64.b64decode(file_content_base64)
        # Create BytesIO Object and write the file content into it
        byte_io = BytesIO()
        byte_io.write(file_content)
        # Seek back to beginning of file
        byte_io.seek(0)
        return send_file(byte_io, as_attachment=True, mimetype='application/octet-stream',download_name =file_name)
    else:
        return jsonify({'error': 'File Not Found'}), 404

@app.route('/backend/download_questions', methods = ['POST','GET'])
def download_questions():
    
    #print(questions_dict)
    #q_file_name = request.args.get('name') +'_reports.docx'
    #questionPath= resumePath + '/Generated_Docx/'
    #file_path = os.path.join(questionPath, q_file_name)
    #if os.path.exists(file_path):
        #return send_file(file_path, as_attachment=True)
    #else:
        #return 'File Not Found'
    file_name = request.args.get('name')
    #print(file_name)
    
    questions_dict=fetch_questions_dict_from_s3(session['id'])
    
    if file_name in questions_dict:
        file_content = questions_dict[file_name]
        #print(file_content)
        # # Parse response and extract data
        # lines = file_content.strip().split('\n')
        # candidate_name = lines[2].strip().split("|")[1].strip()
        # questions_and_scores = [line.strip() for line in lines[4:]]

        # Create a new Document
        doc = Document()

        # sections = doc.sections
        # for section in sections:
        #     section.left_margin = Pt(35)  # Adjust left margin as needed
        #     section.right_margin = Pt(35)  # Adjust right margin as needed
        #     section.top_margin = Pt(35)  # Adjust top margin as needed
        #     section.bottom_margin = Pt(35)  # Adjust bottom margin as needed

        # timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        # # doc.add_paragraph(f"Report Generated At: {timestamp}",style = 'Heading 1')

        # timestamp_paragraph = doc.add_paragraph(f"Report Generated At: {timestamp}")
        # timestamp_run = timestamp_paragraph.runs[0]
        # timestamp_run.bold = True
        # timestamp_run.font.size = Pt(11)  # Adjust the font size as needed

        # # Add candidate name
        # doc.add_heading('Candidate Name:', level=1)
        # doc.add_paragraph(candidate_name)
        doc.add_paragraph(file_content)

        # # Add table
        # table = doc.add_table(rows=1, cols=2)
        # table.style = 'Table Grid'

        # table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # hdr_cells = table.rows[0].cells
        # hdr_cells[0].text = 'Questions'
        # hdr_cells[1].text = 'Score'

        # # Make column names bold
        # for cell in hdr_cells:
        #     for paragraph in cell.paragraphs:
        #         for run in paragraph.runs:
        #             run.bold = True

        # # Populate table
        # for line in questions_and_scores:
        #     parts = line.split('|')
        #     if len(parts) > 1 and parts[1].strip() != '---':  # Check if the line is not just '---'
        #         question = parts[1].strip()
        #         score = parts[2].strip() if len(
        #             parts) > 2 else ""  # Handle the case where score is not provided
        #         row_cells = table.add_row().cells
        #         row_cells[0].text = question
        #         row_cells[1].text = score

        #         # Add space after each question
        #         row_cells[0].paragraphs[0].runs[-1].add_break()

        #         # for cell in row_cells:
        #         #     for paragraph in cell.paragraphs:
        #         #         paragraph.space_after = Pt(10)

        # # Adjust row height to fit content
        # for row in table.rows:
        #     for cell in row.cells:
        #         for paragraph in cell.paragraphs:
        #             paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Set alignment to left
        #             paragraph.space_after = Pt(12)  # Add space after each paragraph (adjust as needed)
        #         cell.height = Pt(60)  # Adjust row height to fit content

        # # Add a border around the entire document
        # for paragraph in doc.paragraphs:
        #     paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        #     paragraph.space_after = Pt(12)
        #     for run in paragraph.runs:
        #         run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black for better visibility
        #         run.font.name = 'Cambria'
        #         run.font.size = Pt(11)

        # # Add space after the headings "Questions" and "Score"
        # for cell in table.rows[0].cells:
        #     cell.paragraphs[0].runs[-1].add_break()

        # for cell in table.columns[1].cells:
        #     cell.width = Pt(30)  # Adjust cell width as needed

        # for cell in table.columns[0].cells:
        #     cell.width = Pt(700)

        # table.border_collapse = True
        
        # Save the docx file to an in-memory BytesIO object
        file_stream = BytesIO()
        doc.save(file_stream)
    
        # Reset the file stream position to the beginning
        file_stream.seek(0)
        
        return send_file(file_stream, as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',download_name ='questions.docx')
    else:
        return jsonify({'error': 'File Not Found'}), 404
        
#Route for download link of excel
@app.route('/backend/download_excel', methods=['GET', 'POST'])
def downExcel():
    noob = fetch_dataframe_from_s3(bucket, f'{session["id"]}/noob.csv')
    print(noob)
    noob1 = noob.copy()
    noob1 = noob1.fillna("")
    noob1 = noob1[['File Name', 'Name', 'Programming Language','GenAI Recommendation','DS Recommendation', 'MLOps Recommendation', 'Recent Experience', 'Machine learning',
                   'Deep Learning', 'MLOps', 'Cloud', 'GenAI', 'Additional Skills', 'Alternate Recommendation',
                   'Projects', 'Candidate Summary', 'Email', 'Phone']]
    styled_df = noob1.style.apply(color_columns, axis=None)
    styled_df = styled_df.apply(
        lambda row: [
            'background-color: #FFFF00' if 'Borderline' in val else 'background-color: #00FF00' if 'Yes' in val else 'background-color: #FF0000' if 'No' in val else ''
            for val in row],
        axis=1,
        subset=['GenAI Recommendation','DS Recommendation', 'MLOps Recommendation']
    )
    info = {'Steps':['This are the steps you should follow',
                 'Step1','Step2']
        }
    info_df = pd.DataFrame(info)
    excel_buffer = BytesIO()
    with pd.ExcelWriter(excel_buffer, engine="xlsxwriter") as writer:
        styled_df.to_excel(writer, sheet_name="Sheet1", index=False)
        workbook = writer.book
        worksheet = writer.sheets['Sheet1']

        header_style = workbook.add_format({'bg_color': '#6495ED', 'border': 1})
        for col_num, value in enumerate(noob1.columns.values):
            worksheet.write(0, col_num, value, header_style)

        data_format = workbook.add_format({'border': 1})
        for row_num, row in enumerate(noob1.values, start=1):
            for col_num, value in enumerate(row):
                worksheet.write(row_num, col_num, value, data_format)

        column_widths = {'File Name': 30, 'Name': 15, 'Programming Language': 15, 'Recent Experience': 15,
                         'Machine learning': 30,
                         'Deep Learning': 30, 'MLOps': 20, 'Cloud': 20, 'GenAI': 20, 'Additional Skills': 15,
                         'Alternate Recommendation': 15,
                         'Projects': 30, 'Candidate Summary': 30, 'GenAI Recommendation': 30, 'DS Recommendation': 30, 'MLOps Recommendation': 35, 'Email': 20, 'Phone': 20}
        for col, width in column_widths.items():
            column_idx = styled_df.columns.get_loc(col)
            worksheet.set_column(column_idx, column_idx, width)

        column_colors = {'File Name': '#DCE6F1', 'Name': '#DCE6F1', 'Programming Language': '#DCE6F1',
                         'Recent Experience': '#DCE6F1', 'Machine learning': '#FDE9D9',
                         'Deep Learning': '#FDE9D9', 'MLOps': '#FDE9D9', 'Cloud': '#FDE9D9', 'GenAI': '#FDE9D9',
                         'Additional Skills': '#FDE9D9', 'Candidate Summary': '#E6B8B7', 'GenAI Recommendation': '#B7DEE8','DS Recommendation': '#B7DEE8',
                         'MLOps Recommendation': '#B7DEE8', 'Email': '#DCE6F1', 'Phone': '#DCE6F1'}
        for col, color in column_colors.items():
            column_idx = noob1.columns.get_loc(col)
            start_row = 1
            end_row = len(noob1)
            for row_num in range(start_row, end_row + 1):
                cell_value = noob1.iloc[row_num - 1, column_idx]
                worksheet.write(row_num, column_idx, cell_value, workbook.add_format({'bg_color': color, 'border': 1}))
        info_df.to_excel(writer, sheet_name="Sheet2", index=False)
    excel_buffer.seek(0)
    timestamp = datetime.now().strftime('%Y/%m/%d/%H:%M:%S')
    return send_file(excel_buffer, download_name=f'iReS_{timestamp}.xlsx', as_attachment=True)
    
###JD Section###   
    
@app.route('/backend/JD_upload_new', methods=['POST', 'GET'])
def upload_jdfiles():    
    try:
        num_files=0
        num_jd_files=0
        
        files = request.files.getlist('jdfile')
        # print(files)
        jd_files = request.files.getlist('jd_files')
        session['api_key'] = request.form.get('apikey')
        # Filter out directories
        files = [f for f in files if f.filename.count('/') < 2]
        session['jdfile'] = [f.filename for f in files if f.filename.count('/') < 2]
        # print(files)
        num_files = len(session['jdfile'])
    
        jd_files = [f for f in jd_files if f.filename.count('/') < 2]
        session['jd_files'] = [f.filename for f in jd_files if f.filename.count('/') < 2]
        print(jd_files)
        num_jd_files = len(session['jd_files'])
    
        #Store file data in session
        resume_dictionary={}
        download_dictionary={}
        #session['files'] = []

        # Create a ThreadPoolExecutor for multi-threading
        with ThreadPoolExecutor(max_workers=10) as executor:
            future_to_file = {executor.submit(upload_file_to_s3, file, session['id']): file for file in files}

        for future in concurrent.futures.as_completed(future_to_file):
            file = future_to_file[future]
            try:
                result = future.result()
            except Exception as exc:
                print(f'{file} generated an exception: {exc}')
            else:
                resume_dictionary.update(result[0])
                download_dictionary.update(result[1])

        save_dictionary_to_s3(resume_dictionary, session['id'], 'resume')
        save_dictionary_to_s3(download_dictionary, session['id'], 'download')   
        
        #print(resume_dictionary)
        jd_dictionary = {}

        with ThreadPoolExecutor(max_workers=10) as executor:
            future_to_file = {executor.submit(upload_file_to_s3, file, session['id']): file for file in jd_files}

        for future in concurrent.futures.as_completed(future_to_file):
            file = future_to_file[future]
            try:
                result = future.result()
            except Exception as exc:
                print(f'{file} generated an exception: {exc}')
            else:
                jd_dictionary.update(result[0])
                

        save_dictionary_to_s3(jd_dictionary, session['id'], 'JD')

        print(resume_dictionary)
        # ques = question_generator(jd_path, text_inp)
        # ques = question_generator(jd_dictionary,text_input,session['api_key'])
        ques = " "
    
        return jsonify({'message': 'Files successfully uploaded', 'numberofresumes': num_files, 'numberofJD': num_jd_files,'questions':ques}) 
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({'error': str(e)}), 500
        
 
 
text_input = ""
@app.route('/backend/jd_question', methods = ['POST','GET'])
def jd_question():
    try:
        # ques = question_generator(jd_path, text_inp)
        ques = question_generator(jd_dictionary,text_input,session['api_key'])
        print(ques)
        return jsonify(ques)
    except Exception as e:
       print(f"Error during file upload: {e}")
       return jsonify({"error": str(e)}), 500
 
@app.route('/backend/get_jd_count', methods=['GET'])
def get_jd_count():
    try:
        # Count the number of job descriptions stored in the session
        # jd_count = len(session.get('jdfile', []))  # Use session.get to avoid KeyError
        jd_count = len(session['files_jd'])
        #print('--------ytghdcnjhv jd_count ---------')
        print(jd_count)
        #print('--------ytghdcnjhv jd_count ---------')
        return jsonify({'count': jd_count}), 200
    except Exception as e:
        print(f'Error getting job description count: {e}')
        return jsonify({'error': 'Unable to get job description count'}), 500

@app.route('/backend/jd_bucket', methods = ['POST','GET'])
def jd_bucket():
    
    try:
        # summary_data = []
        # bucket = driver_code(jd_path, res_folder, text_inp)
        resume_dictionary=fetch_resume_dict_from_s3(session['id'])
        jd_dictionary=fetch_JD_dict_from_s3(session['id'])
        # bucket = driver_code(jd_dictionary,resume_dictionary,text_input,session['api_key'])
        session['jd_ner']= jd_entity(jd_dictionary)
        # bucket = resume_jd_score(resume_dictionary,jd_ner,text_input,session['api_key'])
        # print(bucket)
        i=0
        for jd_ner in session['jd_ner']:
            session['job_title_list'].append(jd_ner.get("job_title",""))
            df=resume_jd_score_multi(resume_dictionary,jd_ner,text_input,session['api_key'])
            save_dataframe_to_s3(df,bucket,f"{session["id"]}/{session['job_title_list'][i]}.csv")
            i+=1
            df_json = df.to_json()
            session['dataframe_list'].append(df_json)
            # session['dataframe_list'].append(df)
        
        for i, df in enumerate(session['dataframe_list'], start=0):
            df_json = session['dataframe_list'][i]
            df = pd.read_json(df_json)
            for index, row in df.iterrows():
                summary_row = {
                    'Resume File Name': row['Resume File Name'],
                    'Candidate Name': row['Candidate Name'],
                    'Email': row['Email'],
                    'Phone': row['Phone'],
                    f'{session['job_title_list'][i]} Recommendation': row['Recommendation'],
                    f'{session['job_title_list'][i]} Fitment Score': row['Fitment Score (%)'],

                }
                session['summary_data'].append(summary_row)

        summary_df = pd.DataFrame(session['summary_data'])
        summary_df = summary_df.groupby(['Resume File Name'], as_index=False).first()
        
        save_dataframe_to_s3(summary_df,bucket,f"{session["id"]}/summary.csv")

        Bucket = summary_df.to_json(orient='records')
        print(summary_df)
        
        schedule_lambda(session['id'])
        return Bucket

    except Exception as e:
       print(f"Error during file upload: {e}")
       schedule_lambda(session['id'])
       return jsonify({"error": str(e)}), 500

@app.route('/backend/jd_excel_download', methods = ['POST','GET'])
def jd_download():
    
    i=0
    for job_title in session['job_title_list']:
        df = fetch_dataframe_from_s3(bucket,f"{session["id"]}/{session['job_title_list'][i]}.csv")
        df_json = df.to_json()
        session['fetched_dataframes'].append(df_json)
        i+=1

    # Create a new Excel workbook
    wb = Workbook()

    for i, df in enumerate(session['fetched_dataframes'], start=0):
        df_json = session['fetched_dataframes'][i]
        df = pd.read_json(df_json)
        # Create a new sheet for each DataFrame
        ws = wb.create_sheet(title=f"{session['job_title_list'][i]}")
        for r in dataframe_to_rows(df, index=False, header=True):
            ws.append(r)

        # Find the index of the "Recommendation" column
        recommendation_col_index = None
        for col in ws.iter_cols(min_row=1, max_row=1):
            for cell in col:
                if cell.value == 'Recommendation':
                    recommendation_col_index = cell.column

        # Color the recommendation column based on suitability
        if recommendation_col_index:
            for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=recommendation_col_index, max_col=recommendation_col_index):
                for cell in row:
                    if cell.value in ["Suitable", "Highly Suitable"]:
                        cell.fill = PatternFill(start_color="00FF00", end_color="00FF00", fill_type="solid")  # Green color
                    elif cell.value == "Not Suitable":
                        cell.fill = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")  # Red color

    # Create the summary sheet
    ws_summary = wb.active
    ws_summary.title = "Summary"

    summary_df = fetch_dataframe_from_s3(bucket,f"{session["id"]}/summary.csv")
    # Write the summary DataFrame to the summary sheet
    for r in dataframe_to_rows(summary_df, index=False, header=True):
        ws_summary.append(r)
    # Set fixed color for column headers to blue
    header_fill = PatternFill(start_color="8adbf5", end_color="8adbf5", fill_type="solid")
    header_font = Font(bold=True)
    for ws in wb.worksheets:
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font

    # Set orange color for specific column headers
    orange_fill = PatternFill(start_color="FFA500", end_color="FFA500", fill_type="solid")
    specific_headers = ["relevant skills", "Work Experience - Relevant", "key projects","Qualification & Certifications","Fitment Score (%)"]
    # Normalize specific headers
    normalized_headers = [header.strip().lower() for header in specific_headers]

    for ws in wb.worksheets:
        for cell in ws[1]:
            if cell.value:
                normalized_cell_value = cell.value.strip().lower()
                print(f"Checking header: {normalized_cell_value}")  # Debugging line
                if normalized_cell_value in normalized_headers:
                    cell.fill = orange_fill

    # Adjust column widths dynamically based on the content and set wrap text with max width of 30
    max_width = 30
    for ws in wb.worksheets:
        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter  # Get the column name
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(cell.value)
                except:
                    pass
                cell.alignment = Alignment(wrap_text=True, vertical='top',horizontal='left')
            adjusted_width = min(max_length + 2, max_width)
            ws.column_dimensions[column].width = adjusted_width
   
    # Save the workbook to a BytesIO buffer
    excel_buffer = BytesIO()
    wb.save(excel_buffer)
    excel_buffer.seek(0)
    timestamp = datetime.now().strftime("%d-%m-%Y %H:%M")
    return send_file(excel_buffer, download_name=f"Candidate Fitment Overview & Detailed Recommendations {timestamp}.xlsx", as_attachment=True)
 
recommend = None
@app.route('/backend/no_recommend', methods = ['POST','GET'])
def recommend_count():
    global recommend
    try:
        recommend = reco_count()
        return jsonify({"No. of resumes recommended":f"{recommend}"}), 200
    except Exception as e:
       print(f"Error during file upload: {e}")
       return jsonify({"error": str(e)}), 500
 
@app.route('/backend/nott_recommend', methods = ['POST','GET'])
def not_recommend_count():
    global res_count
    global recommend
    try:
        print(f"res_count: {res_count}, recommend: {recommend}")
        not_recommend = int(res_count) - int(recommend)
        print(not_recommend)
        return jsonify({"No. of resumes recommended":f"{not_recommend}"}), 200
    except Exception as e:
       print(f"Error during file upload: {e}")
       return jsonify({"error": str(e)}), 500

@app.route('/backend/delete_session', methods=['POST', 'GET'])
def delete_session():
    session_id = request.form.get('session_id')
    
    delete_folder_from_s3(session_id)
    return jsonify({'message': 'Session data successfully deleted', 'session_id': session_id})

############################


#load high quality Sentence Transformer model for accurate embeddings
embedder = SentenceTransformer('all-mpnet-base-v2')
embedding_dim = embedder.get_sentence_embedding_dimension()

def create_composite_text(ner_data: dict) -> str:
    """
      Create a composite string from the structured NER data
      You can adjust the concatenation logic based on what is most relevant
    """

    composite_fields = [
        ner_data.get("job_title",""),
        ner_data.get("technical_skills",""),
        ner_data.get("relevant_skills",""),
        ner_data.get("domain_specific_terminologies",""),
        ner_data.get("exp_required_relevant","exp_relevant"),
        ner_data.get("exp_required_overall","exp_overall"),
        ner_data.get("projects",""),
        ner_data.get("certifications",""),
        ner_data.get("education",""),
        ner_data.get("job_responsibilites",""),
        ner_data.get("location","")
    ]

    # Convert list elements to comma-separated strings
    composite_fields = [', '.join(field) if isinstance(field, list) else field for field in composite_fields]

    #filter out empty strings and join them
    composite_text = " | ".join([field for field in composite_fields if field])
    return composite_text

def get_embedding_from_ner(ner_data: dict) -> np.ndarray:
    """
    Generate and return a normalized embedding vector from the given NER data
    """
    composite_text = create_composite_text(ner_data)
    embedding = embedder.encode([composite_text])[0]

    #Normalize the embedding for cosine similiarity

    norm = np.linalg.norm(embedding)
    if norm > 0:
        embedding=embedding/norm
    return np.array(embedding,dtype='float32')


index = faiss.IndexFlatIP(embedding_dim)

# for mapping index to actual NER data



def add_resume_to_faiss(ner_data: dict):
    """
    Create an embedding from resume NER data and add it to FAISS index
    """

    embedding = get_embedding_from_ner(ner_data)

    #FAISS expects a 2D array

    embedding = np.expand_dims(embedding,axis=0)
    index.add(embedding)

    session['resume_ner_data_list'].append(ner_data)

def calculate_fitment_score(cosine_similarity):
    """
    convert cosine similarity (0 to 1) to a fitment score (0 to 100)
    """

    return cosine_similarity*100

def improve_fitment_score(resume_ner,jd_ner,fitment_score):

    prompt = (
    "You are a candidate evaluation expert. Always use a zero-shot chain-of-thought approach with clear, step-by-step reasoning.Based on the information below,"
    "calculate a final fitment score for the candidate on a scale from 0 to 100. "
    "The score should primarily reflect the semantic similarity between the candidate's resume and the job description. "
    "Consider cosine similarity score as a baseline for the semantic match between the candidate's resume and the job description "
    "Analyze the semantic context of both the job description and the candidate's resume using the extracted NER data. If the candidate's qualifications, skills & achievements and all other things demonstrate a strong and highly relevant alignment with the job requirement, adjust the fitment score upward. Conversely, if there are notable gaps in the must have skills,relevant experience and qualification required by the job, the score should be lowered"
    "More focus should be given to technical_must_have_skills and relevant_skills from Job Description NER Entities."
    "The final fitment score should reflect a holistic semantic evaluation that accurately captures the candidate's true suitability for the role based on Job Description."
    "Do not set the final score to 0 unless there is virtually no match between the resume and the job description.\n\n"
    f"Job Description NER Entities : {jd_ner}\n"
    f"Candidate's resume NER Entities : {resume_ner}\n"
    f"Cosine similarity score (semantic match): {fitment_score:.2f}\n\n"
    "Strictly Provide only the final fitment score as a number between 0 and 100, without any additional commentary."
    # "Provide final fitment score with reasoning"
    )
    
    response_text = model.generate_content(prompt, generation_config=genai.types.GenerationConfig(
                    candidate_count=1,
                    temperature=0.3)
                                                    )

    print(response_text.text)

    return response_text.text

def compute_llm_fitment_score(jd_skills: str, resume_skills: str) -> float:
    """
    Compute a skill match score by comparing list of skills extracted(refined by LLM) from the JD & resume
    The score is the fraction of JD skills present in the resume
    Both inputs are comma separated
    """

    jd_skill_set = set([skill.strip().lower() for skill in jd_skills.split(",") if skill.strip()])
    resume_skill_set = set([skill.strip().lower() for skill in resume_skills.split(",") if skill.strip()])
    print(jd_skill_set)
    print(resume_skill_set)
    if not jd_skill_set:
        return 0.0

    match_count = len(jd_skill_set.intersection(resume_skill_set))
    print(jd_skill_set.intersection(resume_skill_set))
    return match_count/len(jd_skill_set)


def compute_experience_score(jd_required: str, candidate_exp: str) -> float:
    """
     Compute experience score by comparing candidate's relevant experience to the required experience from the JD
     Both are assumed to be string like "5 years" or "5 years 1 month"
     Return 1.0 if the candidate meets or exceeds the requirement
    """

    try:
        # simplistic parsing: take first number(years)
        required_years = float(jd_required.split()[0])
        candidate_years = float(candidate_exp.split()[0])
        print(required_years)
        print(candidate_years)

    except Exception as e:
        return 0.0
    return 1.0 if candidate_years >= required_years else candidate_years/required_years


def calculate_composite_score(similarity_score: float,
                              gemini_skill_score: float,
                              exp_score: float,
                              weights: dict = None  ) -> float:

    """
    Combine various signals to produce a holistic fitment score
    similarity_score: cosine similiarity from FAISS (0 to 1)
    gemini_skill_score: Fraction of JD skills found in resume (0 to 1)
    exp_score: Experience match score (0 to 1)

    Weights default to:
    - similarity: .5
    - gemini_skill: .3
    - experience: .2
    """

    if weights is None:
        weights = {"similiarity": 0.5, "gemini": 0.3, "experience": 0.2}

    final = (weights["similiarity"] * similarity_score +
             weights["gemini"] * gemini_skill_score +
             weights["experience"] * exp_score)

    return final*100

def generate_recommendation_summary(resume_ner,jd_ner,fitment_score):
  prompt= (
        "You are a candidate evaluation expert. Our platform processes resumes and job descriptions through a comprehensive pipeline that includes:"

        "NER Extraction: Structured entities(e.g., technical skills, relevant skills, domain terminologies, work experience, projects, certifications) are extracted form both the job description and candidate's resume."
        "Embedding Generation: A high quality transformer model converts the composite text(constructed from the extracted entities) into normalized embeddings."
        "Semantic Similarity Search: A cosine similarity search using FAISS is performed between the job description and resume embeddings."
        "Fitment score calculation: A composite fitment score is computed by combining the cosine similarity score with additional signals such as skill match, experience match, and industry/domain alignment"

        "Using the outputs from these stage, you are provided with the following inputs:"
        f"Job Description NER Entities: {jd_ner}"
        f"Resume NER Entities: {resume_ner}"
        f"Calculated Fitment Score: {fitment_score} (derived from cosine similarity search)"

        "Tasks:"
        "1. Provide a concise candidate summary (~40 words) that captures the candidate’s overall qualifications, skills, and work experience in direct alignment with the job description."
        "2. Identify and clearly list the top 3 strengths and key skills from the candidate that directly match the job requirements."
        "3. Analyze and summarize the candidate’s work experience in relation to the job description, emphasizing the most relevant roles and responsibilities."
        "4. Based on the fitment score and comparison between the job description and the resume, succinctly describe the candidate’s key areas of excellence and areas needing further development (~40 words)."
        "5. Highlight any significant skill gaps between the candidate’s resume and the job description."
        "6. Evaluate the candidate’s soft skills and cultural fit for the role based on the provided information."
        "7. Identify and summarize key projects, achievements, or contributions from the candidate that align with the job’s responsibilities."
        "8. Evaluate the candidate’s qualifications (education and certifications) in the context of the job requirements."
        "9. Based solely on the above analyses, output a final recommendation using one of these categories: [Highly Suitable, Suitable, Not Suitable]. Provide ONLY one of these options as the recommendation, without additional commentary."
        "10. Based on the overall analyses, suggest a profile for which candidate is best suited like data science/data engineeer etc. Please focus provide only one best suitable profile"
        
        "Output:"
        "Provide your entire response in JSON format with exactly the following keys (no extra keys or lists):"

        '''Format :
         {
          "candidate_summary": "",
          "Fitment Score Analysis": "",
          "top_3_strengths_&_skills": "",
          "work_experience_analysis": "",
          "areas_of_excellence_&_potential_development": "",
          "skill_gaps": "",
          "soft_skills_&_cultural_fit": "",
          "key_projects_&_achievements": "",
          "qualification_evaluation": "",
          "recommendation": "",
          "alternate_recommendation": ""
         }
        '''
      )
  response_text = model.generate_content(prompt, generation_config=genai.types.GenerationConfig(
                candidate_count=1,
                temperature=0.3)
                                                   )

  print(response_text.text)

  cleaned_paragraph = clean_json(response_text.text)
  #   print(cleaned_paragraph)
  json_data_recommendation = json.loads(cleaned_paragraph)


  return json_data_recommendation

def convert_list_to_text(value):
    if isinstance(value, list):
        return ', '.join(map(str, value))
    return value


def search_similar_jd(jd_ner_data: dict, top_k=4):
    jd_embedding = get_embedding_from_ner(jd_ner_data)
    jd_embedding = np.expand_dims(jd_embedding, axis=0)
    score, indices = index.search(jd_embedding, top_k)
    results = []

    for i, idx in enumerate(indices[0]):
        if idx < len(session['resume_ner_data_list']):
            fitment_score = calculate_fitment_score(score[0][i])
            resume = session['resume_ner_data_list'][idx]
            improved_fitment_score = improve_fitment_score(resume,jd_ner_data, score[0][i])
            gemini_skill_score = compute_llm_fitment_score(jd_ner_data.get("technical_must_have_skills", ""), resume.get("relevant_skills", ""))
            exp_score = compute_experience_score(jd_ner_data.get("exp_required_relevant", ""), resume.get("exp_relevant", ""))
            composite_score = calculate_composite_score((int(improved_fitment_score)/100), gemini_skill_score, exp_score)
            recommendation_summary = generate_recommendation_summary(resume,jd_ner_data, round(int(improved_fitment_score),2))
            results.append({
                "resume": resume,
                # "similarity_score": float(score[0][i]),
                "similarity_score": round(int(improved_fitment_score),2),
                "gemini_skill_score": round(gemini_skill_score*100,2),
                "exp_score": round(exp_score*100,2),
                "fitment_score": round(composite_score,2),
                "recommendation_summary": convert_list_to_text(recommendation_summary.get("candidate_summary")),
                "fitment_score_analysis": convert_list_to_text(recommendation_summary.get("Fitment Score Analysis")),
                "top_3_strengths_&_skills": convert_list_to_text(recommendation_summary.get("top_3_strengths_&_skills")),
                "work_experience_analysis": convert_list_to_text(recommendation_summary.get("work_experience_analysis")),
                "key_projects_&_achievements": convert_list_to_text(recommendation_summary.get("key_projects_&_achievements")),
                "qualification_evaluation": convert_list_to_text(recommendation_summary.get("qualification_evaluation")),
                "recommendation": recommendation_summary.get("recommendation",""),
                "alternate_recommendation": recommendation_summary.get("alternate_recommendation","")
            })
    return results

def clean_index():
#   global index, resume_ner_data_list
  index = faiss.IndexFlatIP(embedding_dim)
  session['resume_ner_data_list'] = [] #each element is a dict of the resume's NER output = []

# Convert list values to comma-separated values
def convert_list_to_string2(data):
    if isinstance(data, list):
        return ', '.join(data)
    return data

def jd_entity(jd_dictionary):
  
  for file_name, jd_text in jd_dictionary.items():
    doc3=nlp(jd_text)

    jd_entities = {"PERSON": [], "ORG": [], "Email": [], "Mobile No.": []}
    for ent in doc3.ents:
        if ent.label_ in jd_entities:
            jd_entities[ent.label_].append(ent.text)
    print(jd_entities)
    prompt = (
        "You are an expert HR specialized in recruitment. Always use a zero-shot chain-of-thought approach with clear, step-by-step reasoning by using conditions mentioned below and provide required information accurately."
        f"Below is a Job Description:\n{jd_text}\n\n"
        f"And here are some entities We have extracted: \n{jd_entities}\n\n"
        "Please provide the following entities that you found from Job Description: job_title, technical_must_have_skills, relevant_skills, industry_specific_terminologies, exp_required_relevant, exp_required_overall, projects, certifications, education, job_responsibilities, location,required_experience_and_qualification "
        "If you found multiple values for any details mentioned above, incorporate that only into comma-separated values. Do not use dictionaries/lists."
        '''Note: 1. exp_required_relevant and exp_required_overall strictly must be numerical value specifying years.
                 2. For industry_specific_terminologies include only the industry terminologies like insurance, life sciences, banking. Avoid including technical skills.
        '''
        "Provide the response in JSON format"
        '''Format:{
                  "job_title": "",
                  "technical_must_have_skills": "",
                  "relevant_skills": "",
                  "industry_specific_terminologies": "",
                  "exp_required_relevant": "",
                  "exp_required_overall": "",
                  "projects": "",
                  "certifications": "",
                  "education": "",
                  "job_responsibilites": "",
                  "location": "",
                  "required_experience_and_qualification":""
                  }
        '''
    )

    response_text = model.generate_content(prompt, generation_config=genai.types.GenerationConfig(
                    candidate_count=1,
                    temperature=0.3)
                                                      )

    print(response_text.text)

    cleaned_paragraph = clean_json(response_text.text)
    #   print(cleaned_paragraph)
    json_data_jd = json.loads(cleaned_paragraph)

    print(json_data_jd)

    session['jd_ner_list'].append(json_data_jd)

  return session['jd_ner_list']

def resume_jd_score(resume_dict,json_data_jd,text_input,GOOGLE_API_KEY):
  clean_index()
  for file_name, resume_text in resume_dict.items():
    doc4=nlp(resume_text)

    resume_entities = {"PERSON": [], "ORG": [], "Email": [], "Mobile No.": []}
    for ent in doc4.ents:
        if ent.label_ in resume_entities:
            resume_entities[ent.label_].append(ent.text)
    # print(jd_entities)

    normalize_terms = {
        "JS": "JavaScript",
        "ML": "Machine Learning",
        "Node.JS": "Node.js"
    }

    Domain = 'Data Engineer'

    prompt = (
            f"Below is a resume:\n{resume_text}\n\n"
            f"And here are some entities We have extracted: \n{resume_entities}\n\n"
            f"We have to scan the resume for {Domain} domain."
            f"Normalize the abbreviation or term to its full, standard form based on the resume context. Here are some examples: \n{normalize_terms}\n\n"
            "Please provide a refined comma-separated list of technical skills,comprehensive list of relevant skills, abilities, industry-specific skills present in the resume."
            "Be skill-agnostic and include emerging & domain-specific skills if applicable"
            "Don't inferred skills based on project descriptions and experience, only provide skills which are mentioned in resume."
            "Please provide the below details as well : Person, email,mobile no, exp(relevant & overall only in years), certifications"
            f"Use your understanding of {Domain} domain. Relevant skills(Technical skills only, do not include soft skills/project management skills) and relevant experience must be from the {Domain} domain."
            "If you found multiple values for any details mentioned above, incoporate that only into comma seperated values. Do not use dictiories/list."
            "Note: 1.exp_relevant and exp_overall strictly must be numerical value specifying years."
            "Provide the response in JSON format"
            '''Format:{
                      "technical_skills": "",
                      "relevant_skills": "",
                      "industry_specific_skills ": "",
                      "person": "",
                      "email": "",
                      "mobile_no": "",
                      "Projects": "",
                      "exp_relevant": "",
                      "exp_overall": "",
                      "certifications": ""
                  }'''
        )

    response_text = model.generate_content(prompt, generation_config=genai.types.GenerationConfig(
                    candidate_count=1,
                    temperature=0.3)
                                                      )

    print(response_text.text)

    cleaned_paragraph = clean_json(response_text.text)
    #   print(cleaned_paragraph)
    json_data_resume = json.loads(cleaned_paragraph)

    print(json_data_resume)

    add_resume_to_faiss(json_data_resume)

  similar_resumes = search_similar_jd(json_data_jd,top_k=len(resume_dict))

  for item in similar_resumes:
      for key in item['resume']:
          item['resume'][key] = convert_list_to_string2(item['resume'][key])


  # Convert JSON data to DataFrame
  df = pd.json_normalize(similar_resumes)
  df = df.drop(columns=['resume.Projects'])

  # Flatten the nested dictionary and rename columns
  df.columns = ['Score', 'Skill Score', 'Expeience Score','Final Score','Summary',"Fitment Score Analysis","top_3_strengths_&_skills","work_experience_analysis","key_projects_&_achievements","qualification_evaluation","recommendation",
                'Technical Skills', 'Relevant Skills',
                'Domain Specific Terminologies', 'Person',
                'Email', 'Mobile Numbers',
                'Relevant Experience', 'Overall Experience',
                 'Certifications'
                ]

  # Reorder columns to ensure "Fitment Score" is the last column
  columns_order = ['Technical Skills', 'Relevant Skills',
                  'Domain Specific Terminologies', 'Person',
                  'Email', 'Mobile Numbers',
                  'Relevant Experience', 'Overall Experience',
                  'Certifications',
                  'Score', 'Skill Score', 'Expeience Score','Final Score','Summary',"Fitment Score Analysis","top_3_strengths_&_skills","work_experience_analysis","key_projects_&_achievements","qualification_evaluation","recommendation"]
  df = df[columns_order]
  
  # Convert DataFrame to JSON
  json_data = df.to_json(orient='records')

  return json_data
  # # Save DataFrame to Excel file
  # df.to_excel('resumes.xlsx', index=False)

#   # Create a new Excel workbook and select the active worksheet
#   wb = Workbook()
#   ws = wb.active

#   # Append the DataFrame to the worksheet
#   for r in dataframe_to_rows(df, index=False, header=True):
#       ws.append(r)

#   # Set fixed color for column headers to blue
#   header_fill = PatternFill(start_color="8adbf5", end_color="8adbf5", fill_type="solid")
#   header_font = Font(bold=True)
#   for cell in ws[1]:
#       cell.fill = header_fill
#       cell.font = header_font

#   # Adjust column widths dynamically based on the content and set wrap text with max width of 30
#   max_width = 30
#   for col in ws.columns:
#       max_length = 0
#       column = col[0].column_letter  # Get the column name
#       for cell in col:
#           try:
#               if len(str(cell.value)) > max_length:
#                   max_length = len(cell.value)
#           except:
#               pass
#           cell.alignment = Alignment(wrap_text=True, vertical='top')
#       adjusted_width = min(max_length + 2, max_width)
#       ws.column_dimensions[column].width = adjusted_width

#   # Save the workbook to a file
#   wb.save("resumes_dynamic_width_wrap_text.xlsx")

def resume_jd_score_multi(resume_dict, json_data_jd,text_input,GOOGLE_API_KEY):
    clean_index()
    current_datetime = datetime.now()

    for file_name, resume_text in resume_dict.items():
        doc4=nlp(resume_text)

        resume_entities = {"PERSON": [], "ORG": [], "Email": [], "Mobile No.": []}
        for ent in doc4.ents:
            if ent.label_ in resume_entities:
                resume_entities[ent.label_].append(ent.text)
        normalize_terms = {
            "JS": "JavaScript",
            "ML": "Machine Learning",
            "Node.JS": "Node.js"
        }

        # Domain = 'Data Engineer'
        Domain = json_data_jd.get("job_title","")

        prompt = (
            "You are an expert HR specialized in recruitment. Always use a zero-shot chain-of-thought approach with clear, step-by-step reasoning by using conditions mentioned below and provide required information accurately."
            "Your first task is to arrange the resume content into a well-structured format. Please organize the raw text into the following sections: Professional Summary, Skills, Work Experience, Education, and any other relevant sections typically found in a resume. Ensure each section is clearly labeled and formatted appropriately."
            f"Below is a resume:\n{resume_text}\n\n"
            f"And here are some entities We have extracted: \n{resume_entities}\n\n"
            f"We have to scan the resume for {Domain} domain."
            f"Normalize the abbreviation or term to its full, standard form based on the resume context. Here are some examples: \n{normalize_terms}\n\n"
            f"Present/Today's date and time is {current_datetime}. Calculate years of experience (relevant & overall) upto present/today's date."
            f"For years of relevant experience, calculate how many years of experience candidate has for {Domain} domain/role from work experience. Please verify your answer as this is important step."
            "Please provide a refined comma-separated list of technical skills, comprehensive list of relevant skills, abilities, industry-specific skills present in the resume."
            "Be skill-agnostic and include emerging & domain-specific skills if applicable."
            "Don't infer skills based on project descriptions and experience, only provide skills which are mentioned in section projects/work experience/professional experience."
            "For industry-specific-skills, based on projects and overall resume induce industry-specific-skills like life science, insurance, etc."
            "Please provide the below details as well: Person(name of the candidate), email, mobile no, exp (relevant & overall only in years), certifications, education and qualifications, job history."
            "Job history should include following keys alongwith their values : job_title, duration, job_responsibilities"
            f"Use your understanding of {Domain} domain. Relevant skills (Technical skills only, do not include soft skills/project management skills) and relevant experience must be from the {Domain} domain only."
            "If you found multiple values for any details mentioned above, incorporate that only into comma-separated values. Do not use dictionaries/lists."
            f'''Note: 1. exp_relevant and exp_overall strictly must be numerical value specifying years.
                     2. Your task is to search for the technical skills section or similar sections where candidate usually write their skills as comma or \n separated.
                        Along with this search for the professional summary or similar to summary where candidate gives small introduction about themselves. sometimes it is just a paragraph.
                        You have to exclude that whole section containing skills as comma or \n separated, paragraph of summary or introduction of candidate.
                     3. Next task here is to go through each project done by the candidate, extract relevant skills which are mentioned by the candidate in the project.
                        If you are not able to extract any skill strictly provide none.
                     4. Do not include skills or experience of candidate as an intern. Also use your understanding and do not include skills from profile summary, paragraph explaining summary.
                     5. Technical skills should include all the skills candidate used across all projects and they should be mentioned under Projects/Work Experience/Professional Experience, Relevant skills will be the subset of Technical skills which will
                        include only that skills which are relevant to {Domain} domain/role.
                     6. Please remember that do not include skills from "Skills" / "Professional Summary" sections as technical or relevant skills.
            '''
            "Provide the response in JSON format."
            '''Format:{
                      "technical_skills": "",
                      "relevant_skills": "",
                      "industry_specific_skills": "",
                      "person": "",
                      "email": "",
                      "mobile_no": "",
                      "Projects": "",
                      "exp_relevant": "",
                      "exp_overall": "",
                      "certifications": "",
                      "education_and_qualifications": "",
                      "job_history":""
                  }'''
        )

        response_text = model.generate_content(prompt, generation_config=genai.types.GenerationConfig(
            candidate_count=1,
            temperature=0.1)
        )

        print(response_text.text)

        cleaned_paragraph = clean_json(response_text.text)
        json_data_resume = json.loads(cleaned_paragraph)

        json_data_resume['File Name'] = file_name

        print(json_data_resume)

        add_resume_to_faiss(json_data_resume)

    similar_resumes = search_similar_jd(json_data_jd, top_k=len(resume_dict))
    # # Process multiple JDs
    # all_results = process_multiple_jds(json_data_jd, resume_dict)
    print(similar_resumes)
    # Convert list values to comma-separated values
    # def convert_list_to_string(data):
    #     if isinstance(data, list):
    #         return ', '.join(data)
    #     return data

    def convert_list_to_string(data):
      if isinstance(data, list):
          # Handle nested dictionaries if present
          return ', '.join([', '.join(v.values()) if isinstance(v, dict) else str(v) for v in data])
      return data

    for item in similar_resumes:
        for key in item['resume']:
            item['resume'][key] = convert_list_to_string(item['resume'][key])
    print(similar_resumes)
    # Convert JSON data to DataFrame
    df = pd.json_normalize(similar_resumes)
    df = df.drop(columns=['resume.Projects','resume.education_and_qualifications','resume.job_history'])

    print(df)

    # Flatten the nested dictionary and rename columns
    df.columns = ['Fitment Score (%)', 'JD vs Candidate - Skill Match(%)', 'JD vs Candidate - Experience Match(%)','Composite Fitement Score (%)','Summary',"Fitment Score Analysis","Key Skills Matched","Work Experience Analysis","Key Projects","Qualification & Certifications","Recommendation","Alternate Recommendation",
                  'Technical Skills', 'Relevant Skills',
                  'Domain Specific Skills', 'Candidate Name',
                  'Email', 'Phone',
                  'Work Experience - Relevant', 'Work Experience - Overall',
                  'Certifications','Resume File Name'
                  ]

    df['Candidate Name'] = df['Candidate Name'].str.title()

    # Reorder columns to ensure "Fitment Score" is the last column
    columns_order = ['Resume File Name','Candidate Name',
                    'Email', 'Phone','Technical Skills', 'Relevant Skills',
                    'Domain Specific Skills',
                    'Work Experience - Relevant', 'Work Experience - Overall',"Work Experience Analysis","Key Projects","Qualification & Certifications",
                    'Certifications',
                    'Fitment Score (%)', 'JD vs Candidate - Skill Match(%)', 'JD vs Candidate - Experience Match(%)','Composite Fitement Score (%)',"Fitment Score Analysis","Key Skills Matched",'Summary',"Recommendation","Alternate Recommendation"]
    df = df[columns_order]
    
    df['Alternate Recommendation'] = df.apply(
    lambda row: row['Alternate Recommendation'] if row['Recommendation'] == 'Not Suitable' else '',
    axis=1
)

    return df

########################################
    
def fetch_append_update_json_in_s3(bucket, file_path, new_key, new_value):

    s3 = boto3.client('s3')

    try:
        obj = s3_client.get_object(Bucket=bucket, Key=file_path)
        data = obj['Body'].read().decode('utf-8')
        # Fetch the existing JSON file from S3
        json_data = json.loads(data)

        # Append the new key-value pair to the JSON data
        new_element={new_key : new_value}
        
        json_data.append(new_element)
        # Update the JSON file in S3
        s3_client.put_object(Bucket=bucket, Key=file_path, Body=json.dumps(json_data).encode('utf-8'))

        print(f"Successfully appended '{new_key}: {new_value}' to {Key} and updated in S3.")

    except Exception as e:
        print(f"Error occurred: {str(e)}")
    
def delete_folder_from_s3(session_id):
    bucket = s3_resource.Bucket('ires-v1')
    for obj in bucket.objects.filter(Prefix=session_id):
        s3_client.delete_object(Bucket='ires-v1', Key=obj.key)

DS_Skills = ["Probability and Statistics", "Hypothesis Testing", "Regression Analysis", "Time Series Analysis",
             "Data Cleaning and Preprocessing",
             "Exploratory Data Analysis (EDA)", "Data Visualization (Matplotlib, Seaborn, ggplot2)",
             "Feature Engineering",
             "Supervised Learning (Classification, Regression)",
             "Unsupervised Learning (Clustering, Dimensionality Reduction)",
             "Ensemble Methods (Random Forests, Gradient Boosting)", "Text Processing", "Tokenization",
             "Named Entity Recognition", "Natural Language Processing", "Sentiment Analysis"]

DL_Skills = ["Neural Networks", "Convolutional Neural Networks (CNN)", "Recurrent Neural Networks (RNN)",
             "Long Short-Term Memory (LSTM)",
             "Gated Recurrent Units (GRU)", "Autoencoders", "Reinforcement Learning", "Transfer Learning",
             "Hyperparameter Tuning", "TensorFlow",
             "Keras", "PyTorch", "Theano", "Natural Language Processing (NLP) with Deep Learning Techniques",
             "Computer Vision with Deep Learning Techniques"]

MLOps_Skills = ["Continuous Integration/Continuous Deployment (CI/CD)", "Model Deployment", "Model Monitoring",
                "Model Versioning",
                "Model Orchestration", "Containerization (Docker, Kubernetes)", "Automated Testing",
                "Scalability", "Performance Optimization", "Model Governance", "Data Versioning", "MLOps Practices",
                "Cloud Computing Platforms (AWS, Azure, Google Cloud)",
                "Deployment Automation Tools (TensorFlow Serving, Seldon Core, Kubeflow)",
                "MLOps Platforms (MLflow, TFX, DataRobot)", "Workflow Automation Tools (Airflow, Luigi)"]

GenAI_Skills = ["Generative Adversarial Networks (GAN)", "Variational Autoencoders (VAE)", "Transformers", "DALL-E",
                "Langchain", "Hugging Face",
                "GPT", "Gemini", "OpenAI GPT 3.5", "Palm", " LLMs", "Llama2", "Mistral 8*7", "Prompt Engineering",
                "Transformers", "BERT", "Fine tune LLM using LoRA, QLoRA"]

Example1 = """
         Resume content = "POOJA UGALE JR.DATA SCIENTIST A tech lover Data Scientist with strong math background and 3+ years of experience using predictive modelling, data processing, and data mining algorithms to solve challenging business problems. 
                         poojaugale777@gmail.com +91 7057444357 Pune, Maharashtra WORK EXPERIENCE JR. DATA SCIENTIST IFS INDIA MARCANTILE Pvt. Ltd. NOV 2019 -Present Pune, Maharashtra Project Name: Built predictive based piecewise degradation model of remaining useful life prediction of an Turbofan engine 
                         Domain:Aviation Skills/Tools Used: LSTM, Train test split, CNN,GRU, Matplotlib, RNN, Encoder, Decoder, Attention Based Encoder Decoder GRU Description: • Built a machine learning and deep learning model that can predicts the failure of Machine in upcoming number of days. 
                         • The objective of this project is to implement various predictive maintenance methods and assess the performance of each equipment • Predictive maintenance techniques are used to determine the condition of an equipment to plan the maintenance /Failure ahead of its time. 
                         By doing this the equipment downtime cost can be reduced significantly. Project Name:On Field Survice Resource Monitoring System Domain: Construction Field Skills/Tools Used: Logestic egression ,SVM,KNN,MLP Classifier 
                         Description: • To Built a machine learning model that can used to classifying the performance is in working state,idle or other . • Installing IoT devices to generate the data and creat resouces management mining in construction business onsite and offsite and creat an automated dashboard to view the report. 
                         • Create web application to fetch the data from DB Project Name: Legal Document Curation using NLP Domain: Legal System Skills/Tools Used: • Tokenization,stemming,Lemmatization,train test split, Word cloud,TF-IDF,Gensim,Sumy, Transformer. 
                         Description: • An Important challenge in legal text summarization is identifying the information part while avoiding the irrelevant one. • The goal of thise project is to develop a web application for meaningful legal text summarization an extraction of entities 
                         • Summarization is performed by using Spacy,Gensim,Nltk etc KEY SKILLS Regression Analysis Natural Language Processing SQL Machine Learning Data Pre-processing Visualization Artificial Intelligence (AI) Deep Learning Data Cleaning Social Media Mining Feature Engineering Computer Vision Microsoft Excel Advanced Statistics Python Hypothesis Testing AWS 
                         TECHNICAL EXPERTISE • Programming languages: Python • Packages: Numpy, SciPy, Pandas, NLTK, Matplotlib, Scikit- learn, Tensorflow, Keras • Tools : Jupyter Notebook, Anaconda, MS Word and MS Excel. • Modeling : Regression, Classification like Random forest, Logistic, SVM,Bagging, Clustering -K-Means. • Deep Learning: Natural Language Processing, ANN, CNN, RNN, LSTM, etc. 
                         EDUCATIONAL DETAILS • Bachelor Degree Civil Engineering with 74% (Distinction)from Savitribai Phule Pune University. • HSC with 75.69% in 2014 from State Board, Maharashtra. • SSC with 82.73% in 2012 from State Board, Maharashtra. PERSONAL DETAILS • Date of Birth : 07/07/1996 • Current Address : Panchvati Society, R.no.401, B-wing, Sinhgad campus, Ambegaon,Pune-46 
                         • Languages : English, Hindi, Marathi SOFT SKILLS Complex Problem Solver Critical Thinking Team Player Good Product Understanding Quick Learner INTERESTS Learning New Skills Investing Spirituality Sports "

        Here are the details of the candidates:

        1. Programming Languages: Python
        2. Recent Experience: 4 years
        3. Machine learning skills: Logestic regression ,SVM, Tokenization,stemming,Lemmatization,train test split, Word cloud,TF-IDF,Gensim,Sumy,Transformer
        4. Deep Learning skills: LSTM, Train test split, CNN,GRU, Matplotlib, RNN, Encoder, Decoder, Attention Based Encoder Decoder GRU
        5. MLOps skills: None
        6. Cloud skills: AWS
        7. GenAI skills: None
        8. List of projects: 
        1) Built predictive based piecewise degradation model of remaining useful life prediction of an Turbofan engine
        2) On Field Survice Resource Monitoring System 
        3) Legal Document Curation using NLP
        9. Recommendation: Recommend for interview. The candidate has a knowledge of Python and has 4 years of recent experience. 
                           Candidate has Machine learning skills and Deep learning skills and done at least 3 projects, in addition they have cloud skills.
        10. Name of the candidate: POOJA UGALE
        11. count of projects: 3
        12. additional skills: None
        13. email: poojaugale777@gmail.com
        14. phone: +91 7057444357

        """

Example2 = """
        Resume content = "Jayshree Sharma Indian Institute of Information Technology Guwahati (+91) 8254812553/(+91)7099215869 bpsdworld@gmail.com 
                         OBJECTIVE WORK EXPERIENCES  Working as Data Scientist: Advanced Analytics in IBM, Bengaluru from 16/02/2021 to present. 
                         Project Digital Solutions Asset Performance Management Description In this project, we need to bring out an Artificial Intelligence based solution for asset management for the power grid company. 
                         Employer IBM India Private Limited Responsibilities  Working as Data Scientist, implementing business rules through IBM ODM (Operational Decision Management). \
                          Creating models and deployment with docker. Environment Python, Machine learning Models, ODM.  Worked as full time Data Science-intern at TNG Innovation Labs Pvt Ltd, Bangalore from June 2020 to 30 November 2020. 
                          Worked for fetching and analyzing data for live AI projects like fleet management, routing etc  Worked for Machine learning, AI projects for videos: Getting camera pose from dash cam video and finding path of vehicle including left/right turns and detecting accident/non accident. 
                         Used Random forest, CNN and other models.  Worked for Machine learning, AI projects for time series data: Used LSTM, Bi-LSTM, CNN-LSTM models. Project Accident Détection(Jido) Description As part of this project, we detect the accident or non-accidenttrips and type of accident (front/back/side(left/right) hit) using the sensor data collected from vehicle and using dashcam video. 
                         Employer TNG Innovation Labs Pvt. Ltd. Responsibilities  Getting camera pose from dash cam video using CNN and other models.  Finding path of vehicle including left/right turns and detecting accident/non accident using different approaches and finding best suitable one. 
                          Getting high performance using supervised and unsupervised approaches along with statistical methods. Environment Python, Machine learning Models Project Vehicle route tracking Description The objective of this project is to track the vehicle live route using sensor data collected from vehicle. 
                         Employer TNG Innovation Labs Pvt. Ltd. Responsibilities  Using reverse geo coding and finding the total length travelled by a vehicle within a state. 
                          Writing a script for finding at what time vehicle crosses the border and how many litres of fuel consumed by vehicle within a state Environment Python, Machine learning Models Project Employee Transportation system Description As part of this project, we have find best route for employee transportation with minimum time and least fuel consumption. Employer TNG Innovation Labs Pvt. Ltd. Responsibilities 
                          Analysing the Speed limit data in third party APIs like Mapbox navigation app, OSM, Tomtom Navigation App, Bing navigation app and finding the best app suitable as per use case.  Representing and comparing the data from all the possible app available in the market for navigation with speed limit.  Presenting the above details to seniors.  Doing research work on getting algorithms for bus stop selection with given constrains and delivering same to teammates. 
                         Environment Python, Machine learning Models  Worked as Computer Science Faculty in Jawahar Navodaya Vidyalaya, Nongstoin, Meghalaya from Aug 2016 to April 2017 
                          Worked as a Database Administrator at ISS (Integrated System Services) from May 2014 to May 2016. Worked on projects like SOCIO-ECONOMIC CASTE CENSUS (SECC). My responsibility was to manage and supervise a team of 50 resources across 12 districts of Assam under the SECC project working on solving technical issues like database issues, software problems and other non-technicalissues. 
                         Project Socio-Economic Caste census 2011 Description The objective of this project is to digitize the census data and correct it as per requirement. Employer Integrated Systems & Services. Responsibilities  My responsibility was to manage and supervise a team of 50 resources across 12 districts of Assam under the SECC project working on solving technical issues like database issues, software problems and other non-technicalissues. Environment Mysql. PAPER PUBLISHED 
                          An IoT based Smart Assistive Devicefor the Visually Impaired. 2020 IEEERegion 10 Symposium (TENSYMP) TECHNICAL SKILLS  PRIMARY SKILLS: MACHINE LEARNING, DEEP LEARNING, PYTHON, COMPUTER VISION, AI  SECONDARY SKILLS: DATA SCIENCE, BIG DATA, SPARK, AWS  Programming Skills: C, Python, C++, Scala (basics)  Other languages: MySQL,HTML, CSS  Operating System: Windows, Ubuntu  Software Tools/Frameworks: Jupiter Notebook, Anaconda, Pycharm, Eclipse, MatLab, Netbeans, Spark, Keras, Tensorflow, AWS. 
                          Data Science, Machine learning, AI using Numpy, Pandas, Scikit learn, Regression techniques, Clustering algorithms, SVM, Random forest, CNN, LSTM, geopandas, geocoding etc , . ACADEMIC DETAILS  M.Tech in Computer Science and Engineering [2020]: Indian Institute of Information Technology Guwahati (IIIT). CGPA:7.6  Bachelor of Technology (B. Tech IT) [ 2012]: Department of Information Technology, Assam University Silchar (Central University) Percentage: 66.25% 
                          Higher Secondary Examination [2006]: Jawahar Navodaya Vidyalaya (CBSE), Tinsukia, Assam Percentage: 72.00%  High School Leaving Certificate Examination [2004]: Jawahar Navodaya Vidyalaya (CBSE), Tinsukia Percentage: 79.40% PERSONAL SKILLS Comprehensive problem-solving abilities, excellent verbal and written communication skills, ability to interact with people efficiently, willingness to learn. Brilliant team player, self-driven, well at collaborating with the team all the time. CERTIFICATION 
                          Certification from NPTEL on Cloud Computing under Prof. Adrijit Goswami, IITKharagpur. EXTRA CURRICULAR ACTIVITIES AND ACHIEVEM ENTS  Qualified GATE 2011,2018,2019.  NCC “B” Certificate holder as “Senior Under Officer”  Secured “A” grade in All India Level Intelligence, Aptitude and Scholastic test.  Secured “A” grade in Indian Institute of personality development test at all India level.  Qualified 43rdAnnual All India UN Information Test.  Got “NEHRU AWARD” by Tea Board India.
                           Got 3rd prize in 8th National „Youth Parliament‟ Competition at Regional level.  Secured position in various competitions and eventslike racing, essay writing, song, badminton etc. ACADEMIC PROJECT  Human Activity Recognition using feature learning and Transfer learning (June 2019 –June 2020) M.Tech- Project- Under the guidance of Dr. Dip Sankar Banerjee, IIITG The project primarily aims at finding a good approach for recognizing human activity on sensor dataset. 
                          Focus on dynamic feature learning and Transfer Learning for activity recognition. Technology Used: Python3, keras with Tensorflow backend.  Anomaly Detection using MLlib (Jan -Apr 2019) Instructor Dr. Dip Sankar Banerjee, IIIT Guwahati The Credit Card Fraud Detection by applying different Machine Learning Techniqueson past credit card transaction dataset and analyzing the result. The aim here is to detect the fraudulent transactions by using different machine learning techniques. 
                          Technology Used: pySpark  IoT (Internet of Thing) project “Smart blind stick” Instructor Dr. Rakesh Matam, IIIT Guwahati (Jan-April 2019) The aim of the project is to build an intelligent walking stick that provide an efficient movement aid for the blind persons which gives a sense of vision by providing alert when any object comes in front of them and also notify family or friend when help is needed. Technology Used: Raspberry pi, Raspbian OS. 
                           Network on Chip ( 2012) Under the guidanceof Asst. Prof. Abhijit Biswas, TS School of Technology, Assam University Silchar. The project involved the study of delay optimization in Network on Chip.  Medical shop Management System ( 2011) Under the guidanceof Prof. Ajoy Kumar Khan, TS Schoolof Technology, Assam University Silchar. The project involved management of all records related to a medical shop. Technology Used: Java, MySQL INTERNSHIP 
                           Undertaken Industrial Training a t IOCL, Guwahati Refinery, Title: Networking REPOSITORIES AND NETWORK  LinkedIn: https://www.linkedin.com/in/jayshree -sharma-022bb672/  Github: https://github.com/Sharmajayshree REFERENCES References are available on request."

        Here are the details of the candidates:

        1. Programming Languages: Python
        2. Recent Experience: 3 years
        3. Machine learning skills: Python, Machine learning Models
        4. Deep Learning skills: None
        5. MLOps skills: None
        6. Cloud skills: AWS
        7. GenAI skills: None
        8. List of projects: 
        1) Digital Solutions Asset Performance Management                  
        9. Recommendation: Not Recommended for interview. Candidate have very less Machine learning skills and don't have deep learning, MLOps, cloud and genai skills.
                           Candidate don't have any good Machine learning projects.
        10. Name of the candidate: Jayshree Sharma 
        11. Count of projects: 1      
        12. additional skills: None
        13. email: bpsdworld@gmail.com
        14. phone: (+91) 8254812553/(+91)7099215869


        """


def generate_resume_summary(resumes_dict, Optional_skills,DS_Skills,DL_Skills,MLOps_Skills,GenAI_Skills,selectfield,GOOGLE_API_KEY):
    data = []
    questions_dict = {}
    index = 0
    total_resumes = len(resumes_dict)
    genai.configure(api_key=GOOGLE_API_KEY)
    model = genai.GenerativeModel('gemini-1.5-flash')
    for file_name, ( text_content) in resumes_dict.items():

        prompt5 = f"""
                Given is the resume content of the candidate. Analyze the resume.

                STEP 1:
                    Your task is to search for the technical skills section or similar sections where candidate usually write their skills as comma separated.
                    Along with this search for the professional summary or similar to summary where candidate gives small introduction about themselves. sometimes it is just a paragraph.
                    Here, use your understanding that you need to avoid excluding project details or information about projects.
                STEP 2:
                    You have to exclude that whole section containing skills as comma separated, paragraph of summary or introduction of candidate and give everything other than this in proper format.
                    Include name, Email ID and phone number.


                Resume content:{text_content} 
                """

        try:

            try:
                response_text5 = model.generate_content(prompt5, generation_config=genai.types.GenerationConfig(
                    candidate_count=1,
                    temperature=0.3)
                                                        )
                                                                                                                                                                                                                                                               
            except Exception as e:
                if "Resource has been exhausted" in str(e):
                    print("Rate limit exceeded. Please try again later.")
                    time.sleep(60)
                    response_text5 = model.generate_content(prompt5, generation_config=genai.types.GenerationConfig(
                    candidate_count=1,
                    temperature=0.3)
                                                        )
                else:
                    raise e

            #print(response_text5.text)
            # model_info = genai.get_model("models/gemini-1.5-flash")
            # print(f"{model_info.input_token_limit=}")
            # print(f"{model_info.output_token_limit=}")
            # print(response_text5.usage_metadata)

            prompt4 = f"""
                Your task is to analyze the given resume. You have to think step by step as mentioned and give answer of the questions in just 1 to 2 words asked on the resume.

                I am providing you example of both recommended candidate as well as not recommended candidate. Please go through it and understand how skills and other things are extracted.
                I want you to follow the same logic while giving answers for the following questions. 

                Example of recommended candidate = {Example1}

                Example of not recommended candidate = {Example2}. Here Candidate has experience as an intern so we excluded everything.

                In the above both examples, user search for powerbi as optional skill but both candidate haven't worked on it that's why additional skills are marked as 'No'.

                Ensure accuracy while providing the answers. Avoid giving skills mentioned in the questions, used it only for your reference.

                STEP 1:
                    Question: Does candidate mentioned or used 'Python' or 'R' programming language? Provide programming langauge used only.

                    If answer is 'Yes' then only move to next part else simply return 'No programming language'.

                STEP 2:
                    Question:  Does the candidate have at least 2 years of recent experience in NLP, Machine learning, Machine learning, Deep learning or GenAI ?

                    If yes then only move to next question else provide 'Less experience'.

                    Question: If 'Yes', does the candidate have used any of the NLP, Machine learning, Machine learning, Deep learning or GenAI skills during this recent experience? 

                    Question: How many years of recent experience candidate have? Recent experience means last company or present company in which candidate is/was working. Consider present as Feb 2024.

                STEP 3:
                    For this step only consider resume content which talks about candidates actual work experience in Machine learning, deep learning, genai, mlops. Please focus only to answer skills related questions below based on such work experience/projects only. Don't extract skills which are mentioned explicitly under technical skills or primary skills or secondary skills and which are not part of sections related to work experience and project experience.
                    For your understanding following are the examples of skills mapping.
                    You should analyze the content of the resume and extract skills based on patterns and keywords commonly associated with these domains.
                    Avoid giving skills mentioned below, used it only for your reference. 

                    Machine learning skills : {DS_Skills}
                    Deep Learning skills : {DL_Skills}
                    MLOps skills : {MLOps_Skills} 
                    GenAI skills : {GenAI_Skills}

                    Your task here is to go through each project done by the candidate, extract relevant skills which are mentioned by the candidate in the project. 
                    If you are not able to extract any skill strictly provide none.

                    For following questions, Only extract projects/work experience/professional experience  related to machine learning, deep learning, mlops, genai ONLY if he/she 'developed' or 'build' something in the project. 

                    Question: What are the skills mentioned in sections like technical skills or primary skills or secondary skills? 
                    Please do not include this skills in the analysis.

                    Question: Does the candidate has any work experience as an intern? What skills he used as an intern ?
                    Please do not include this skills in the analysis.

                    Question: What are the specific Machine learning skills candidate has utilized or developed while working on projects except as an intern?
                    question: Are you sure ? Ensure accuracy while giving response and include only technical skills.
                    Question: What are the specific deep learning skills candidate has utilized or developed while working on projects except as an intern?
                    question: Are you sure ? Ensure accuracy while giving response and include only technical skills.
                    Question: What are the specific MLOps skills candidate has utilized or developed while working on projects?
                    question: Are you sure ? Ensure accuracy while giving response and include only technical skills.
                    Question: What are the specific cloud skills candidate has utilized or developed while working on projects?
                    question: Are you sure ? Ensure accuracy while giving response and include only technical skills.
                    Question: What are the specific GenAI skills candidate has utilized or developed while working on projects?
                    question: Are you sure ? Ensure accuracy while giving response and include only technical skills.
                    
                    If {Optional_skills} is empty then provide additional skills as 'None'.
                    check if candidate has mentioned or used skills {Optional_skills} in the resume, if yes then provide that skill as additional skills else return 'None'.
                    Avoid extracting skills other than {Optional_skills} in additional skills.
                    
                    Always verify your answer and check above mentioned conditions are met, then provide updated answers.

                STEP 4:
                    Question: Provide the machine learning/deep learning/genai related project names done by the candidate?
                    Question: Provide the total number of  machine learning/deep learning/genai related project done by the candidate as mentioned above?

                STEP 5:
                    You are SME in Machine learning and genai and your task is to evaluate given resume based on step 1 to step 5. 
                    If the candidate has the skills in Machine learning and deep learning in step 3 and at least 3 projects as calculated in step 4, recommend that candidate for job with detailed reasoning referring to the responses from step 1 to step 3. If the candidate is rejected give the proper reasoning.
                    Use your understanding while giving recommendation.

                STEP 6:
                    Provide the name of the candidate. If you are unable to find name then provide None.
                    Provide the Email ID of the candidate. If you are unable to find Email ID then provide None.
                    Provide the Phone number of the candidate. If you are unable to find Phone number then provide None.


                Please follow the below format (Use Plain text only, avoid using Bold format) for providing the response, Use the latest skills fetched for related asked question:


                1. Programming Languages:
                2. Recent Experience:
                3. Machine learning skills:
                4. Deep Learning skills:
                5. MLOps skills:
                6. Cloud skills:
                7. GenAI skills:
                8. List of projects:
                9. Recommendation:
                10. Name of the candidate:
                11. count of projects:
                12. additional skills:
                13. email:
                14. phone:


                Resume content:{response_text5.text}
                """

            try:
                response_text = model.generate_content(prompt4, generation_config=genai.types.GenerationConfig(
                candidate_count=1,
                temperature=0.3)
                                                   )
            except Exception as e:
                if "Resource has been exhausted" in str(e):
                    print("Rate limit exceeded. Please try again later.")
                    time.sleep(60)
                    response_text = model.generate_content(prompt4, generation_config=genai.types.GenerationConfig(
                        candidate_count=1,
                        temperature=0.3)
                                                    )
                else:
                    raise e

            # st.write(response_text.text)

            language_pattern = r"\*{0,2}Programming Languages:\*{0,2}\s(.*?)(?=\d+\. |\*{0,2}Recent Experience:\*{0,2}|\n)"

            duration_pattern = r"\*{0,2}Recent Experience:\*{0,2}\s(.*?)(?=\*{0,2}Machine learning skills:\*{0,2}|\n)"

            ds_pattern = r"\*{0,2}Machine learning skills:\*{0,2}\s(.*?)(?=\*{0,2}Deep Learning skills:\*{0,2}|\n)"

            dl_pattern = r"\*{0,2}Deep Learning skills:\*{0,2}\s(.*?)(?=\*{0,2}MLOps skills:\*{0,2}|\n)"

            MLOps_pattern = r"\*{0,2}MLOps skills:\*{0,2}\s(.*?)(?=\*{0,2}Cloud skills:\*{0,2}|\n)"

            Cloud_pattern = r"\*{0,2}Cloud skills:\*{0,2}\s(.*?)(?=\*{0,2}GenAI skills:\*{0,2}|\n)"

            genai_pattern = r"\*{0,2}GenAI skills:\*{0,2}\s(.*?)(?=\*{0,2}List of projects:\*{0,2}|\n)"

            projects_pattern = r"\*{0,2}List of projects:\*{0,2}\s(.*?)(?=\*{0,2}\d+\. Recommendation:\*{0,2}|\Z)"

            recommendation_pattern = r"\*{0,2}Recommendation:\*{0,2}\s(.*?)(?=\*{0,2}Name of the candidate:\*{0,2}|\n)"

            name_pattern = r"\*{0,2}Name of the candidate:\*{0,2}\s(.*?)(?=\*{0,2}count of projects:\*{0,2}|\n)"

            count_pattern = r"\*{0,2}count of projects:\*{0,2}\s(.*?)(?=\*{0,2}additional skills:\*{0,2}|\n)"

            additional_pattern = r"\*{0,2}additional skills:\*{0,2}\s(.*?)(?=\*{0,2}email:\*{0,2}|\n)"

            email_pattern = r"\*{0,2}email:\*{0,2}\s(.*?)(?=\*{0,2}phone:\*{0,2}|\n)"
            
            phone_pattern = r"\*{0,2}phone:\*{0,2}\s(.*?)$"


            # Extract text following each label using regular expressions

            extract_text = lambda pattern, response: re.search(pattern, response, re.DOTALL).group(

                1).strip() if re.search(pattern, response, re.DOTALL) else "Could not read"

            language = extract_text(language_pattern, response_text.text)

            duration = extract_text(duration_pattern, response_text.text)

            ds_skills = extract_text(ds_pattern, response_text.text)

            dl_skills = extract_text(dl_pattern, response_text.text)

            mlops_skills = extract_text(MLOps_pattern, response_text.text)

            cloud_skills = extract_text(Cloud_pattern, response_text.text)

            genai_skills = extract_text(genai_pattern, response_text.text)

            projects = extract_text(projects_pattern, response_text.text)
            # st.write(projects)

            recommendation = extract_text(recommendation_pattern, response_text.text)

            additional = extract_text(additional_pattern, response_text.text)

            name = extract_text(name_pattern, response_text.text)
            
            email = extract_text(email_pattern, response_text.text)

            phone = extract_text(phone_pattern, response_text.text)
            
            # print(phone)
            phone = phone.strip()
            phone = phone.replace('+91',"")
            phone = phone.replace('-',"")
            phone = phone.replace('+',"").replace('(',"").replace(")","")
            phone = phone.strip()
            # print(phone)

            if str(name).lower() in ['none', 'not provided', 'not mentioned', 'could not read', 'not available',
                                     'not mentioned in the resume']:
                name = ' '.join(
                    re.sub(r'\[.*?\]', '', os.path.splitext(file_name)[0]).replace('+', ' ').replace('_', ' ').replace(
                        'Naukri', '').replace('PHENOM', '').split()[:2]).title()
            else:
                name = name.title()

            count = extract_text(count_pattern, response_text.text)

            if 'None' in str(count) or 'Not' in str(count) or 'No' in str(count) or 'not' in str(count) or str(
                    count) == "":
                count = 0
            else:
                count = count

            ds_list = ds_skills.split(',')
            ds_list = [item.strip() for item in ds_list]
            dl_list = dl_skills.split(',')
            dl_list = [item.strip() for item in dl_list]
            genai_list = genai_skills.split(',')
            genai_list = [item.strip() for item in genai_list]

            llm_list = ['large language models (LLMs)', 'LLM', 'GPT', 'GPT 3.5', 'GPT-3', 'BERT', 'ChatGPT',
                        'GPT-4', 'Generative AI', 'GenAI', 'Prompts', 'AzureOpenAI',
                        'langchain', 'Llamaindex framework', 'LLMs', 'Vector DB', 'Open-AI Da-Vinci LLM',
                        'Development of LLMs and Generative AI applications', "LLM's",
                        'PaLM', 'Gen AI', 'Prompt Engineering', 'Fine tune LLM using LoRA', 'QLoRA']

            for item in llm_list:
                if item in ds_list:
                    ds_list.remove(item)
                    if item not in genai_list:
                        genai_list.append(item)
                if item in dl_list:
                    dl_list.remove(item)
                    if item not in genai_list:
                        genai_list.append(item)

            ds_skills = ','.join(ds_list)
            dl_skills = ','.join(dl_list)
            genai_skills = ','.join(genai_list)
            genai_skills = genai_skills.replace('None,', '')

            def calculate_skill_score(skills_str):
                skill_count = 0
                skill_count = skills_str.count(',') + 1

                score = 0
                if skill_count == 1 and (
                        'None' in skills_str or 'Not' in skills_str or 'No' in skills_str or 'not' in skills_str or skills_str == ""):
                    skill_count = 0
                    score = 0


                elif skill_count >= 1 and skill_count <= 8:
                    score = 0.25 * skill_count
                else:
                    score = 2

                return [skill_count, score]

            def calculate_ds_skill_score(skills_str):
                skill_count = 0
                skill_count = skills_str.count(',') + 1

                score = 0
                if skill_count == 1 and (
                        'None' in skills_str or 'Not' in skills_str or 'No' in skills_str or 'not' in skills_str or skills_str == ""):
                    skill_count = 0
                    score = 0


                elif skill_count >= 1 and skill_count <= 10:
                    score = 0.2 * skill_count
                else:
                    score = 2

                return [skill_count, score]

            def calculate_dl_skill_score(skills_str):
                skill_count = 0
                skill_count = skills_str.count(',') + 1

                score = 0
                if skill_count == 1 and (
                        'None' in skills_str or 'Not' in skills_str or 'No' in skills_str or 'not' in skills_str or skills_str == ""):
                    skill_count = 0
                    score = 0


                elif skill_count >= 1 and skill_count <= 10:
                    score = 0.25 * skill_count
                else:
                    score = 2

                return [skill_count, score]

            def calculate_MLOps_skill_score(skills_str):
                skill_count = 0
                skill_count = skills_str.count(',') + 1

                score = 0
                if skill_count == 1 and (
                        'None' in skills_str or 'Not' in skills_str or 'No' in skills_str or 'not' in skills_str or skills_str == ""):
                    skill_count = 0
                    score = 0


                elif skill_count >= 1 and skill_count <= 3:
                    score = 0.2 * skill_count
                else:
                    score = 0.6

                return [skill_count, score]

            def calculate_cloud_skill_score(skills_str):

                skill_count = 0
                skill_count = skills_str.count(',') + 1

                score = 0
                if skill_count == 1 and (
                        'None' in skills_str or 'Not' in skills_str or 'No' in skills_str or 'not' in skills_str or skills_str == ""):
                    skill_count = 0
                    score = 0


                elif skill_count >= 1 and skill_count <= 3:
                    score = 0.2 * skill_count
                else:
                    score = 0.6

                return [skill_count, score]

            def calculate_combine_skill_score(combine_count):

                if combine_count == 0:
                    score = 0

                elif combine_count >= 1 and combine_count <= 4:
                    score = 0.25 * combine_count
                else:
                    score = 1

                return score

            def calculate_project_score(count):

                if count == 0:
                    score = 0

                elif count >= 1 and count <= 10:
                    score = 0.1 * count
                else:
                    score = 1

                return score

            genAI_count = calculate_skill_score(str(genai_skills))
            ds_count = calculate_ds_skill_score(str(ds_skills))
            dl_count = calculate_dl_skill_score(str(dl_skills))
            mlops_count = calculate_MLOps_skill_score(str(mlops_skills))
            cloud_count = calculate_cloud_skill_score(str(cloud_skills))

            sum1 = dl_count[0] + ds_count[0]+genAI_count[0]
            # st.write(sum1)

            if genAI_count[0] == 0:
                genai_status = "No"
            else:
                genai_status = "Yes"

            if ds_count[0] == 0:
                ds_status = "No"
            else:
                ds_status = "Yes"

            if dl_count[0] == 0:
                dl_status = "No"
            else:
                dl_status = "Yes"

            if mlops_count[0] == 0:
                mlops_status = "No"
            else:
                mlops_status = "Yes"

            if cloud_count[0] == 0:
                cloud_status = "No"
            else:
                cloud_status = "Yes"

            # ds_list = ds_skills.split(',')
            # dl_list = dl_skills.split(',')
            # combine_skills = ds_list + dl_list
            # combine_count = len(combine_skills)

            combine_score = ds_count[1] + dl_count[1]

            project_score = calculate_project_score(int(count))

            total_score = combine_score + project_score + cloud_count[1] + genAI_count[1] + mlops_count[1]
            # st.write(total_score)

            rerun = False

            if total_score == 0 and not rerun:
                prompt5 = f"""
                                Given is the resume content of the candidate. Analyze the resume.

                                STEP 1:
                                    Your task is to search for the technical skills section or similar sections where candidate usually write their skills as comma separated.
                                    Along with this search for the professional summary or similar to summary where candidate gives small introduction about themselves. sometimes it is just a paragraph.
                                    Here, use your understanding that you need to avoid excluding project details or information about projects.
                                STEP 2:
                                    You have to exclude that whole section containing skills as comma separated, paragraph of summary or introduction of candidate and give everything other than this in proper format.
                                    Include name, Email ID and phone number.


                                Resume content:{text_content} 
                                """

                try:

                    try:
                        response_text5 = model.generate_content(prompt5, generation_config=genai.types.GenerationConfig(
                            candidate_count=1,
                            temperature=0.3)
                                                                )
                    except Exception as e:
                        if "Resource has been exhausted" in str(e):
                            print("Rate limit exceeded. Please try again later.")
                            time.sleep(60)
                            response_text5 = model.generate_content(prompt5, generation_config=genai.types.GenerationConfig(
                            candidate_count=1,
                            temperature=0.3)
                                                                )
                        else:
                            raise e

                    # st.write(response_text5.text)

                    prompt4 = f"""
                                Your task is to analyze the given resume. You have to think step by step as mentioned and give answer of the questions in just 1 to 2 words asked on the resume.

                                I am providing you example of both recommended candidate as well as not recommended candidate. Please go through it and understand how skills and other things are extracted.
                                I want you to follow the same logic while giving answers for the following questions. 

                                Example of recommended candidate = {Example1}

                                Example of not recommended candidate = {Example2}. Here Candidate has experience as an intern so we excluded everything.

                                In the above both examples, user search for powerbi as optional skill but both candidate haven't worked on it that's why additional skills are marked as 'No'.

                                Ensure accuracy while providing the answers. Avoid giving skills mentioned in the questions, used it only for your reference.

                                STEP 1:
                                    Question: Does candidate mentioned or used 'Python' or 'R' programming language? Provide programming langauge used only.

                                    If answer is 'Yes' then only move to next part else simply return 'No programming language'.

                                STEP 2:
                                    Question:  Does the candidate have at least 2 years of recent experience in NLP, Machine learning, Machine learning, Deep learning or GenAI ?

                                    If yes then only move to next question else provide 'Less experience'.

                                    Question: If 'Yes', does the candidate have used any of the NLP, Machine learning, Machine learning, Deep learning or GenAI skills during this recent experience? 

                                    Question: How many years of recent experience candidate have? Recent experience means last company or present company in which candidate is/was working. Consider present as Feb 2024.

                                STEP 3:
                                    For this step only consider resume content which talks about candidates actual work experience in Machine learning, deep learning, genai, mlops. Please focus only to answer skills related questions below based on such work experience/projects only. Don't extract skills which are mentioned explicitly under technical skills or primary skills or secondary skills and which are not part of sections related to work experience and project experience.
                                    For your understanding following are the examples of skills mapping.
                                    You should analyze the content of the resume and extract skills based on patterns and keywords commonly associated with these domains. 
                                    Avoid giving skills mentioned below, used it only for your reference.

                                    Machine learning skills : {DS_Skills}
                                    Deep Learning skills : {DL_Skills}
                                    MLOps skills : {MLOps_Skills} 
                                    GenAI skills : {GenAI_Skills}

                                    Your task here is to go through each project done by the candidate, extract relevant skills which are mentioned by the candidate in the project. 
                                    If you are not able to extract any skill strictly provide none.

                                    For following questions, Only extract projects/work experience/professional experience  related to machine learning, deep learning, mlops, genai ONLY if he/she 'developed' or 'build' something in the project. 

                                    Question: What are the skills mentioned in sections like technical skills or primary skills or secondary skills? 
                                    Please do not include this skills in the analysis.

                                    Question: Does the candidate has any work experience as an intern? What skills he used as an intern ?
                                    Please do not include this skills in the analysis.

                                    Question: What are the specific Machine learning skills candidate has utilized or developed while working on projects except as an intern?
                                    question: Are you sure ? Ensure accuracy while giving response and include only technical skills. 
                                    Question: What are the specific deep learning skills candidate has utilized or developed while working on projects except as an intern?
                                    question: Are you sure ? Ensure accuracy while giving response and include only technical skills.
                                    Question: What are the specific MLOps skills candidate has utilized or developed while working on projects?
                                    question: Are you sure ? Ensure accuracy while giving response and include only technical skills.
                                    Question: What are the specific cloud skills candidate has utilized or developed while working on projects?
                                    question: Are you sure ? Ensure accuracy while giving response and include only technical skills.
                                    Question: What are the specific GenAI skills candidate has utilized or developed while working on projects?
                                    question: Are you sure ? Ensure accuracy while giving response and include only technical skills.
                                    
                                    If {Optional_skills} is empty then provide additional skills as 'None'.
                                    check if candidate has mentioned or used skills {Optional_skills} in the resume, if yes then provide that skill as additional skills else return 'None'.
                                    Avoid extracting skills other than {Optional_skills} in additional skills.
                    
                                    Always verify your answer and check above mentioned conditions are met, then provide updated answers.

                                STEP 4:
                                    Question: Provide the machine learning/deep learning/genai related project names done by the candidate?
                                    Question: Provide the total number of  machine learning/deep learning/genai related project done by the candidate as mentioned above?

                                STEP 5:
                                    You are SME in Machine learning and genai and your task is to evaluate given resume based on step 1 to step 5. 
                                    If the candidate has the skills in Machine learning and deep learning in step 3 and at least 3 projects as calculated in step 4, recommend that candidate for job with detailed reasoning referring to the responses from step 1 to step 3. If the candidate is rejected give the proper reasoning.
                                    Use your understanding while giving recommendation.

                                STEP 6:
                                    Provide the name of the candidate. If you are unable to find name then provide None.
                                    Provide the Email ID of the candidate. If you are unable to find Email ID then provide None.
                                    Provide the Phone number of the candidate. If you are unable to find Phone number then provide None.

                                Please follow the below format (Use Plain text only, avoid using Bold format) for providing the response, Use the latest skills fetched for related asked question:

                                1. Programming Languages:
                                2. Recent Experience:
                                3. Machine learning skills:
                                4. Deep Learning skills:
                                5. MLOps skills:
                                6. Cloud skills:
                                7. GenAI skills:
                                8. List of projects:
                                9. Recommendation:
                                10. Name of the candidate:
                                11. count of projects:
                                12. additional skills:
                                13. email:
                                14. phone:


                                Resume content:{response_text5.text}
                                """

                    try:
                        response_text = model.generate_content(prompt4, generation_config=genai.types.GenerationConfig(
                        candidate_count=1,
                        temperature=0.3)
                                                        )
                    except Exception as e:
                        if "Resource has been exhausted" in str(e):
                            print("Rate limit exceeded. Please try again later.")
                            time.sleep(60)
                            response_text = model.generate_content(prompt4, generation_config=genai.types.GenerationConfig(
                                candidate_count=1,
                                temperature=0.3)
                                                        )
                        else:
                            raise e

                    # st.write(response_text.text)

                    language_pattern = r"\*{0,2}Programming Languages:\*{0,2}\s(.*?)(?=\d+\. |\*{0,2}Recent Experience:\*{0,2}|\n)"

                    duration_pattern = r"\*{0,2}Recent Experience:\*{0,2}\s(.*?)(?=\*{0,2}Machine learning skills:\*{0,2}|\n)"

                    ds_pattern = r"\*{0,2}Machine learning skills:\*{0,2}\s(.*?)(?=\*{0,2}Deep Learning skills:\*{0,2}|\n)"

                    dl_pattern = r"\*{0,2}Deep Learning skills:\*{0,2}\s(.*?)(?=\*{0,2}MLOps skills:\*{0,2}|\n)"

                    MLOps_pattern = r"\*{0,2}MLOps skills:\*{0,2}\s(.*?)(?=\*{0,2}Cloud skills:\*{0,2}|\n)"

                    Cloud_pattern = r"\*{0,2}Cloud skills:\*{0,2}\s(.*?)(?=\*{0,2}GenAI skills:\*{0,2}|\n)"

                    genai_pattern = r"\*{0,2}GenAI skills:\*{0,2}\s(.*?)(?=\*{0,2}List of projects:\*{0,2}|\n)"

                    projects_pattern = r"\*{0,2}List of projects:\*{0,2}\s(.*?)(?=\*{0,2}\d+\. Recommendation:\*{0,2}|\Z)"

                    recommendation_pattern = r"\*{0,2}Recommendation:\*{0,2}\s(.*?)(?=\*{0,2}Name of the candidate:\*{0,2}|\n)"

                    name_pattern = r"\*{0,2}Name of the candidate:\*{0,2}\s(.*?)(?=\*{0,2}count of projects:\*{0,2}|\n)"

                    count_pattern = r"\*{0,2}count of projects:\*{0,2}\s(.*?)(?=\*{0,2}additional skills:\*{0,2}|\n)"

                    additional_pattern = r"\*{0,2}additional skills:\*{0,2}\s(.*?)(?=\*{0,2}email:\*{0,2}|\n)"

                    email_pattern = r"\*{0,2}email:\*{0,2}\s(.*?)(?=\*{0,2}phone:\*{0,2}|\n)"
                    
                    phone_pattern = r"\*{0,2}phone:\*{0,2}\s(.*?)$"

                    # Extract text following each label using regular expressions

                    extract_text = lambda pattern, response: re.search(pattern, response, re.DOTALL).group(

                        1).strip() if re.search(pattern, response, re.DOTALL) else "Could not read"

                    language = extract_text(language_pattern, response_text.text)

                    duration = extract_text(duration_pattern, response_text.text)

                    ds_skills = extract_text(ds_pattern, response_text.text)

                    dl_skills = extract_text(dl_pattern, response_text.text)

                    mlops_skills = extract_text(MLOps_pattern, response_text.text)

                    cloud_skills = extract_text(Cloud_pattern, response_text.text)

                    genai_skills = extract_text(genai_pattern, response_text.text)

                    projects = extract_text(projects_pattern, response_text.text)
                    # st.write(projects)

                    recommendation = extract_text(recommendation_pattern, response_text.text)

                    additional = extract_text(additional_pattern, response_text.text)

                    name = extract_text(name_pattern, response_text.text)
                    
                    email = extract_text(email_pattern, response_text.text)

                    phone = extract_text(phone_pattern, response_text.text)
                    
                    # print(phone)
                    phone = phone.strip()
                    phone = phone.replace('+91',"")
                    phone = phone.replace('-',"")
                    phone = phone.replace('+',"").replace('(',"").replace(")","")
                    phone = phone.strip()
                    # print(phone)


                    if str(name).lower() in ['none', 'not provided', 'not mentioned', 'could not read', 'not available',
                                             'not mentioned in the resume']:
                        name = ' '.join(
                            re.sub(r'\[.*?\]', '', os.path.splitext(file_name)[0]).replace('+', ' ').replace('_',
                                                                                                             ' ').replace(
                                'Naukri', '').replace('PHENOM', '').split()[:2]).title()
                    else:
                        name = name.title()

                    count = extract_text(count_pattern, response_text.text)

                    if 'None' in str(count) or 'Not' in str(count) or 'No' in str(count) or 'not' in str(count) or str(
                            count) == "":
                        count = 0
                    else:
                        count = count

                    ds_list = ds_skills.split(',')
                    ds_list = [item.strip() for item in ds_list]
                    dl_list = dl_skills.split(',')
                    dl_list = [item.strip() for item in dl_list]
                    genai_list = genai_skills.split(',')
                    genai_list = [item.strip() for item in genai_list]

                    llm_list = ['large language models (LLMs)', 'LLM', 'GPT', 'GPT 3.5', 'GPT-3', 'BERT', 'ChatGPT',
                                'GPT-4', 'Generative AI', 'GenAI', 'Prompts', 'AzureOpenAI',
                                'langchain', 'Llamaindex framework', 'LLMs', 'Vector DB', 'Open-AI Da-Vinci LLM',
                                'Development of LLMs and Generative AI applications', "LLM's",
                                'PaLM', 'Gen AI', 'Prompt Engineering', 'Fine tune LLM using LoRA', 'QLoRA']

                    for item in llm_list:
                        if item in ds_list:
                            ds_list.remove(item)
                            if item not in genai_list:
                                genai_list.append(item)
                        if item in dl_list:
                            dl_list.remove(item)
                            if item not in genai_list:
                                genai_list.append(item)

                    ds_skills = ','.join(ds_list)
                    dl_skills = ','.join(dl_list)
                    genai_skills = ','.join(genai_list)
                    genai_skills = genai_skills.replace('None,', '')

                    def calculate_skill_score(skills_str):
                        skill_count = 0
                        skill_count = skills_str.count(',') + 1

                        score = 0
                        if skill_count == 1 and (
                                'None' in skills_str or 'Not' in skills_str or 'No' in skills_str or 'not' in skills_str or skills_str == ""):
                            skill_count = 0
                            score = 0


                        elif skill_count >= 1 and skill_count <= 8:
                            score = 0.25 * skill_count
                        else:
                            score = 2

                        return [skill_count, score]

                    def calculate_ds_skill_score(skills_str):
                        skill_count = 0
                        skill_count = skills_str.count(',') + 1

                        score = 0
                        if skill_count == 1 and (
                                'None' in skills_str or 'Not' in skills_str or 'No' in skills_str or 'not' in skills_str or skills_str == ""):
                            skill_count = 0
                            score = 0


                        elif skill_count >= 1 and skill_count <= 10:
                            score = 0.2 * skill_count
                        else:
                            score = 2

                        return [skill_count, score]

                    def calculate_dl_skill_score(skills_str):
                        skill_count = 0
                        skill_count = skills_str.count(',') + 1

                        score = 0
                        if skill_count == 1 and (
                                'None' in skills_str or 'Not' in skills_str or 'No' in skills_str or 'not' in skills_str or skills_str == ""):
                            skill_count = 0
                            score = 0


                        elif skill_count >= 1 and skill_count <= 10:
                            score = 0.2 * skill_count
                        else:
                            score = 2

                        return [skill_count, score]

                    def calculate_MLOps_skill_score(skills_str):
                        skill_count = 0
                        skill_count = skills_str.count(',') + 1

                        score = 0
                        if skill_count == 1 and (
                                'None' in skills_str or 'Not' in skills_str or 'No' in skills_str or 'not' in skills_str or skills_str == ""):
                            skill_count = 0
                            score = 0


                        elif skill_count >= 1 and skill_count <= 3:
                            score = 0.2 * skill_count
                        else:
                            score = 0.6

                        return [skill_count, score]

                    def calculate_cloud_skill_score(skills_str):

                        skill_count = 0
                        skill_count = skills_str.count(',') + 1

                        score = 0
                        if skill_count == 1 and (
                                'None' in skills_str or 'Not' in skills_str or 'No' in skills_str or 'not' in skills_str or skills_str == ""):
                            skill_count = 0
                            score = 0


                        elif skill_count >= 1 and skill_count <= 3:
                            score = 0.2 * skill_count
                        else:
                            score = 0.6

                        return [skill_count, score]

                    def calculate_combine_skill_score(combine_count):

                        if combine_count == 0:
                            score = 0

                        elif combine_count >= 1 and combine_count <= 4:
                            score = 0.25 * combine_count
                        else:
                            score = 1

                        return score

                    def calculate_project_score(count):

                        if count == 0:
                            score = 0

                        elif count >= 1 and count <= 10:
                            score = 0.1 * count
                        else:
                            score = 1

                        return score

                    genAI_count = calculate_skill_score(str(genai_skills))
                    ds_count = calculate_ds_skill_score(str(ds_skills))
                    dl_count = calculate_dl_skill_score(str(dl_skills))
                    mlops_count = calculate_MLOps_skill_score(str(mlops_skills))
                    cloud_count = calculate_cloud_skill_score(str(cloud_skills))

                    sum1 = dl_count[0] + ds_count[0]+genAI_count[0]
                    # st.write(sum1)

                    if genAI_count[0] == 0:
                        genai_status = "No"
                    else:
                        genai_status = "Yes"

                    if ds_count[0] == 0:
                        ds_status = "No"
                    else:
                        ds_status = "Yes"

                    if dl_count[0] == 0:
                        dl_status = "No"
                    else:
                        dl_status = "Yes"

                    if mlops_count[0] == 0:
                        mlops_status = "No"
                    else:
                        mlops_status = "Yes"

                    if cloud_count[0] == 0:
                        cloud_status = "No"
                    else:
                        cloud_status = "Yes"

                    # ds_list = ds_skills.split(',')
                    # dl_list = dl_skills.split(',')
                    # combine_skills = ds_list + dl_list
                    # combine_count = len(combine_skills)

                    combine_score = ds_count[1] + dl_count[1]

                    project_score = calculate_project_score(int(count))

                    total_score = combine_score + project_score + cloud_count[1] + genAI_count[1] + mlops_count[1]
                    # st.write(total_score)

                except Exception:
                    total_score = combine_score + project_score + cloud_count[1] + genAI_count[1] + mlops_count[1]

                rerun = True

            

            # if level == "Normal":
            #     threshold1 = 1.8
            #     threshold2 = 1
            #     threshold3 = 1.5
            #     genai_count_threshold1 = 2
            #     genai_count_threshold2 = 4
            #     combine_score_threshold = 1
            #     mlops_count_threshold = 2
            #     cloud_count_threshold = 0
 
            # else:
            #     threshold1 = 2.3
            #     threshold2 = 1.5
            #     threshold3 = 2
            #     genai_count_threshold1 = 4
            #     genai_count_threshold2 = 6
            #     combine_score_threshold = 2
            #     mlops_count_threshold = 3
            #     cloud_count_threshold = 1
 
 
 
            if selectfield == 'GenAI/DataScience':

                if total_score >= 1.8 and sum(
                        x == 'None' for x in [ds_skills, dl_skills, mlops_skills, cloud_skills, genai_skills]) <= 4:
                    if (genAI_count[0]) > 2:
                        new_recommend_genai = "Yes.Recommended for GenAI profile."
                        if (combine_score) >= 1:
                            new_recommend_ds = "Yes.Recommended for Data Science profile."
                        else:
                            new_recommend_ds = "No.Candidate doesn't have essential skills."
                        if (mlops_count[0]) >=2 and cloud_count[0] > 0:
                            mlops_recommend = "Yes.Recommended for MLOps profile."
                        else:
                            mlops_recommend = "No.Candidate doesn't have essential skills."
                    else:
                        new_recommend_ds = "Yes.Recommended for Data Science profile."
                        new_recommend_genai = "No.Candidate doesn't have essential skills."
                        if (mlops_count[0]) >=2 and cloud_count[0] > 0:
                            mlops_recommend = "Yes.Recommended for MLOps profile."
                        else:
                            mlops_recommend = "No.Candidate doesn't have essential skills."

                elif total_score >= 1 and sum(
                        x == 'None' for x in [ds_skills, dl_skills, mlops_skills, cloud_skills, genai_skills]) <= 2:
                    if (genAI_count[0]) > 2:
                        new_recommend_genai = "Yes.Recommended for GenAI profile."
                        if (combine_score) >= 1:
                            new_recommend_ds = "Yes.Recommended for Data Science profile."
                        else:
                            new_recommend_ds = "No.Candidate doesn't have essential skills."
                        if (mlops_count[0]) >=2 and cloud_count[0] > 0:
                            mlops_recommend = "Yes.Recommended for MLOps profile."
                        else:
                            mlops_recommend = "No.Candidate doesn't have essential skills."
                    else:
                        new_recommend_ds = "Yes.Recommended for Data Science profile."
                        new_recommend_genai = "No.Candidate doesn't have essential skills."
                        if (mlops_count[0]) >=2 and cloud_count[0] > 0:
                            mlops_recommend = "Yes.Recommended for MLOps profile."
                        else:
                            mlops_recommend = "No.Candidate doesn't have essential skills."

                elif sum1 >= 5 and total_score >= 1.5 and sum(
                        x == 'None' for x in [ds_skills, dl_skills, mlops_skills, cloud_skills, genai_skills]) <= 3:
                    if (genAI_count[0]) > 2:
                        new_recommend_genai = "Yes.Recommended for GenAI profile."
                        if (combine_score) >= 1:
                            new_recommend_ds = "Yes.Recommended for Data Science profile."
                        else:
                            new_recommend_ds = "No.Candidate doesn't have essential skills."
                        if (mlops_count[0]) >=2 and cloud_count[0] > 0:
                            mlops_recommend = "Yes.Recommended for MLOps profile."
                        else:
                            mlops_recommend = "No.Candidate doesn't have essential skills."
                    else:
                        new_recommend_ds = "Yes.Recommended for Data Science profile."
                        new_recommend_genai = "No.Candidate doesn't have essential skills."
                        if (mlops_count[0]) >=2 and cloud_count[0] > 0:
                            mlops_recommend = "Yes.Recommended for MLOps profile."
                        else:
                            mlops_recommend = "No.Candidate doesn't have essential skills."

                elif (genAI_count[0]) > 4:
                    new_recommend_genai = "Yes.Recommended for GenAI profile."
                    if (combine_score) >= 1:
                        new_recommend_ds = "Yes.Recommended for Data Science profile."
                    else:
                        new_recommend_ds = "No.Candidate doesn't have essential skills."
                    if (mlops_count[0]) >= 2 and cloud_count[0] > 0:
                        mlops_recommend = "Yes.Recommended for MLOps profile."
                    else:
                        mlops_recommend = "No.Candidate doesn't have essential skills."

                else:
                    new_recommend_ds = "No.Candidate doesn't have essential skills."
                    new_recommend_genai = "No.Candidate doesn't have essential skills."
                    if (mlops_count[0]) >= 2 and cloud_count[0] > 0:
                        mlops_recommend = "Yes.Recommended for MLOps profile."
                    else:
                        mlops_recommend = "No.Candidate doesn't have essential skills."

                prompt2 = f"""
                Your task is to generate positive summary if  {new_recommend_genai} =="Recommended for GenAI profile." or {new_recommend_ds} =="Recommended for Data Science profile." for the candidate based on the given skills.
                Your task is also to generate negative summary if {new_recommend_genai} == "No. Candidate doesn't have essential skills." or {new_recommend_ds} == "No. Candidate doesn't have essential skills." for the candidate based on the given skills if candidate has less or no skills or worked on less than 2 projects.

                provide in detail positive,negative summary in following format(Use Plain text only, avoid using Bold format) .

                For Example:

                 1. Positive Summary:
                 2. Negative Summary: 


                Machine learning skills = {ds_skills}
                Deep learning skills = {dl_skills}
                MLOps skills = {mlops_skills}
                Cloud skills = {cloud_skills}
                GenAI skills = {genai_skills}

                """

                # response_text1 = model.generate_content(prompt2, generation_config=genai.types.GenerationConfig(
                #     candidate_count=1,
                #     temperature=0.3)
                #                                         )

                # st.write(response_text1.text)

            elif selectfield == 'Machine Learning':

                if (ds_count[0]) > 6 and (int(count)) > 2:
                    new_recommend_ds = "Yes.Recommended for Machine Learning profile."
                    new_recommend_genai =""
                    if (mlops_count[0]) >= 2 and cloud_count[0] > 0:
                        mlops_recommend = "Yes.Recommended for MLOps profile."
                    else:
                        mlops_recommend = "No.Candidate doesn't have essential skills."

                else:
                    new_recommend_ds = "No.Candidate doesn't have essential skills."
                    new_recommend_genai = ""
                    if (mlops_count[0]) >= 2 and cloud_count[0] > 0:
                        mlops_recommend = "Yes.Recommended for MLOps profile."
                    else:
                        mlops_recommend = "No.Candidate doesn't have essential skills."

                prompt2 = f"""
                Your task is to generate positive summary if  {new_recommend_ds} =="Recommended for Machine Learning profile." for the candidate based on the given skills.
                Your task is also to generate negative summary if {new_recommend_ds} == "No. Candidate doesn't have essential skills." for the candidate based on the given skills if candidate has less or no skills or worked on less than 2 projects.

                provide in detail positive,negative summary in following format(Use Plain text only, avoid using Bold format) .

                For Example:

                 1. Positive Summary:
                 2. Negative Summary: 


                Machine learning skills = {ds_skills}

                """

                response_text1 = model.generate_content(prompt2, generation_config=genai.types.GenerationConfig(
                    candidate_count=1,
                    temperature=0.3)
                                                        )

                # st.write(response_text1.text)

            else:

                if (dl_count[0]) > 6 and (int(count)) > 2:
                    new_recommend_ds = "Yes.Recommended for Deep Learning profile."
                    new_recommend_genai = ""
                    if (mlops_count[0]) >= 2 and cloud_count[0] > 0:
                        mlops_recommend = "Yes.Recommended for MLOps profile."
                    else:
                        mlops_recommend = "No.Candidate doesn't have essential skills."

                else:
                    new_recommend_ds = "No.Candidate doesn't have essential skills."
                    new_recommend_genai = ""
                    if (mlops_count[0]) >= 2 and cloud_count[0] > 0:
                        mlops_recommend = "Yes.Recommended for MLOps profile."
                    else:
                        mlops_recommend = "No.Candidate doesn't have essential skills."

                prompt2 = f"""
                Your task is to generate positive summary if  {new_recommend_ds} =="Recommended for Deep Learning profile." for the candidate based on the given skills.
                Your task is also to generate negative summary if {new_recommend_ds} == "No. Candidate doesn't have essential skills." for the candidate based on the given skills if candidate has less or no skills or worked on less than 2 projects.

                provide in detail positive,negative summary in following format(Use Plain text only, avoid using Bold format) .

                For Example:

                 1. Positive Summary:
                 2. Negative Summary: 

                Deep learning skills = {dl_skills}

                """

                response_text1 = model.generate_content(prompt2, generation_config=genai.types.GenerationConfig(
                    candidate_count=1,
                    temperature=0.3)
                                                        )

                # st.write(response_text1.text)

            positive_pattern = r"\*{0,2}Positive Summary:\*{0,2}\s(.*?)(?=\*{0,2}Negative Summary:\*{0,2}|\Z)"

            negative_pattern = r"\*{0,2}Negative Summary:\*{0,2}\s(.*?)$"

            extract_text = lambda pattern, response: re.search(pattern, response, re.DOTALL).group(

                1).strip() if re.search(pattern, response, re.DOTALL) else "Could not read"

            # positive = extract_text(positive_pattern, response_text1.text)
            positive="Candidate has essential skills"
            # st.write(positive)

            # negative = extract_text(negative_pattern, response_text1.text)
            negative="Candidate doesn't have essential skills"
            # st.write(negative)

            if 'No.' in new_recommend_ds or 'No.' in new_recommend_genai:
                recommendation = negative

            else:
                recommendation = positive
            # st.write(recommendation)

            # if 'None' in additional:
            #     additional = ''
            additional_list = additional.split(',')

            for i in range(len(additional_list)):
                additional_list[i] = additional_list[i].lower()

            for i in range(len(Optional_skills)):
                Optional_skills[i] = Optional_skills[i].lower()

            if ('Recommended' in new_recommend_ds or 'Recommended' in new_recommend_genai) and any(
                    skill in additional_list for skill in Optional_skills):
                alter = 'Yes'
                Optional_skills = ','.join(Optional_skills).capitalize()
                if 'Recommended' in new_recommend_ds:
                    new_recommend_ds = new_recommend_ds + f" Also recommended for {Optional_skills}"
                if 'Recommended' in new_recommend_genai:
                    new_recommend_genai = new_recommend_genai + f" Also recommended for {Optional_skills}"
            else:
                alter = 'No'
                additional = ''
            # alter = 'No'

            # if total_score == 0:
            #     language = duration = ds_skills = ds_status = dl_skills = dl_status = mlops_skills = mlops_status = cloud_skills = cloud_status = genai_skills = genai_status = additional = projects = recommendation = alter = ""
            #     new_recommend = "Need Human Intervention"

            if 'Recommended' in new_recommend_ds or 'Recommended' in new_recommend_genai:
                prompt99 = f"""
                        Consider you are an interviewer, your task is to ask potential questions to be asked from the candidate whose resume is {text_content}.

                        Generate a set of 10 comprehensive interview questions based on resume provided. Questions must be related to Data Science skills and/or GenAI skills for a candidate applying for a role either of Data Science and/or Generative AI. These questions should assess the candidate's understanding, experience, and technical skills across candidate's major skills:

                        1. Projects and Technical Aspects (5 Questions): Questions should delve into the candidate's practical experience and the technical challenges they've navigated in their projects. Generate case study based technical questions as an interviewer based on projects/work experiences done by the candidates. The dept of Questions should increase complexity wise from question 1 to 5. Each question should also have a section to ask specific technical package/skills used related to python language.

                        2. Python Specific Technical Skills (3 Questions): Ask conceptual programming questions to check candidate's programming ability and problem solving power in python.

                        3. Additional questions (2 questions): Ask questions based on other skills/cloud if mentioned in resume. for example, "How would you deploy machine learning model on AWS?"

                        Provide response in the tabular format. First row should contain only one column where name of the candidate should be mentioned.
                        Next rows should contain 2 columns , first column for questions with question number like Q1,Q2,etc, second column for score and it should be empty always.

                        Exclude the provided example from the analysis, Use the example only for the formatting.

                        For example:
                        "
                        | Name of the candidate |
                        | --- |
                        | DEBASIS SAMAL |
                        | Questions | Score |
                        | --- | --- |
                        | Q1.Walk me through the process of building the movie recommender system that you build. How did you train the NLP vectorizer model and predict similar movies based on cosine similarity? |  |
                        | Q2.Describe your experience in developing the multilingual chatbot using Generative AI (OpenAI) on AWS SageMaker. How did you handle the challenges of multilingual processing and knowledge extraction from PDF documents? |  |
                        | --- | --- |
                        | ... |  |
                        "


                        Name of the candidate = {name}
                        Machine learning skills = {ds_skills}
                        Deep learning skills = {dl_skills}
                        MLOps skills = {mlops_skills}
                        Cloud skills = {cloud_skills}
                        GenAI skills = {genai_skills}

                        """
                # questions_text = model.generate_content(prompt99, generation_config=genai.types.GenerationConfig(
                #     candidate_count=1,
                #     temperature=0.3)
                #                                         )
                
                # st.write(questions_text.text)
                # st.write(type(questions_text.text))
                
                # questions_dict[file_name]= questions_text.text
                questions_dict[file_name]= "HI"

                # # Parse response and extract data
                # lines = questions_text.text.strip().split('\n')
                # candidate_name = lines[2].strip().split("|")[1].strip()
                # questions_and_scores = [line.strip() for line in lines[4:]]

                # # Create a new Document
                # doc = Document()

                # sections = doc.sections
                # for section in sections:
                #     section.left_margin = Pt(35)  # Adjust left margin as needed
                #     section.right_margin = Pt(35)  # Adjust right margin as needed
                #     section.top_margin = Pt(35)  # Adjust top margin as needed
                #     section.bottom_margin = Pt(35)  # Adjust bottom margin as needed

                # timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                # # doc.add_paragraph(f"Report Generated At: {timestamp}",style = 'Heading 1')

                # timestamp_paragraph = doc.add_paragraph(f"Report Generated At: {timestamp}")
                # timestamp_run = timestamp_paragraph.runs[0]
                # timestamp_run.bold = True
                # timestamp_run.font.size = Pt(11)  # Adjust the font size as needed

                # # Add candidate name
                # doc.add_heading('Candidate Name:', level=1)
                # doc.add_paragraph(candidate_name)

                # # Add table
                # table = doc.add_table(rows=1, cols=2)
                # table.style = 'Table Grid'

                # table.alignment = WD_TABLE_ALIGNMENT.LEFT

                # hdr_cells = table.rows[0].cells
                # hdr_cells[0].text = 'Questions'
                # hdr_cells[1].text = 'Score'

                # # Make column names bold
                # for cell in hdr_cells:
                #     for paragraph in cell.paragraphs:
                #         for run in paragraph.runs:
                #             run.bold = True

                # # Populate table
                # for line in questions_and_scores:
                #     parts = line.split('|')
                #     if len(parts) > 1 and parts[1].strip() != '---':  # Check if the line is not just '---'
                #         question = parts[1].strip()
                #         score = parts[2].strip() if len(
                #             parts) > 2 else ""  # Handle the case where score is not provided
                #         row_cells = table.add_row().cells
                #         row_cells[0].text = question
                #         row_cells[1].text = score

                #         # Add space after each question
                #         row_cells[0].paragraphs[0].runs[-1].add_break()

                #         # for cell in row_cells:
                #         #     for paragraph in cell.paragraphs:
                #         #         paragraph.space_after = Pt(10)

                # # Adjust row height to fit content
                # for row in table.rows:
                #     for cell in row.cells:
                #         for paragraph in cell.paragraphs:
                #             paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Set alignment to left
                #             paragraph.space_after = Pt(12)  # Add space after each paragraph (adjust as needed)
                #         cell.height = Pt(60)  # Adjust row height to fit content

                # # Add a border around the entire document
                # for paragraph in doc.paragraphs:
                #     paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                #     paragraph.space_after = Pt(12)
                #     for run in paragraph.runs:
                #         run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black for better visibility
                #         run.font.name = 'Cambria'
                #         run.font.size = Pt(11)

                # # Add space after the headings "Questions" and "Score"
                # for cell in table.rows[0].cells:
                #     cell.paragraphs[0].runs[-1].add_break()

                # for cell in table.columns[1].cells:
                #     cell.width = Pt(30)  # Adjust cell width as needed

                # for cell in table.columns[0].cells:
                #     cell.width = Pt(700)

                # table.border_collapse = True

                # # new_folder = os.path.join(resume_folder, 'Generated_Docx')
                # # os.makedirs(new_folder, exist_ok=True)
                # # docx_filename = os.path.join(new_folder, f'{name}_reports.docx')
                # # doc.save(docx_filename)
                # # download_link = f"file:///{docx_filename}"
                # # docx_link = f'=HYPERLINK("{download_link}", "Open Report")'
                docx_link = ''
            else:
                docx_link = ''

            # prompt1 = f"""
            # Based on the following data,return 'Yes' or 'No' to indicate whether the profile is suitable for a data science-related position.

            # Following are the conditions for positive recommendation:

            #     1.Candidate must have any of the 'Python' or 'R' in 'Skills' or 'Programming Languages'.

            #     2.Candidate must have either of 'GenAI Related' or 'Data Science Related' skills.

            # Indicate 'Yes' only if the candidate satisfies two conditions above else 'No'.   
            # Indicate 'No' if only 'Machine Learning' is present in 'Data Science Related' skills and has 'None' in 'GenAI Related' skills.

            # Also, provide a short reason for your recommendation.
            # If 'GenAI Related' skills present then recommend candidate for GenAI profile.
            # If only 'Data Science Related' skills present then recommend candidate for Data Science profile.
            # Avoid using phrases like 'Answer:'.
            # If Recommendation is 'No', provide reason as 'Candidate doesn't have essential skills'.

            # Please focus exclude the provided example from the analysis.

            # For example:
            # Yes. Recommended for Data Science profile.
            # Reason: Given the strong foundation and candidate demonstrates expertise in Data Science skills, the candidate is well-suited for positions related to Data Science.
            # The candidate possesses proficiency in the Programming Languages, specifically in 'Python.' 
            # Notably, the candidate has cloud skills of AWS,Azure.

            # Focus solely on the relevant information requested and avoid additional details.

            # Summary: {response_text1.text}

            # """

            # recommendation = openai.Completion.create(

            #     engine="text-davinci-003",

            #     prompt=prompt1,

            #     max_tokens=500,

            #     temperature=0

            # )

            # recommendation_text = recommendation.choices[0].text.strip()

            # recommendation_text = model.generate_content(prompt1, generation_config=genai.types.GenerationConfig(
            #     candidate_count=1,
            #     top_p=0.6,
            #     top_k=5,
            #     temperature=0)
            #                                              )

            # recommendation_text=process1(response_text,prompt1)

            # role, experience,datascience_skills[1], recommendation,pdf_link,docx_link

            pdf_link =''
            data.append([file_name, name, language, duration, ds_skills, ds_status, dl_skills, dl_status, mlops_skills,
                         mlops_status, cloud_skills, cloud_status, genai_skills, genai_status, additional, alter,
                         projects, recommendation, new_recommend_genai,new_recommend_ds,mlops_recommend, genAI_count[0], ds_count[0], dl_count[0],
                         mlops_count[0], cloud_count[0], count,
                         combine_score, project_score, cloud_count[1], genAI_count[1], mlops_count[1], total_score,
                         pdf_link, docx_link, email, phone])

            # Print statements to check progress
            # st.write(data)
            print(f"Processed {file_name}")
            # time.sleep(9)
            # Emit progress update
            progress = (index + 1) / total_resumes * 100    
            # print(progress)    
            socketio.emit('progress', {'progress': progress})
            index = index + 1
            

        except Exception as e:

            # st.error(f"Error processing summary for {file_name}: {e}")
            print(f"Error processing summary for {file_name}: {e}")

            # Emit progress update
            progress = (index + 1) / total_resumes * 100    
            # print(progress)    
            socketio.emit('progress', {'progress': progress})
            index = index + 1

            data.append(
                [file_name, "Error in processing", "Error in processing", "Error in processing", "Error in processing",
                 "Error in processing", "Error in processing", "Error in processing",
                 "Error in processing","Error in processing","Error in processing",
                 "Error in processing", "Error in processing", "Error in processing", "Error in processing",
                 "Error in processing", "Error in processing","Error in processing",
                 "Error in processing", "Error in processing", "Error in processing", "Error in processing",
                 "Error in processing", "Error in processing","Error in processing",
                 "Error in processing", "Error in processing", "Error in processing", "Error in processing",
                 "Error in processing", "Error in processing", "Error in processing",
                 "Error in processing", "Error in processing", 0, "Error in processing", "Error in processing"])

    return data,questions_dict

def dataFr(selectfield,Optional_skills,resume_dict,GOOGLE_API_KEY):

    # resumes_dict = load_resumes_as_dict2(resume_folder)
    # # st.write(resumes_dict)
    summary_data,questions_dict = generate_resume_summary(resume_dict, Optional_skills,DS_Skills,DL_Skills,MLOps_Skills,GenAI_Skills,selectfield,GOOGLE_API_KEY)

    # batch_size = 10
    # files = os.listdir(resume_folder)
    # num_filess = len(files)
    # num_batches = (num_filess + batch_size - 1) // batch_size

    # for i in range(num_batches):
    #     start_i = i * batch_size
    #     end_i = min((i + 1) * batch_size, num_filess)
    #     batch_files = files[start_i:end_i]
    #     resumes_dict = load_resumes_as_dict2(resume_folder, batch_files)
    #     # st.write(resumes_dict)
    #     summary_data = generate_resume_summary(resumes_dict, Optional_skills,DS_Skills,DL_Skills,MLOps_Skills,GenAI_Skills,resume_folder,selectfield)
        # summary_data = generate_resume_summary(resumes_dict, Optional_skills, DS_Skills, DL_Skills, MLOps_Skills,GenAI_Skills, resume_folder,user_score,weight_ml,weight_dl,weight_genai,weight_mlops,weight_cloud)

    df_summaries = pd.DataFrame(summary_data,
                                      columns=['File Name','Name','Programming Language','Recent Experience','Machine learning','ML Status',
                                               'Deep Learning','DL Status','MLOps','MLops Status','Cloud','Cloud Status','GenAI','GenAI Status',
                                               'Additional Skills','Alternate Recommendation','Projects','Candidate Summary','GenAI Recommendation','DS Recommendation','MLOps Recommendation','GenAI Count','ML Count',
                                               'DL Count','MLops Count','Cloud Count','Project Count','ML DL Score','Project Score','Cloud Score','GenAI Score',
                                               'MLOps Score','Total Score', 'View PDF','View Docx', 'Email', 'Phone'])

    df_summaries['Total Score'] = pd.to_numeric(df_summaries['Total Score'], errors='coerce')

    df_summaries = df_summaries.sort_values(by=['Total Score'], ascending=False)
    # df_summaries['Additional Skills'] = ''
    # df_summaries['Alternate Recommendation'] = 'No'  # Default to 'No'
    #
    # for i in Optional_skills:
    #     mask = df_summaries['Skills'].str.contains(i, case=False)
    #     recommendation_mask = df_summaries['Recommendation'].str.startswith('Yes', na=False)
    #
    #     # Check if the skill is not already present in 'Additional Skills'
    #     not_present_mask = ~df_summaries['Additional Skills'].str.contains(i, case=False) | df_summaries[
    #         'Additional Skills'].isna()
    #
    #     # Update 'Alternate Recommendation' based on conditions
    #     df_summaries.loc[mask & recommendation_mask & not_present_mask, 'Alternate Recommendation'] = 'Yes'
    #
    #     # Append the skill only if it's not already present
    #     df_summaries.loc[mask & recommendation_mask & not_present_mask, 'Additional Skills'] = df_summaries.loc[
    #         mask & recommendation_mask & not_present_mask, 'Additional Skills'].apply(
    #         lambda x: x + ',' + i if pd.notna(x) and x != '' else i)
    #
    # # Update 'Recommendation' to 'No' if 'Additional Skills' is empty
    # df_summaries.loc[df_summaries['Additional Skills'].eq('') | df_summaries[
    #     'Additional Skills'].isna(), 'Alternate Recommendation'] = 'No'

    return  df_summaries,questions_dict

# Running app
if __name__ == '__main__':
    socketio.run(app,host='0.0.0.0', port=3001, debug=True)
    # socketio.run(app)